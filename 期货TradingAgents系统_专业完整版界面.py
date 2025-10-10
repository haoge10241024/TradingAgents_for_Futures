#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
å•†å“æœŸè´§ Trading Agentsç³»ç»ŸStreamlitç•Œé¢
ä¿®å¤ç¼–ç é—®é¢˜çš„åŒæ—¶ä¿æŒå®Œæ•´åŠŸèƒ½å’Œä¸“ä¸šæ€§

åŠŸèƒ½ç‰¹ç‚¹ï¼š
1. å®Œæ•´çš„æ•°æ®ç®¡ç†å’Œæ›´æ–°åŠŸèƒ½
2. ä¸“ä¸šçš„åˆ†æé…ç½®å’Œæ‰§è¡Œ
3. å¤šç©ºè¾©è®ºé£æ§å†³ç­–ï¼ˆå«äº¤æ˜“å‘˜ç¯èŠ‚ï¼‰
4. WordæŠ¥å‘Šç”Ÿæˆå’Œå¯¼å‡º
5. æ‰€æœ‰åŸæœ‰çš„ä¸“ä¸šåŠŸèƒ½

ä½œè€…: 7haoge  
ç‰ˆæœ¬: å®Œæ•´å¯ç”¨ç‰ˆV1
åˆ›å»ºæ—¶é—´: 2025-01-19
è”ç³»æ–¹å¼: 953534947@qq.com
"""

import streamlit as st
import pandas as pd
import json
import asyncio
import time
import io
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Any, Set
import plotly.express as px
import plotly.graph_objects as go

# å›¾è¡¨ä¸‹è½½åŠŸèƒ½å·²ç¦ç”¨ï¼ˆé¿å…ç³»ç»Ÿrunningé—®é¢˜ï¼‰
CHART_DOWNLOAD_AVAILABLE = False
import subprocess
import sys
import threading
import os

# Wordæ–‡æ¡£ç”Ÿæˆç›¸å…³
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("âš ï¸ æœªå®‰è£…python-docxåº“ï¼ŒWordæŠ¥å‘ŠåŠŸèƒ½ä¸å¯ç”¨ã€‚è¯·è¿è¡Œ: pip install python-docx")

# ============================================================================
# å·¥å…·å‡½æ•°
# ============================================================================

def convert_confidence_to_level(confidence_score):
    """å°†æ•°å€¼ä¿¡å¿ƒåº¦è½¬æ¢ä¸ºç›´è§‚ç­‰çº§"""
    try:
        score = float(confidence_score)
        if score >= 0.8:
            return "é«˜ä¿¡å¿ƒ", "ğŸ’ª", "success"
        elif score >= 0.6:
            return "ä¸­ç­‰ä¿¡å¿ƒ", "ğŸ‘", "warning" 
        elif score >= 0.4:
            return "ä½ä¿¡å¿ƒ", "ğŸ¤”", "info"
        else:
            return "è°¨æ…è§‚æœ›", "âš ï¸", "error"
    except:
        return "æœªçŸ¥", "â“", "info"

def check_commodity_local_data(commodity: str, data_status: Dict) -> Dict[str, bool]:
    """æ£€æŸ¥æŒ‡å®šå“ç§åœ¨å„ä¸ªæ¨¡å—ä¸­çš„æœ¬åœ°æ•°æ®æƒ…å†µ"""
    module_mapping = {
        "technical": "æŠ€æœ¯é¢åˆ†æ",
        "basis": "åŸºå·®åˆ†æ",
        "inventory": "åº“å­˜æ•°æ®", 
        "positioning": "æŒä»“å¸­ä½",
        "term_structure": "æœŸé™ç»“æ„",
        "receipt": "ä»“å•æ•°æ®"
    }
    
    result = {}
    
    # ç›´æ¥ä½¿ç”¨æ–‡ä»¶ç³»ç»Ÿæ£€æŸ¥ï¼Œæ›´å¯é 
    for module_key, module_name in module_mapping.items():
        # ä¼˜å…ˆä½¿ç”¨ç›´æ¥æ£€æŸ¥æ–¹æ³•
        has_data = check_commodity_data_direct(commodity, module_key)
        result[module_key] = has_data
    
    return result

def check_specific_commodity_data(commodity: str, module_key: str) -> bool:
    """æ£€æŸ¥ç‰¹å®šå“ç§åœ¨æŒ‡å®šæ¨¡å—ä¸­æ˜¯å¦æœ‰æ•°æ®"""
    try:
        # æ£€æŸ¥æ•°æ®ç®¡ç†å™¨ä¸­æ˜¯å¦æœ‰è¯¥å“ç§çš„æ•°æ®
        if hasattr(st.session_state, 'data_manager'):
            data_manager = st.session_state.data_manager
            
            # æ ¹æ®æ¨¡å—ç±»å‹æ£€æŸ¥ä¸åŒçš„æ•°æ®æºï¼ˆä½¿ç”¨æ­£ç¡®çš„é…ç½®é”®åï¼‰
            if module_key == "technical":
                # æ£€æŸ¥æŠ€æœ¯åˆ†ææ•°æ®
                config = data_manager.modules_config.get("technical_analysis", {})
                tech_path = config.get("path")
                if tech_path and tech_path.exists():
                    commodity_path = tech_path / commodity
                    if commodity_path.exists():
                        # æ£€æŸ¥æ˜¯å¦æœ‰æ•°æ®æ–‡ä»¶
                        data_file = commodity_path / config.get("data_file", "ohlc_data.csv")
                        return data_file.exists() and data_file.stat().st_size > 0
            
            elif module_key == "basis":
                # æ£€æŸ¥åŸºå·®åˆ†ææ•°æ®
                config = data_manager.modules_config.get("basis", {})
                basis_path = config.get("path")
                if basis_path and basis_path.exists():
                    commodity_path = basis_path / commodity
                    if commodity_path.exists():
                        # æ£€æŸ¥æ˜¯å¦æœ‰æ•°æ®æ–‡ä»¶
                        data_file = commodity_path / config.get("data_file", "basis_data.csv")
                        return data_file.exists() and data_file.stat().st_size > 0
            
            elif module_key == "inventory":
                # æ£€æŸ¥åº“å­˜æ•°æ®
                config = data_manager.modules_config.get("inventory", {})
                inv_path = config.get("path")
                if inv_path and inv_path.exists():
                    commodity_path = inv_path / commodity
                    if commodity_path.exists():
                        # æ£€æŸ¥æ˜¯å¦æœ‰æ•°æ®æ–‡ä»¶
                        data_file = commodity_path / config.get("data_file", "inventory.csv")
                        return data_file.exists() and data_file.stat().st_size > 0
            
            elif module_key == "positioning":
                # æ£€æŸ¥æŒä»“æ•°æ®
                config = data_manager.modules_config.get("positioning", {})
                pos_path = config.get("path")
                if pos_path and pos_path.exists():
                    commodity_path = pos_path / commodity
                    if commodity_path.exists():
                        # æ£€æŸ¥æ˜¯å¦æœ‰æ•°æ®æ–‡ä»¶
                        data_file = commodity_path / config.get("data_file", "long_position_ranking.csv")
                        return data_file.exists() and data_file.stat().st_size > 0
            
            elif module_key == "term_structure":
                # æ£€æŸ¥æœŸé™ç»“æ„æ•°æ®
                config = data_manager.modules_config.get("term_structure", {})
                term_path = config.get("path")
                if term_path and term_path.exists():
                    commodity_path = term_path / commodity
                    if commodity_path.exists():
                        # æ£€æŸ¥æ˜¯å¦æœ‰æ•°æ®æ–‡ä»¶
                        data_file = commodity_path / config.get("data_file", "term_structure.csv")
                        return data_file.exists() and data_file.stat().st_size > 0
            
            elif module_key == "receipt":
                # æ£€æŸ¥ä»“å•æ•°æ®
                config = data_manager.modules_config.get("receipt", {})
                receipt_path = config.get("path")
                if receipt_path and receipt_path.exists():
                    commodity_path = receipt_path / commodity
                    if commodity_path.exists():
                        # æ£€æŸ¥æ˜¯å¦æœ‰æ•°æ®æ–‡ä»¶
                        data_file = commodity_path / config.get("data_file", "receipt.csv")
                        return data_file.exists() and data_file.stat().st_size > 0
        else:
            # å¦‚æœæ²¡æœ‰data_managerï¼Œç›´æ¥æ£€æŸ¥ç¡¬ç¼–ç è·¯å¾„
            return check_commodity_data_direct(commodity, module_key)
        
        return False
        
    except Exception as e:
        st.error(f"æ£€æŸ¥å“ç§æ•°æ®æ—¶å‡ºé”™: {e}")
        # å‡ºé”™æ—¶ä¹Ÿå°è¯•ç›´æ¥æ£€æŸ¥
        return check_commodity_data_direct(commodity, module_key)

def check_commodity_data_direct(commodity: str, module_key: str) -> bool:
    """ç›´æ¥æ£€æŸ¥å“ç§æ•°æ®ï¼ˆä¸ä¾èµ–session_stateï¼‰"""
    try:
        from pathlib import Path
        
        # ç¡¬ç¼–ç æ•°æ®åº“è·¯å¾„
        database_root = Path("qihuo/database")
        
        # æ¨¡å—è·¯å¾„æ˜ å°„
        module_paths = {
            "technical": database_root / "technical_analysis",
            "basis": database_root / "basis", 
            "inventory": database_root / "inventory",
            "positioning": database_root / "positioning",
            "term_structure": database_root / "term_structure",
            "receipt": database_root / "receipt"
        }
        
        # æ•°æ®æ–‡ä»¶æ˜ å°„
        data_files = {
            "technical": "ohlc_data.csv",
            "basis": "basis_data.csv",
            "inventory": "inventory.csv", 
            "positioning": "long_position_ranking.csv",
            "term_structure": "term_structure.csv",
            "receipt": "receipt.csv"
        }
        
        if module_key in module_paths:
            module_path = module_paths[module_key]
            if module_path.exists():
                commodity_path = module_path / commodity
                if commodity_path.exists():
                    data_file = commodity_path / data_files[module_key]
                    return data_file.exists() and data_file.stat().st_size > 0
        
        return False
        
    except Exception as e:
        st.error(f"ç›´æ¥æ£€æŸ¥å“ç§æ•°æ®æ—¶å‡ºé”™: {e}")
        return False

def get_market_data_info(commodity: str) -> Dict:
    """è·å–çœŸå®çš„å¸‚åœºæ•°æ®ä¿¡æ¯"""
    try:
        # è·å–å½“å‰ä»·æ ¼
        current_price = get_commodity_current_price(commodity)
        
        # è·å–ä»·æ ¼åŒºé—´
        price_range = futures_price_provider.get_price_range(commodity, 30)
        
        # è®¡ç®—ä»·æ ¼ä½ç½®
        avg_price = price_range.get('avg', current_price)
        if current_price > avg_price * 1.02:
            price_position = "åé«˜"
        elif current_price < avg_price * 0.98:
            price_position = "åä½"  
        else:
            price_position = "ä¸­æ€§"
            
        return {
            "current_price": current_price,
            "price_range": price_range,
            "price_position": price_position,
            "data_quality": "è‰¯å¥½" if current_price > 0 else "å¼‚å¸¸"
        }
    except Exception as e:
        st.error(f"è·å–{commodity}å¸‚åœºæ•°æ®å¤±è´¥: {e}")
        return {
            "current_price": 0,
            "price_range": {"low": 0, "high": 0, "avg": 0},
            "price_position": "æœªçŸ¥",
            "data_quality": "å¼‚å¸¸"
        }

# ============================================================================
# å¯¼å…¥ç³»ç»Ÿæ¨¡å—
# ============================================================================
try:
    from æœŸè´§TradingAgentsç³»ç»Ÿ_ç¬¬ä¸‰é˜¶æ®µå®Œæ•´ç‰ˆ import CompleteFuturesTradingExecution
    from æœŸè´§TradingAgentsç³»ç»Ÿ_åŸºç¡€æ¶æ„ import FuturesTradingAgentsConfig, FuturesAnalysisIntegrator
    from æœŸè´§TradingAgentsç³»ç»Ÿ_å·¥å…·æ¨¡å— import TimeUtils
    # å¯¼å…¥ä»·æ ¼æ•°æ®è·å–å™¨
    from ä»·æ ¼æ•°æ®è·å–å™¨ import get_commodity_current_price, futures_price_provider
except ImportError as e:
    st.error(f"å¯¼å…¥ç³»ç»Ÿæ¨¡å—å¤±è´¥: {e}")
    st.stop()

# ============================================================================
# é¡µé¢é…ç½®
# ============================================================================

st.set_page_config(
    page_title="å•†å“æœŸè´§ Trading Agentsç³»ç»Ÿ",
    page_icon="ğŸ’¹",
    layout="wide",
    initial_sidebar_state="expanded"
)

# è‡ªå®šä¹‰CSSæ ·å¼
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        color: #FF6B6B;
        text-align: center;
        margin-bottom: 2rem;
        background: linear-gradient(90deg, #FF6B6B, #4ECDC4);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .stage-header {
        font-size: 1.5rem;
        color: #2E86AB;
        margin: 1rem 0;
        padding: 0.5rem;
        border-left: 4px solid #2E86AB;
        background-color: #f0f8ff;
    }
    .analysis-card {
        background-color: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border-left: 5px solid #4ECDC4;
    }
    .debate-card {
        background-color: #fff8e1;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border-left: 5px solid #ff9800;
    }
    .risk-card {
        background-color: #f3e5f5;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border-left: 5px solid #9c27b0;
    }
    .decision-card {
        background-color: #e8f5e8;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin: 1rem 0;
        border-left: 5px solid #4caf50;
    }
    .metric-container {
        display: flex;
        justify-content: space-around;
        margin: 1rem 0;
    }
    .metric-item {
        text-align: center;
        padding: 1rem;
        background-color: white;
        border-radius: 8px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    .progress-bar {
        background-color: #e0e0e0;
        border-radius: 10px;
        padding: 3px;
        margin: 10px 0;
    }
    .progress-fill {
        background: linear-gradient(90deg, #4ECDC4, #44A08D);
        height: 20px;
        border-radius: 7px;
        transition: width 0.3s ease;
    }
    .page-selector {
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 10px;
        margin: 15px 0;
        border: 2px solid #e9ecef;
    }
    .analyst-card {
        background: white;
        padding: 20px;
        border-radius: 12px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin: 15px 0;
        border-left: 5px solid #667eea;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .analyst-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 15px rgba(0,0,0,0.15);
    }
    .content-section {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #007bff;
        margin: 15px 0;
        line-height: 1.6;
    }
    .citation-section {
        background-color: #e7f3ff;
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #bee5eb;
        margin: 15px 0;
    }
    .navigation-hint {
        background: linear-gradient(45deg, #f39c12, #e67e22);
        color: white;
        padding: 15px;
        border-radius: 8px;
        text-align: center;
        margin: 20px 0;
        font-size: 1.1em;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 5px;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: white;
        border-radius: 8px;
        font-weight: 500;
        padding: 8px 16px;
        border: 1px solid #e9ecef;
        transition: all 0.3s ease;
    }
    .stTabs [data-baseweb="tab"]:hover {
        background-color: #e7f3ff;
        border-color: #007bff;
        transform: translateY(-1px);
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #007bff 0%, #0056b3 100%) !important;
        color: white !important;
        border-color: #0056b3 !important;
        box-shadow: 0 4px 8px rgba(0,123,255,0.3);
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# æ•°æ®ç®¡ç†å™¨ï¼ˆä¿æŒå®Œæ•´åŠŸèƒ½ï¼‰
# ============================================================================

class StreamlitDataManager:
    """Streamlitæ•°æ®ç®¡ç†å™¨"""
    
    def __init__(self):
        self.database_root = Path("D:/Cursor/cursoré¡¹ç›®/TradingAgent/qihuo/database")
        self.modules_config = {
            "inventory": {
                "name": "åº“å­˜æ•°æ®",
                "path": self.database_root / "inventory",
                "data_file": "inventory.csv",
                "date_column": "date",
                "date_format": "%Y-%m-%d",
                "update_script": "å¢é‡æ›´æ–°_å®Œæ•´å¢å¼ºç‰ˆ.py",
                "structure_type": "by_commodity"  # æŒ‰å“ç§åˆ†æ–‡ä»¶å¤¹
            },
            "positioning": {
                "name": "æŒä»“å¸­ä½",
                "path": self.database_root / "positioning", 
                "data_file": "long_position_ranking.csv",
                "date_column": "date",
                "date_format": "%Y%m%d",
                "update_script": "å®Œæ•´æœŸè´§æŒä»“æ•°æ®ç®¡ç†ç³»ç»Ÿ.py",
                "structure_type": "by_commodity"  # æŒ‰å“ç§åˆ†æ–‡ä»¶å¤¹
            },
            "term_structure": {
                "name": "æœŸé™ç»“æ„",
                "path": self.database_root / "term_structure",
                "data_file": "term_structure.csv",
                "date_column": "date",
                "date_format": "%Y%m%d",
                "update_script": "å®Œæ•´ä¿®æ­£ç‰ˆæœŸé™ç»“æ„æ•°æ®åº“æ›´æ–°å™¨.py",
                "structure_type": "by_commodity"  # æŒ‰å“ç§åˆ†æ–‡ä»¶å¤¹
            },
            "technical_analysis": {
                "name": "æŠ€æœ¯é¢åˆ†æ",
                "path": self.database_root / "technical_analysis",
                "data_file": "ohlc_data.csv",
                "date_column": "æ—¶é—´",
                "date_format": "%Y-%m-%d",
                "update_script": "æœŸè´§æŠ€æœ¯åˆ†ææ•°æ®è·å–ç³»ç»Ÿ.py",
                "structure_type": "by_commodity"  # æŒ‰å“ç§åˆ†æ–‡ä»¶å¤¹
            },
            "basis": {
                "name": "åŸºå·®åˆ†æ",
                "path": self.database_root / "basis",
                "data_file": "basis_data.csv",
                "date_column": "date",
                "date_format": "%Y-%m-%d",
                "update_script": "æ™ºèƒ½åŸºå·®æ•°æ®æ›´æ–°ç³»ç»Ÿ_æ”¹è¿›ç‰ˆ.py",
                "structure_type": "by_commodity"  # æŒ‰å“ç§åˆ†æ–‡ä»¶å¤¹
            },
            "receipt": {
                "name": "ä»“å•æ•°æ®",
                "path": self.database_root / "receipt",
                "data_file": "receipt.csv",
                "date_column": "date",
                "date_format": "%Y-%m-%d",
                "update_script": "ä»“å•æ•°æ®é‡‡é›†ç³»ç»Ÿ_ç»Ÿä¸€å‘½åç‰ˆ.py",
                "structure_type": "by_commodity"  # æŒ‰å“ç§åˆ†æ–‡ä»¶å¤¹
            }
        }
    
    @st.cache_data(ttl=300)
    def get_data_status(_self) -> Dict:
        """è·å–æ•°æ®çŠ¶æ€"""
        status_data = {
            "modules": {},
            "summary": {
                "total_commodities": set(),
                "common_commodities": set(),
                "last_update": None
            }
        }
        
        all_commodities = []
        module_commodity_sets = []
        
        for module_name, config in _self.modules_config.items():
            module_status = {
                "name": config["name"],
                "status": "checking",
                "path": str(config["path"]),
                "last_update": "æœªçŸ¥",
                "commodities_count": 0,
                "total_records": 0,
                "latest_date": "æœªçŸ¥",
                "data_quality": "æ£€æŸ¥ä¸­"
            }
            
            try:
                data_path = config["path"]
                
                if not data_path.exists():
                    module_status["status"] = "error"
                    module_status["error_message"] = "æ•°æ®ç›®å½•ä¸å­˜åœ¨"
                else:
                    # æ£€æŸ¥æ˜¯å¦ä¸ºæŒ‰å“ç§åˆ†æ–‡ä»¶å¤¹çš„ç»“æ„
                    if config.get("structure_type") == "by_commodity":
                        # æ‰«æå“ç§æ–‡ä»¶å¤¹
                        commodity_folders = []
                        total_records = 0
                        latest_dates = []
                        
                        for item in data_path.iterdir():
                            if item.is_dir():
                                commodity = item.name
                                data_file = item / config["data_file"]
                                
                                if data_file.exists():
                                    try:
                                        df = pd.read_csv(data_file, encoding='utf-8')
                                        if not df.empty:
                                            commodity_folders.append(commodity)
                                            total_records += len(df)
                                            
                                            # å¤„ç†æ—¥æœŸåˆ—
                                            if config["date_column"] in df.columns:
                                                try:
                                                    date_col = df[config["date_column"]]
                                                    date_format = config.get("date_format", "%Y-%m-%d")
                                                    
                                                    if date_format == "%Y%m%d":
                                                        if date_col.dtype in ['int64', 'float64']:
                                                            df[config["date_column"]] = pd.to_datetime(date_col.astype(str), format='%Y%m%d', errors='coerce')
                                                        else:
                                                            df[config["date_column"]] = pd.to_datetime(date_col, format='%Y%m%d', errors='coerce')
                                                    else:
                                                        df[config["date_column"]] = pd.to_datetime(date_col, format=date_format, errors='coerce')
                                                    
                                                    valid_dates = df[config["date_column"]].dropna()
                                                    if not valid_dates.empty:
                                                        latest_dates.append(valid_dates.max())
                                                        
                                                except Exception:
                                                    pass
                                    except Exception:
                                        pass
                        
                        if commodity_folders:
                            module_status["commodities_count"] = len(commodity_folders)
                            module_status["total_records"] = total_records
                            module_status["status"] = "success"
                            module_status["data_quality"] = "è‰¯å¥½"
                            
                            # è®°å½•å“ç§åˆ—è¡¨
                            all_commodities.extend(commodity_folders)
                            module_commodity_sets.append(set(commodity_folders))
                            
                            # æœ€æ–°æ—¥æœŸ
                            if latest_dates:
                                latest_date = max(latest_dates)
                                module_status["latest_date"] = latest_date.strftime("%Y-%m-%d")
                                module_status["last_update"] = latest_date.strftime("%Y-%m-%d")
                        else:
                            module_status["status"] = "warning"
                            module_status["error_message"] = "æœªæ‰¾åˆ°æœ‰æ•ˆçš„å“ç§æ•°æ®"
                    
                    else:
                        # å•ä¸€æ–‡ä»¶ç»“æ„ï¼ˆä¿æŒåŸæœ‰é€»è¾‘ï¼‰
                        data_file = data_path / config["data_file"]
                        
                        if not data_file.exists():
                            module_status["status"] = "error"
                            module_status["error_message"] = "æ•°æ®æ–‡ä»¶ä¸å­˜åœ¨"
                        else:
                            try:
                                df = pd.read_csv(data_file, encoding='utf-8')
                                
                                if not df.empty:
                                    module_status["total_records"] = len(df)
                                    
                                    # å¤„ç†æ—¥æœŸåˆ—
                                    if config["date_column"] in df.columns:
                                        try:
                                            date_col = df[config["date_column"]]
                                            date_format = config.get("date_format", "%Y-%m-%d")
                                            
                                            if date_format == "%Y%m%d":
                                                if date_col.dtype in ['int64', 'float64']:
                                                    df[config["date_column"]] = pd.to_datetime(date_col.astype(str), format='%Y%m%d', errors='coerce')
                                                else:
                                                    df[config["date_column"]] = pd.to_datetime(date_col, format='%Y%m%d', errors='coerce')
                                            else:
                                                df[config["date_column"]] = pd.to_datetime(date_col, format=date_format, errors='coerce')
                                            
                                            valid_dates = df[config["date_column"]].dropna()
                                            if not valid_dates.empty:
                                                latest_date = valid_dates.max()
                                                module_status["latest_date"] = latest_date.strftime("%Y-%m-%d")
                                                module_status["last_update"] = latest_date.strftime("%Y-%m-%d")
                                        
                                        except Exception as date_error:
                                            module_status["latest_date"] = f"æ—¥æœŸè§£æé”™è¯¯: {date_error}"
                                    
                                    # ç»Ÿè®¡å“ç§æ•°é‡
                                    commodity_columns = ['symbol', 'commodity', 'å“ç§', 'åˆçº¦']
                                    commodity_col = None
                                    for col in commodity_columns:
                                        if col in df.columns:
                                            commodity_col = col
                                            break
                                    
                                    if commodity_col:
                                        unique_commodities = df[commodity_col].unique()
                                        module_status["commodities_count"] = len(unique_commodities)
                                        all_commodities.extend(unique_commodities)
                                        module_commodity_sets.append(set(unique_commodities))
                                    
                                    module_status["status"] = "success"
                                    module_status["data_quality"] = "è‰¯å¥½"
                                    
                                else:
                                    module_status["status"] = "warning"
                                    module_status["error_message"] = "æ•°æ®æ–‡ä»¶ä¸ºç©º"
                                    
                            except Exception as read_error:
                                module_status["status"] = "error"
                                module_status["error_message"] = f"è¯»å–æ•°æ®å¤±è´¥: {read_error}"
                        
            except Exception as e:
                module_status["status"] = "error"
                module_status["error_message"] = f"æ£€æŸ¥æ¨¡å—å¤±è´¥: {e}"
            
            status_data["modules"][module_name] = module_status
        
        # è®¡ç®—æ±‡æ€»ä¿¡æ¯
        if all_commodities:
            status_data["summary"]["total_commodities"] = set(all_commodities)
            
            # æ‰¾å‡ºæ‰€æœ‰æ¨¡å—éƒ½æœ‰çš„å“ç§
            if module_commodity_sets:
                common_commodities = module_commodity_sets[0]
                for commodity_set in module_commodity_sets[1:]:
                    common_commodities = common_commodities.intersection(commodity_set)
                status_data["summary"]["common_commodities"] = common_commodities
        
        return status_data
    
    def _get_module_detailed_info(self, module_key: str) -> Dict:
        """è·å–æ¨¡å—çš„è¯¦ç»†æ•°æ®ä¿¡æ¯"""
        if module_key not in self.modules_config:
            return {}
        
        config = self.modules_config[module_key]
        detailed_info = {}
        
        try:
            data_path = config["path"]
            if not data_path.exists():
                return {}
            
            # æ£€æŸ¥æ˜¯å¦ä¸ºæŒ‰å“ç§åˆ†æ–‡ä»¶å¤¹çš„ç»“æ„
            if config.get("structure_type") == "by_commodity":
                for item in data_path.iterdir():
                    if item.is_dir():
                        commodity = item.name
                        data_file = item / config["data_file"]
                        
                        if data_file.exists():
                            try:
                                df = pd.read_csv(data_file, encoding='utf-8')
                                if not df.empty:
                                    info = {
                                        'record_count': len(df),
                                        'start_date': 'æœªçŸ¥',
                                        'end_date': 'æœªçŸ¥'
                                    }
                                    
                                    # å°è¯•è§£ææ—¥æœŸä¿¡æ¯
                                    if config["date_column"] in df.columns:
                                        try:
                                            date_col = df[config["date_column"]]
                                            date_format = config.get("date_format", "%Y-%m-%d")
                                            
                                            if date_format == "%Y%m%d":
                                                if date_col.dtype in ['int64', 'float64']:
                                                    df[config["date_column"]] = pd.to_datetime(date_col.astype(str), format='%Y%m%d', errors='coerce')
                                                else:
                                                    df[config["date_column"]] = pd.to_datetime(date_col, format='%Y%m%d', errors='coerce')
                                            else:
                                                df[config["date_column"]] = pd.to_datetime(date_col, format=date_format, errors='coerce')
                                            
                                            valid_dates = df[config["date_column"]].dropna()
                                            if not valid_dates.empty:
                                                info['start_date'] = valid_dates.min().strftime("%Y-%m-%d")
                                                info['end_date'] = valid_dates.max().strftime("%Y-%m-%d")
                                        except Exception:
                                            pass
                                    
                                    detailed_info[commodity] = info
                            except Exception:
                                pass
            
            return detailed_info
        except Exception:
            return {}
    
    def get_available_commodities(self) -> List[str]:
        """è·å–æ‰€æœ‰å¯ç”¨çš„å“ç§åˆ—è¡¨"""
        try:
            # è·å–æ•°æ®çŠ¶æ€
            status = self.get_data_status()
            
            # è¿”å›æœ‰æ•°æ®çš„å“ç§åˆ—è¡¨
            if "summary" in status and "common_commodities" in status["summary"]:
                commodities = list(status["summary"]["common_commodities"])
                
                # å¦‚æœæ²¡æœ‰å…±åŒå“ç§ï¼Œè¿”å›æ‰€æœ‰å“ç§
                if not commodities and "total_commodities" in status["summary"]:
                    commodities = list(status["summary"]["total_commodities"])
                
                return sorted(commodities) if commodities else []
            else:
                return []
                
        except Exception as e:
            print(f"âŒ è·å–å¯ç”¨å“ç§å¤±è´¥: {e}")
            return []
    
    def get_module_supported_varieties(self, module_name: str = None) -> Dict[str, Set[str]]:
        """è·å–å„æ¨¡å—æ”¯æŒçš„å“ç§åˆ—è¡¨"""
        
        # æ£€æŸ¥åº“å­˜æ•°æ®æ”¯æŒçš„å“ç§
        inventory_varieties = set()
        inventory_path = self.database_root / "inventory"
        if inventory_path.exists():
            for item in inventory_path.iterdir():
                if item.is_dir() and (item / "inventory.csv").exists():
                    inventory_varieties.add(item.name)
        
        # æ£€æŸ¥ä»“å•æ•°æ®æ”¯æŒçš„å“ç§
        receipt_varieties = set()
        receipt_path = self.database_root / "receipt"
        if receipt_path.exists():
            for item in receipt_path.iterdir():
                if item.is_dir() and (item / "receipt.csv").exists():
                    receipt_varieties.add(item.name)
        
        # æ£€æŸ¥æŒä»“æ•°æ®æ”¯æŒçš„å“ç§
        positioning_varieties = set()
        positioning_path = self.database_root / "positioning"
        if positioning_path.exists():
            for item in positioning_path.iterdir():
                if item.is_dir() and (item / "long_position_ranking.csv").exists():
                    positioning_varieties.add(item.name)
        
        # æ£€æŸ¥åŸºå·®æ•°æ®æ”¯æŒçš„å“ç§
        basis_varieties = set()
        basis_path = self.database_root / "basis"
        if basis_path.exists():
            for item in basis_path.iterdir():
                if item.is_dir() and (item / "basis_data.csv").exists():
                    basis_varieties.add(item.name)
        
        # æ£€æŸ¥æœŸé™ç»“æ„æ•°æ®æ”¯æŒçš„å“ç§
        term_structure_varieties = set()
        term_structure_path = self.database_root / "term_structure"
        if term_structure_path.exists():
            for item in term_structure_path.iterdir():
                if item.is_dir() and (item / "term_structure.csv").exists():
                    term_structure_varieties.add(item.name)
        
        # å®šä¹‰å„æ¨¡å—æ”¯æŒçš„å“ç§
        module_varieties = {
            # åº“å­˜ä»“å•åˆ†æï¼šæ•°æ®å¢å¼ºç‰ˆæ”¯æŒä»»ä½•å“ç§ï¼ˆæœ¬åœ°æ•°æ®+è”ç½‘æœç´¢ï¼‰
            "åº“å­˜ä»“å•åˆ†æ": self._get_all_futures_varieties(),
            
            # æŒä»“å¸­ä½åˆ†æï¼šéœ€è¦æŒä»“æ•°æ®ï¼Œå¯ä»¥è”ç½‘è¡¥å……
            "æŒä»“å¸­ä½åˆ†æ": positioning_varieties,
            
            # åŸºå·®åˆ†æï¼šè”ç½‘å¢å¼ºç‰ˆæ”¯æŒä»»ä½•å“ç§ï¼ˆæœ¬åœ°æ•°æ®+è”ç½‘æœç´¢ï¼‰
            "åŸºå·®åˆ†æ": self._get_all_futures_varieties(),
            
            # æœŸé™ç»“æ„åˆ†æï¼šéœ€è¦æœŸé™ç»“æ„æ•°æ®
            "æœŸé™ç»“æ„åˆ†æ": term_structure_varieties,
            
            # æŠ€æœ¯åˆ†æï¼šç†è®ºä¸Šæ”¯æŒæ‰€æœ‰å“ç§ï¼ˆé€šè¿‡è”ç½‘è·å–ï¼‰
            "æŠ€æœ¯åˆ†æ": inventory_varieties | receipt_varieties | positioning_varieties | basis_varieties | term_structure_varieties,
            
            # æ–°é—»åˆ†æï¼šç†è®ºä¸Šæ”¯æŒæ‰€æœ‰å“ç§ï¼ˆé€šè¿‡è”ç½‘è·å–ï¼‰
            "æ–°é—»åˆ†æ": inventory_varieties | receipt_varieties | positioning_varieties | basis_varieties | term_structure_varieties,
        }
        
        if module_name:
            return module_varieties.get(module_name, set())
        else:
            return module_varieties
    
    def _get_all_futures_varieties(self) -> Set[str]:
        """è·å–æ‰€æœ‰æœŸè´§å“ç§ä»£ç ï¼ˆç”¨äºæ”¯æŒä»»ä½•å“ç§çš„æ¨¡å—ï¼‰"""
        # ä¸­å›½æœŸè´§å¸‚åœºä¸»è¦å“ç§ä»£ç 
        all_varieties = {
            # è´µé‡‘å±
            'AG', 'AU',
            # æœ‰è‰²é‡‘å±
            'CU', 'AL', 'ZN', 'PB', 'NI', 'SN', 'BC',
            # é»‘è‰²é‡‘å±
            'RB', 'HC', 'I', 'JM', 'J', 'SS', 'WR',
            # èƒ½æºåŒ–å·¥
            'BU', 'FU', 'LU', 'RU', 'NR', 'BR', 'SP', 'L', 'V', 'PP', 'EG', 'TA', 'SA', 'MA', 'FG', 'EB', 'PG', 'LG', 'SC',
            # å†œäº§å“
            'CF', 'SR', 'RM', 'OI', 'M', 'Y', 'A', 'B', 'C', 'CS', 'JD', 'PK', 'UR', 'LH', 'AP', 'CJ', 'WH', 'PM', 'RI', 'LR', 'JR',
            # æ–°ææ–™
            'SF', 'SM', 'ZC', 'PF', 'CY', 'AO', 'PS', 'LC', 'SI', 'PX', 'PR',
            # å…¶ä»–
            'SH'
        }
        return all_varieties
    
    def run_data_update_direct(self, module_name: str) -> Dict:
        """ç›´æ¥è¿è¡Œæ•°æ®æ›´æ–°è„šæœ¬ï¼Œè®©ç”¨æˆ·ä¸è„šæœ¬äº¤äº’"""
        config = self.modules_config.get(module_name)
        if not config or not config.get("update_script"):
            return {"status": "error", "message": "æ— æ›´æ–°è„šæœ¬"}
        
        script_path = Path(config["update_script"])
        if not script_path.exists():
            return {"status": "error", "message": f"æ›´æ–°è„šæœ¬ä¸å­˜åœ¨: {script_path}"}
        
        try:
            # åœ¨Windowsä¸‹å¯åŠ¨æ–°çš„å‘½ä»¤è¡Œçª—å£è¿è¡Œè„šæœ¬
            import subprocess
            import sys
            
            # æ„å»ºå‘½ä»¤
            cmd = f'start cmd /k "cd /d {Path.cwd()} && python {script_path} && pause"'
            
            # ä½¿ç”¨shell=Trueåœ¨æ–°çª—å£ä¸­æ‰§è¡Œ
            result = subprocess.run(cmd, shell=True, cwd=Path.cwd())
            
            return {
                "status": "success",
                "message": f"å·²å¯åŠ¨ {config['name']} æ•°æ®æ›´æ–°",
                "details": "è¯·åœ¨æ–°æ‰“å¼€çš„å‘½ä»¤è¡Œçª—å£ä¸­å®Œæˆäº¤äº’æ“ä½œ"
            }
            
        except Exception as e:
            return {"status": "error", "message": f"å¯åŠ¨æ›´æ–°è„šæœ¬å¤±è´¥: {e}"}

# ============================================================================
# å®Œæ•´åˆ†æç®¡ç†å™¨
# ============================================================================

class StreamlitAnalysisManager:
    """Streamlitåˆ†ææ‰§è¡Œç®¡ç†å™¨"""
    
    def __init__(self):
        try:
            self.config = FuturesTradingAgentsConfig("æœŸè´§TradingAgentsç³»ç»Ÿ_é…ç½®æ–‡ä»¶.json")
            self.system = None
            self.current_analysis = None
            self.analysis_results = {}
        except Exception as e:
            st.error(f"é…ç½®åˆå§‹åŒ–å¤±è´¥: {e}")
            self.config = None
    
    def _check_commodity_local_data(self, commodity: str, data_status: Dict) -> Dict[str, bool]:
        """æ£€æŸ¥æŒ‡å®šå“ç§åœ¨å„ä¸ªæ¨¡å—ä¸­çš„æœ¬åœ°æ•°æ®æƒ…å†µ"""
        module_mapping = {
            "technical": "ä¸“ä¸šAIæŠ€æœ¯åˆ†æç³»ç»Ÿ",
            "basis": "ä¸“ä¸šåŸºå·®åˆ†æç³»ç»Ÿ",
            "inventory": "åº“å­˜åˆ†æç³»ç»Ÿ", 
            "positioning": "æŒä»“åˆ†æç³»ç»Ÿ",
            "term_structure": "æœŸé™ç»“æ„åˆ†æç³»ç»Ÿ"
        }
        
        result = {}
        
        for module_key, module_name in module_mapping.items():
            has_data = False
            
            # æ£€æŸ¥è¯¥æ¨¡å—æ˜¯å¦æœ‰è¯¥å“ç§çš„æ•°æ®
            if module_name in data_status.get("modules", {}):
                module_info = data_status["modules"][module_name]
                
                # æ£€æŸ¥å“ç§æ˜¯å¦åœ¨è¯¥æ¨¡å—çš„å“ç§åˆ—è¡¨ä¸­
                if "commodities" in module_info:
                    commodities = module_info["commodities"]
                    if isinstance(commodities, (list, set)):
                        has_data = commodity in commodities
                    elif isinstance(commodities, dict):
                        has_data = commodity in commodities.keys()
                
                # å¦‚æœæ²¡æœ‰å“ç§åˆ—è¡¨ï¼Œæ£€æŸ¥æ˜¯å¦æœ‰è®°å½•æ•°
                if not has_data and module_info.get("commodities_count", 0) > 0:
                    # è¿›ä¸€æ­¥æ£€æŸ¥æ˜¯å¦çœŸçš„æœ‰è¯¥å“ç§çš„æ•°æ®
                    has_data = self._check_specific_commodity_data(commodity, module_key)
            
            result[module_key] = has_data
        
        return result
    
    def _check_specific_commodity_data(self, commodity: str, module_key: str) -> bool:
        """æ£€æŸ¥ç‰¹å®šå“ç§åœ¨æŒ‡å®šæ¨¡å—ä¸­æ˜¯å¦æœ‰æ•°æ®"""
        try:
            # è¿™é‡Œå¯ä»¥æ ¹æ®å®é™…çš„æ•°æ®å­˜å‚¨ç»“æ„æ¥æ£€æŸ¥
            # ç®€åŒ–å®ç°ï¼šæ£€æŸ¥æ•°æ®ç®¡ç†å™¨ä¸­æ˜¯å¦æœ‰è¯¥å“ç§çš„æ•°æ®
            if hasattr(st.session_state, 'data_manager'):
                data_manager = st.session_state.data_manager
                
                # æ ¹æ®æ¨¡å—ç±»å‹æ£€æŸ¥ä¸åŒçš„æ•°æ®æº
                if module_key == "technical":
                    # æ£€æŸ¥æŠ€æœ¯åˆ†ææ•°æ®
                    tech_path = data_manager.modules_config.get("ä¸“ä¸šAIæŠ€æœ¯åˆ†æç³»ç»Ÿ", {}).get("path")
                    if tech_path and tech_path.exists():
                        commodity_path = tech_path / commodity
                        return commodity_path.exists() and any(commodity_path.iterdir())
                
                elif module_key == "basis":
                    # æ£€æŸ¥åŸºå·®åˆ†ææ•°æ®
                    basis_path = data_manager.modules_config.get("ä¸“ä¸šåŸºå·®åˆ†æç³»ç»Ÿ", {}).get("path")
                    if basis_path and basis_path.exists():
                        commodity_path = basis_path / commodity
                        return commodity_path.exists() and any(commodity_path.iterdir())
                
                elif module_key == "inventory":
                    # æ£€æŸ¥åº“å­˜æ•°æ®
                    inv_path = data_manager.modules_config.get("åº“å­˜åˆ†æç³»ç»Ÿ", {}).get("path")
                    if inv_path and inv_path.exists():
                        data_file = inv_path / "inventory_data.csv"
                        if data_file.exists():
                            import pandas as pd
                            df = pd.read_csv(data_file)
                            return commodity in df.get('variety', []).values if not df.empty else False
                
                elif module_key == "positioning":
                    # æ£€æŸ¥æŒä»“æ•°æ®
                    pos_path = data_manager.modules_config.get("æŒä»“åˆ†æç³»ç»Ÿ", {}).get("path")
                    if pos_path and pos_path.exists():
                        data_file = pos_path / "positioning_data.csv"
                        if data_file.exists():
                            import pandas as pd
                            df = pd.read_csv(data_file)
                            return commodity in df.get('variety', []).values if not df.empty else False
                
                elif module_key == "term_structure":
                    # æ£€æŸ¥æœŸé™ç»“æ„æ•°æ®
                    term_path = data_manager.modules_config.get("æœŸé™ç»“æ„åˆ†æç³»ç»Ÿ", {}).get("path")
                    if term_path and term_path.exists():
                        commodity_path = term_path / commodity
                        return commodity_path.exists() and any(commodity_path.iterdir())
            
            return False
            
        except Exception as e:
            print(f"æ£€æŸ¥å“ç§æ•°æ®æ—¶å‡ºé”™: {e}")
            return False

    def initialize_system(self):
        """åˆå§‹åŒ–ç³»ç»Ÿ"""
        try:
            if not self.config:
                return False, "é…ç½®æ–‡ä»¶æœªæ­£ç¡®åŠ è½½"
            
            self.system = CompleteFuturesTradingExecution(
                config_file="æœŸè´§TradingAgentsç³»ç»Ÿ_é…ç½®æ–‡ä»¶.json"
            )
            return True, "ç³»ç»Ÿåˆå§‹åŒ–æˆåŠŸ"
        except Exception as e:
            return False, f"ç³»ç»Ÿåˆå§‹åŒ–å¤±è´¥: {e}"
    
    def run_integrated_analysis(self, commodity: str, analysis_date: str, 
                              selected_modules: List[str], analysis_mode: str,
                              debate_rounds: int = 3) -> Dict:
        """è¿è¡Œå®Œæ•´é›†æˆåˆ†æ"""
        
        if not self.system:
            success, message = self.initialize_system()
            if not success:
                return {"status": "error", "message": message}
        
        try:
            # åˆ›å»ºæ•°æ®æ•´åˆå™¨
            integrator = FuturesAnalysisIntegrator(
                data_root_dir=self.config.get("paths", {}).get("data_root_dir", "D:/Cursor/cursoré¡¹ç›®/TradingAgent/qihuo/database"),
                config=self.config.to_dict()
            )
            
            # ä½¿ç”¨åŒæ­¥æ–¹å¼æˆ–çº¿ç¨‹æ± æ‰§è¡Œå¼‚æ­¥ä»»åŠ¡ï¼Œé¿å…äº‹ä»¶å¾ªç¯å†²çª
            import concurrent.futures
            import threading
            
            def run_async_in_thread(coro):
                """åœ¨æ–°çº¿ç¨‹ä¸­è¿è¡Œå¼‚æ­¥åç¨‹"""
                def thread_target():
                    # åœ¨æ–°çº¿ç¨‹ä¸­åˆ›å»ºäº‹ä»¶å¾ªç¯
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    try:
                        return loop.run_until_complete(coro)
                    finally:
                        loop.close()
                
                # ä½¿ç”¨çº¿ç¨‹æ‰§è¡Œ
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(thread_target)
                    return future.result()
            
            if analysis_mode == "complete_flow":
                # å®Œæ•´æµç¨‹ï¼šåˆ†æå¸ˆ + è¾©è®º + äº¤æ˜“å‘˜ + é£æ§ + CIOå†³ç­–
                # é˜¶æ®µ1ï¼šæ‰§è¡Œ6å¤§åˆ†ææ¨¡å—
                st.info("ğŸ” é˜¶æ®µ1ï¼šæ‰§è¡Œåˆ†æå¸ˆå›¢é˜Ÿåˆ†æ...")
                analysis_state = run_async_in_thread(
                    integrator.collect_all_analyses(commodity, analysis_date, selected_modules)
                )
                
                # é˜¶æ®µ2ï¼šæ‰§è¡Œå®Œæ•´å†³ç­–æµç¨‹ï¼ˆå«äº¤æ˜“å‘˜ç¯èŠ‚ï¼‰
                st.info(f"ğŸ­ é˜¶æ®µ2ï¼šæ‰§è¡Œ{debate_rounds}è½®å®Œæ•´å†³ç­–æµç¨‹ï¼ˆå«äº¤æ˜“å‘˜ç¯èŠ‚ï¼‰...")
                debate_result = run_async_in_thread(
                    integrator.run_optimized_debate_risk_decision(analysis_state, debate_rounds)
                )
                
                return {
                    "status": "success",
                    "result": {
                        "analysis_state": analysis_state,
                        "debate_result": debate_result
                    },
                    "type": "complete_flow",
                    "execution_summary": {
                        "stages_completed": 5,  # åˆ†æå¸ˆ + è¾©è®º + äº¤æ˜“å‘˜ + é£æ§ + CIO
                        "analysis_modules": len(selected_modules),
                        "debate_rounds": debate_rounds,
                        "final_decision": debate_result.get("decision_section", {}).get("final_decision", "æœªçŸ¥") if debate_result and debate_result.get("success") else "åˆ†æå¤±è´¥",
                        "confidence_score": debate_result.get("decision_section", {}).get("confidence", "0%") if debate_result and debate_result.get("success") else "0%"
                    }
                }
                
            else:
                # ä»…æ‰§è¡Œåˆ†æå¸ˆæµç¨‹
                st.info("ğŸ“Š æ‰§è¡Œåˆ†æå¸ˆå›¢é˜Ÿåˆ†æ...")
                analysis_state = run_async_in_thread(
                    integrator.collect_all_analyses(commodity, analysis_date, selected_modules)
                )
                
                return {
                    "status": "success", 
                    "result": analysis_state,
                    "type": "analyst_only"
                }
                
        except Exception as e:
            # æ›´è¯¦ç»†çš„é”™è¯¯ä¿¡æ¯
            error_msg = str(e)
            if "'NoneType' object has no attribute 'post'" in error_msg:
                error_msg = "DeepSeek APIå®¢æˆ·ç«¯åˆå§‹åŒ–å¤±è´¥ï¼Œè¯·æ£€æŸ¥APIå¯†é’¥é…ç½®"
            elif "debate_result" in error_msg:
                error_msg = "è¾©è®ºåˆ†æå¤±è´¥ï¼ŒAPIè°ƒç”¨å¼‚å¸¸å¯¼è‡´ç»“æœä¸ºç©º"
            
            return {"status": "error", "message": error_msg}
    
    def start_analysis(self, commodities: List[str], analysis_date: str, 
                      modules: List[str], analysis_mode: str, config: Dict):
        """å¯åŠ¨åˆ†æ"""
        
        self.current_analysis = {
            "commodities": commodities,
            "analysis_date": analysis_date,
            "modules": modules,
            "analysis_mode": analysis_mode,
            "config": config,
            "status": "running",
            "results": {},
            "current_commodity": None,
            "start_time": datetime.now()
        }
        
        # æ¸…ç©ºä¹‹å‰çš„ç»“æœ
        self.analysis_results = {}
    
    def execute_analysis_for_commodity(self, commodity: str) -> Dict:
        """ä¸ºå•ä¸ªå“ç§æ‰§è¡Œåˆ†æ"""
        if not self.current_analysis:
            return {"status": "error", "message": "æœªå¯åŠ¨åˆ†æ"}
        
        try:
            # æ›´æ–°å½“å‰åˆ†æå“ç§
            self.current_analysis["current_commodity"] = commodity
            
            # æ‰§è¡Œåˆ†æ
            result = self.run_integrated_analysis(
                commodity=commodity,
                analysis_date=self.current_analysis["analysis_date"],
                selected_modules=self.current_analysis["modules"],
                analysis_mode=self.current_analysis["analysis_mode"],
                debate_rounds=self.current_analysis["config"].get("debate_rounds", 3)
            )
            
            # ä¿å­˜ç»“æœ
            self.analysis_results[commodity] = result
            
            return result
            
        except Exception as e:
            error_result = {"status": "error", "message": str(e)}
            self.analysis_results[commodity] = error_result
            return error_result
    
    def process_all_commodities(self):
        """å¤„ç†æ‰€æœ‰å“ç§çš„åˆ†æ"""
        if not self.current_analysis:
            return
        
        commodities = self.current_analysis["commodities"]
        
        for commodity in commodities:
            if self.current_analysis.get("status") != "running":
                break
                
            st.info(f"ğŸ” æ­£åœ¨åˆ†æ {commodity}...")
            result = self.execute_analysis_for_commodity(commodity)
            
            if result.get("status") == "success":
                st.success(f"âœ… {commodity} åˆ†æå®Œæˆ")
            else:
                st.error(f"âŒ {commodity} åˆ†æå¤±è´¥: {result.get('message', 'æœªçŸ¥é”™è¯¯')}")
        
        # æ ‡è®°åˆ†æå®Œæˆ
        if self.current_analysis:
            self.current_analysis["status"] = "completed"
    
    def is_analysis_running(self) -> bool:
        """æ£€æŸ¥æ˜¯å¦æœ‰åˆ†æåœ¨è¿è¡Œ"""
        return (self.current_analysis and 
                self.current_analysis.get("status") == "running")
    
    def get_analysis_progress(self) -> Dict:
        """è·å–åˆ†æè¿›åº¦"""
        if not self.current_analysis:
            return {"progress": 0, "message": "æœªå¼€å§‹åˆ†æ"}
        
        total_commodities = len(self.current_analysis["commodities"])
        completed_commodities = len(self.analysis_results)
        
        progress = (completed_commodities / total_commodities) * 100 if total_commodities > 0 else 0
        
        return {
            "progress": progress,
            "completed": completed_commodities,
            "total": total_commodities,
            "current": self.current_analysis.get("current_commodity"),
            "message": f"å·²å®Œæˆ {completed_commodities}/{total_commodities} ä¸ªå“ç§"
        }
    
    def display_progress(self):
        """æ˜¾ç¤ºåˆ†æè¿›åº¦"""
        if not self.current_analysis:
            st.info("ğŸ‘ˆ è¯·åœ¨åˆ†æé…ç½®é¡µé¢è®¾ç½®å‚æ•°å¹¶å¼€å§‹åˆ†æ")
            return
        
        progress_info = self.get_analysis_progress()
        
        # æ˜¾ç¤ºè¿›åº¦æ¡
        progress_bar = st.progress(progress_info["progress"] / 100)
        st.write(f"ğŸ“Š {progress_info['message']}")
        
        if progress_info.get("current"):
            st.write(f"ğŸ” å½“å‰åˆ†æ: {progress_info['current']}")
        
        # æ˜¾ç¤ºåˆ†æé…ç½®ä¿¡æ¯ - ä½¿ç”¨ç‹¬ç«‹çš„expander
        st.markdown("#### ğŸ“‹ åˆ†æé…ç½®ä¿¡æ¯")
        config = self.current_analysis
        col1, col2 = st.columns(2)
        
        with col1:
            st.write(f"**å“ç§åˆ—è¡¨**: {', '.join(config['commodities'])}")
            st.write(f"**åˆ†ææ—¥æœŸ**: {config['analysis_date']}")
            st.write(f"**åˆ†ææ¨¡å¼**: {config['analysis_mode']}")
        
        with col2:
            st.write(f"**åˆ†ææ¨¡å—**: {', '.join(config['modules'])}")
            st.write(f"**å¼€å§‹æ—¶é—´**: {config['start_time'].strftime('%H:%M:%S')}")
            st.write(f"**çŠ¶æ€**: {config['status']}")
    
    def display_results(self):
        """æ˜¾ç¤ºåˆ†æç»“æœ - ä¼˜åŒ–å¸ƒå±€ç‰ˆæœ¬"""
        if not self.analysis_results:
            st.info("â³ æš‚æ— åˆ†æç»“æœï¼Œè¯·ç­‰å¾…åˆ†æå®Œæˆ...")
            return
        
        st.subheader("ğŸ“ˆ åˆ†æç»“æœ")
        
        # ä¸ºæ¯ä¸ªå“ç§åˆ›å»ºç‹¬ç«‹çš„é¡µé¢å±•ç¤º
        for commodity, result in self.analysis_results.items():
            # ä½¿ç”¨å¤§æ ‡é¢˜åŒºåˆ†ä¸åŒå“ç§
            st.markdown(f"""
            <div style="background: linear-gradient(90deg, #FF6B6B, #4ECDC4); 
                       padding: 20px; margin: 20px 0; border-radius: 10px; 
                       text-align: center;">
                <h2 style="color: white; margin: 0;">ğŸ“Š {commodity} å®Œæ•´åˆ†ææŠ¥å‘Š</h2>
            </div>
            """, unsafe_allow_html=True)
            
            if result.get("status") == "success":
                result_type = result.get("type", "unknown")
                
                if result_type == "complete_flow":
                    # å®Œæ•´æµç¨‹ç»“æœ - åˆ†é¡µæ˜¾ç¤º
                    self._display_complete_flow_result_paginated(commodity, result)
                elif result_type == "analyst_only":
                    # ä»…åˆ†æå¸ˆç»“æœ - åˆ†é¡µæ˜¾ç¤º
                    self._display_analyst_only_result_paginated(commodity, result)
                else:
                    st.write(f"âœ… {commodity} åˆ†æå®Œæˆ")
                    st.json(result)
            
            else:
                st.error(f"âŒ {commodity} åˆ†æå¤±è´¥: {result.get('message', 'æœªçŸ¥é”™è¯¯')}")
            
            # æ·»åŠ æ˜æ˜¾çš„å“ç§åˆ†éš”
            st.markdown("---")
            st.markdown("<br><br>", unsafe_allow_html=True)
    
    def _display_complete_flow_result(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºå®Œæ•´æµç¨‹ç»“æœ"""
        
        st.success(f"âœ… {commodity} å®Œæ•´5é˜¶æ®µæµç¨‹åˆ†æå®Œæˆ")
        
        # æ‰§è¡Œæ‘˜è¦
        execution_summary = result.get("execution_summary", {})
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("å®Œæˆé˜¶æ®µ", f"{execution_summary.get('stages_completed', 0)}/5")
        with col2:
            st.metric("åˆ†ææ¨¡å—", execution_summary.get('analysis_modules', 0))
        with col3:
            st.metric("è¾©è®ºè½®æ•°", execution_summary.get('debate_rounds', 0))
        
        # æœ€ç»ˆå†³ç­–
        final_decision = execution_summary.get('final_decision', 'æœªçŸ¥')
        if final_decision != 'æœªçŸ¥':
            decision_color = "ğŸŸ¢" if "çœ‹å¤š" in final_decision or "ä¹°å…¥" in final_decision else "ğŸ”´" if "çœ‹ç©º" in final_decision or "å–å‡º" in final_decision else "ğŸŸ¡"
            st.markdown(f"### {decision_color} æœ€ç»ˆå†³ç­–: **{final_decision}**")
        
        # è¯¦ç»†ç»“æœ
        analysis_result = result.get("result", {})
        
        # åˆ†æå¸ˆå›¢é˜Ÿç»“æœ
        if "analysis_state" in analysis_result:
            st.markdown("#### ğŸ“Š åˆ†æå¸ˆå›¢é˜ŸæŠ¥å‘Š")
            analysis_state = analysis_result["analysis_state"]
            
            # ç»Ÿè®¡æˆåŠŸçš„æ¨¡å—
            module_attrs = [
                ("inventory_analysis", "ğŸ“¦ åº“å­˜ä»“å•åˆ†æ"),
                ("positioning_analysis", "ğŸ¯ æŒä»“å¸­ä½åˆ†æ"),
                ("technical_analysis", "ğŸ“ˆ æŠ€æœ¯åˆ†æ"),
                ("basis_analysis", "ğŸ’° åŸºå·®åˆ†æ"),
                ("term_structure_analysis", "ğŸ“Š æœŸé™ç»“æ„åˆ†æ"),
                ("news_analysis", "ğŸ“° æ–°é—»åˆ†æ")
            ]
            
            successful_modules = 0
            for attr_name, module_name in module_attrs:
                if hasattr(analysis_state, attr_name):
                    module_result = getattr(analysis_state, attr_name)
                    # ä¿®å¤çŠ¶æ€åˆ¤æ–­ - æ”¯æŒæšä¸¾å’Œå­—ç¬¦ä¸²ä¸¤ç§å½¢å¼
                    if module_result and (
                        module_result.status == "completed" or 
                        (hasattr(module_result.status, 'value') and module_result.status.value == "completed") or
                        str(module_result.status).lower() == "completed"
                    ):
                        successful_modules += 1
            
            st.success(f"âœ… {successful_modules}å¤§åˆ†ææ¨¡å—å·²å®Œæˆ")
            
            # æ˜¾ç¤ºæ¯ä¸ªæ¨¡å—çš„è¯¦ç»†æŠ¥å‘Š
            for attr_name, module_name in module_attrs:
                if hasattr(analysis_state, attr_name):
                    module_result = getattr(analysis_state, attr_name)
                    
                    st.markdown(f"##### {module_name}")
                    
                    # ä¿®å¤çŠ¶æ€åˆ¤æ–­ - æ”¯æŒæšä¸¾å’Œå­—ç¬¦ä¸²ä¸¤ç§å½¢å¼
                    if module_result and (
                        module_result.status == "completed" or 
                        (hasattr(module_result.status, 'value') and module_result.status.value == "completed") or
                        str(module_result.status).lower() == "completed"
                    ):
                        st.success("âœ… åˆ†æå®Œæˆ")
                        
                        # æ˜¾ç¤ºåˆ†æå†…å®¹
                        result_data = module_result.result_data
                        analysis_content = None
                        
                        # å°è¯•ä»å¤šä¸ªå¯èƒ½çš„å­—æ®µè·å–åˆ†æå†…å®¹
                        if result_data:
                            # å°è¯•ä¸åŒçš„å†…å®¹å­—æ®µï¼Œä¼˜å…ˆçº§è°ƒæ•´ä¸ºåº“å­˜åˆ†æçš„å­—æ®µ
                            content_fields = ["ai_comprehensive_analysis", "ai_analysis", "analysis_content", "speculative_analysis", "content", "report", "analysis"]
                            for field in content_fields:
                                if field in result_data and result_data[field]:
                                    analysis_content = result_data[field]
                                    break
                            
                            # å¦‚æœæ˜¯å­—å…¸ï¼Œå°è¯•è·å–åµŒå¥—çš„å†…å®¹
                            if not analysis_content and isinstance(result_data, dict):
                                for key, value in result_data.items():
                                    if isinstance(value, dict) and "analysis_content" in value:
                                        analysis_content = value["analysis_content"]
                                        break
                        
                        # æ˜¾ç¤ºåˆ†æå†…å®¹
                        if analysis_content and isinstance(analysis_content, str) and len(analysis_content.strip()) > 0:
                            # æˆªå–å‰500å­—ç¬¦ä½œä¸ºæ‘˜è¦
                            summary = analysis_content[:500] + ("..." if len(analysis_content) > 500 else "")
                            st.write("**åˆ†ææ‘˜è¦:**")
                            st.write(summary)
                            
                            # å¦‚æœæœ‰å®Œæ•´å†…å®¹ï¼Œæä¾›å±•å¼€é€‰é¡¹
                            if len(analysis_content) > 500:
                                if st.button(f"æŸ¥çœ‹{module_name}å®Œæ•´æŠ¥å‘Š", key=f"expand_{attr_name}_{commodity}"):
                                    st.text_area(f"{module_name}å®Œæ•´æŠ¥å‘Š", analysis_content, height=400, key=f"full_{attr_name}_{commodity}")
                        else:
                            # è°ƒè¯•ä¿¡æ¯ï¼šæ˜¾ç¤ºå¯ç”¨çš„å­—æ®µ
                            if st.session_state.get('debug_mode', False) and result_data:
                                st.caption(f"è°ƒè¯•: å¯ç”¨å­—æ®µ - {list(result_data.keys()) if isinstance(result_data, dict) else type(result_data)}")
                        
                        # æ˜¾ç¤ºå…¶ä»–å…³é”®ä¿¡æ¯ï¼ˆç§»é™¤æ•°æ®è´¨é‡æ˜¾ç¤ºï¼‰
                        if result_data:
                            if "confidence_score" in result_data:
                                confidence = result_data['confidence_score']
                                try:
                                    if isinstance(confidence, (int, float)):
                                        st.write(f"**åˆ†æä¿¡å¿ƒåº¦**: {confidence:.1%}")
                                    else:
                                        st.write(f"**åˆ†æä¿¡å¿ƒåº¦**: {confidence}")
                                except (ValueError, TypeError):
                                    st.write(f"**åˆ†æä¿¡å¿ƒåº¦**: {confidence}")
                            
                            if "key_findings" in result_data and result_data["key_findings"]:
                                st.write("**å…³é”®å‘ç°:**")
                                findings = result_data["key_findings"]
                                if isinstance(findings, list):
                                    for finding in findings[:3]:  # æ˜¾ç¤ºå‰3ä¸ªå…³é”®å‘ç°
                                        st.write(f"â€¢ {finding}")
                    
                    else:
                        st.error(f"âŒ {module_name}åˆ†æå¤±è´¥")
                        error_msg = module_result.error_message if module_result else "æ— æ•°æ®"
                        st.write(f"é”™è¯¯ä¿¡æ¯: {error_msg}")
                    
                    st.markdown("---")
        
        # è¾©è®ºé£æ§å†³ç­–ç»“æœ
        if "debate_result" in analysis_result:
            debate_result = analysis_result["debate_result"]
            
            if debate_result.get("success"):
                # è¾©è®ºç»“æœ
                if "debate_section" in debate_result:
                    st.markdown("#### ğŸ­ æ¿€çƒˆè¾©è®ºç»“æœ")
                    debate = debate_result["debate_section"]
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.write(f"**è¾©è®ºèƒœè€…**: {debate.get('winner', 'æœªçŸ¥')}")
                    with col2:
                        scores = debate.get('scores', {})
                        st.write(f"**å¤šå¤´å¾—åˆ†**: {scores.get('bull', 0):.1f}")
                    with col3:
                        st.write(f"**ç©ºå¤´å¾—åˆ†**: {scores.get('bear', 0):.1f}")
                    
                    st.write(f"**è¾©è®ºæ€»ç»“**: {debate.get('summary', 'æš‚æ— æ€»ç»“')}")
                    
                    # æ˜¾ç¤ºè¯¦ç»†è¾©è®ºè½®æ¬¡
                    if "rounds" in debate and debate["rounds"]:
                        st.markdown("##### ğŸ—£ï¸ è¯¦ç»†è¾©è®ºå†…å®¹")
                        
                        for i, round_data in enumerate(debate["rounds"], 1):
                            st.markdown(f"**ç¬¬{i}è½®è¾©è®º**")
                            
                            # å¤šå¤´å‘è¨€
                            st.markdown("ğŸ‚ **å¤šå¤´è§‚ç‚¹:**")
                            bull_argument = round_data.get("bull_argument", "æš‚æ— å†…å®¹")
                            st.write(bull_argument[:300] + ("..." if len(bull_argument) > 300 else ""))
                            if len(bull_argument) > 300:
                                if st.button(f"æŸ¥çœ‹ç¬¬{i}è½®å¤šå¤´å®Œæ•´å‘è¨€", key=f"bull_full_{i}_{commodity}"):
                                    st.text_area(f"ç¬¬{i}è½®å¤šå¤´å®Œæ•´å‘è¨€", bull_argument, height=200, key=f"bull_text_{i}_{commodity}")
                            
                            st.write(f"*å¾—åˆ†: {round_data.get('bull_score', 0):.1f}åˆ†*")
                            
                            # ç©ºå¤´å‘è¨€
                            st.markdown("ğŸ» **ç©ºå¤´è§‚ç‚¹:**")
                            bear_argument = round_data.get("bear_argument", "æš‚æ— å†…å®¹")
                            st.write(bear_argument[:300] + ("..." if len(bear_argument) > 300 else ""))
                            if len(bear_argument) > 300:
                                if st.button(f"æŸ¥çœ‹ç¬¬{i}è½®ç©ºå¤´å®Œæ•´å‘è¨€", key=f"bear_full_{i}_{commodity}"):
                                    st.text_area(f"ç¬¬{i}è½®ç©ºå¤´å®Œæ•´å‘è¨€", bear_argument, height=200, key=f"bear_text_{i}_{commodity}")
                            
                            st.write(f"*å¾—åˆ†: {round_data.get('bear_score', 0):.1f}åˆ†*")
                            
                            # æœ¬è½®ç»“æœ
                            round_result = round_data.get('round_result', 'æœªçŸ¥')
                            audience_reaction = round_data.get('audience_reaction', 'æš‚æ— ')
                            st.write(f"**æœ¬è½®ç»“æœ**: {round_result}")
                            st.write(f"**è§‚ä¼—ååº”**: {audience_reaction}")
                            
                            if i < len(debate["detailed_rounds"]):
                                st.markdown("---")
                    else:
                        st.info("ğŸ’¡ è¯¦ç»†è¾©è®ºå†…å®¹æš‚æœªè®°å½•")
                
                # äº¤æ˜“å‘˜å†³ç­– - é‡æ–°è®¾è®¡çš„ä¸“ä¸šç•Œé¢
                if "trading_section" in debate_result:
                    st.markdown("#### ğŸ’¼ ä¸“ä¸šäº¤æ˜“å‘˜å†³ç­–")
                    trading = debate_result["trading_section"]
                    
                    # æ ¸å¿ƒå†³ç­–åŒºåŸŸ - ä½¿ç”¨metricç»„ä»¶å±•ç¤ºå…³é”®ä¿¡æ¯
                    st.markdown("##### ğŸ¯ æ ¸å¿ƒå†³ç­–")
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        strategy_type = trading.get('strategy_type', 'æœªçŸ¥')
                        st.metric("ç­–ç•¥ç±»å‹", strategy_type, help="åŸºäºå¤šç©ºè¾©è®ºç»“æœé€‰æ‹©çš„äº¤æ˜“ç­–ç•¥")
                    
                    with col2:
                        position_size = trading.get('position_size', 'æœªçŸ¥')
                        st.metric("å»ºè®®ä»“ä½", position_size, help="é£é™©æ§åˆ¶ä¸‹çš„æœ€ä¼˜ä»“ä½é…ç½®")
                    
                    with col3:
                        risk_reward = trading.get('risk_reward_ratio', 'æœªçŸ¥')
                        st.metric("é£é™©æ”¶ç›Šæ¯”", risk_reward, help="é¢„æœŸæ”¶ç›Šä¸é£é™©çš„æ¯”ä¾‹")
                    
                    
                    # äº¤æ˜“é€»è¾‘ - ä½¿ç”¨expanderç»„ç»‡
                    with st.expander("ğŸ§  äº¤æ˜“é€»è¾‘ä¸å†³ç­–ä¾æ®", expanded=True):
                        reasoning = trading.get('reasoning', 'æš‚æ— é€»è¾‘')
                        if reasoning and reasoning != 'æš‚æ— é€»è¾‘':
                            st.markdown(f"**å†³ç­–é€»è¾‘ï¼š**\n\n{reasoning}")
                        else:
                            st.warning("âš ï¸ äº¤æ˜“é€»è¾‘ä¿¡æ¯ä¸å®Œæ•´")
                    
                    # è¿›åœº/å‡ºåœºç­–ç•¥ - å¹¶åˆ—æ˜¾ç¤º
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("##### ğŸ“ˆ è¿›åœºç­–ç•¥")
                        entry_points = trading.get('entry_points', [])
                        if entry_points and isinstance(entry_points, list) and len(entry_points) > 0:
                            for i, point in enumerate(entry_points, 1):
                                if point and str(point).strip():
                                    st.success(f"**{i}.** {point}")
                        elif entry_points and not isinstance(entry_points, list):
                            st.success(f"**ç­–ç•¥ï¼š** {entry_points}")
                        else:
                            st.info("ğŸ“ è¿›åœºç­–ç•¥å¾…å®Œå–„")
                        
                    with col2:
                        st.markdown("##### ğŸ“‰ å‡ºåœºç­–ç•¥")
                        exit_points = trading.get('exit_points', [])
                        if exit_points and isinstance(exit_points, list) and len(exit_points) > 0:
                            for i, point in enumerate(exit_points, 1):
                                if point and str(point).strip():
                                    st.error(f"**{i}.** {point}")
                        elif exit_points and not isinstance(exit_points, list):
                            st.error(f"**ç­–ç•¥ï¼š** {exit_points}")
                        else:
                            st.info("ğŸ“ å‡ºåœºç­–ç•¥å¾…å®Œå–„")
                    
                    # æ‰§è¡Œè®¡åˆ’å’Œå¸‚åœºæ¡ä»¶ - ä½¿ç”¨tabsç»„ç»‡
                    tab1, tab2, tab3 = st.tabs(["ğŸ“Š äº¤æ˜“åˆçº¦", "ğŸ¯ æ‰§è¡Œè®¡åˆ’", "ğŸŒ å¸‚åœºæ¡ä»¶"])
                    
                    with tab1:
                        contracts = trading.get('specific_contracts', [])
                        if contracts and isinstance(contracts, list) and len(contracts) > 0:
                            for i, contract in enumerate(contracts, 1):
                                if contract and str(contract).strip():
                                    st.write(f"**{i}.** {contract}")
                        elif contracts and not isinstance(contracts, list):
                            st.write(f"**æ¨èåˆçº¦ï¼š** {contracts}")
                        else:
                            st.info("ğŸ“ å…·ä½“åˆçº¦å»ºè®®å¾…å®Œå–„")
                    
                    with tab2:
                        execution_plan = trading.get('execution_plan', '')
                        if execution_plan and str(execution_plan).strip() and execution_plan != 'æš‚æ— æ‰§è¡Œè®¡åˆ’':
                            st.markdown(f"**æ‰§è¡Œæ–¹æ¡ˆï¼š**\n\n{execution_plan}")
                        else:
                            st.info("ğŸ“ æ‰§è¡Œè®¡åˆ’å¾…å®Œå–„")
                    
                    with tab3:
                        market_conditions = trading.get('market_conditions', '')
                        if market_conditions and str(market_conditions).strip() and market_conditions != 'æš‚æ— å¸‚åœºæ¡ä»¶åˆ†æ':
                            st.markdown(f"**å¸‚åœºç¯å¢ƒï¼š**\n\n{market_conditions}")
                        else:
                            st.info("ğŸ“ å¸‚åœºæ¡ä»¶åˆ†æå¾…å®Œå–„")
                
                # é£é™©è¯„ä¼° - é‡æ–°è®¾è®¡çš„ä¸“ä¸šç•Œé¢
                if "risk_section" in debate_result:
                    st.markdown("#### ğŸ›¡ï¸ ä¸“ä¸šé£æ§è¯„ä¼°")
                    risk = debate_result["risk_section"]
                    
                    # é£é™©ç­‰çº§æ˜¾ç¤º - ç”¨é¢œè‰²åŒºåˆ†
                    overall_risk = risk.get('overall_risk', 'æœªçŸ¥')
                    risk_color = {
                        'ä½é£é™©': 'success',
                        'ä¸­é£é™©': 'warning', 
                        'é«˜é£é™©': 'error',
                        'æé«˜é£é™©': 'error'
                    }.get(overall_risk, 'info')
                    
                    # æ ¸å¿ƒé£é™©æŒ‡æ ‡
                    st.markdown("##### âš¡ é£é™©æŒ‡æ ‡")
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if overall_risk != 'æœªçŸ¥':
                            if risk_color == 'success':
                                st.success(f"**é£é™©ç­‰çº§**\n{overall_risk}")
                            elif risk_color == 'warning':
                                st.warning(f"**é£é™©ç­‰çº§**\n{overall_risk}")
                            else:
                                st.error(f"**é£é™©ç­‰çº§**\n{overall_risk}")
                        else:
                            st.info(f"**é£é™©ç­‰çº§**\n{overall_risk}")
                    
                    with col2:
                        position_limit = risk.get('position_limit', 'æœªçŸ¥')
                        st.metric("å»ºè®®ä»“ä½", position_limit, help="åŸºäºé£é™©ç­‰çº§çš„ä»“ä½ä¸Šé™å»ºè®®")
                    
                    with col3:
                        stop_loss = risk.get('stop_loss', 'æœªçŸ¥')
                        st.metric("æ­¢æŸå»ºè®®", stop_loss, help="åŸºäºæŠ€æœ¯åˆ†æçš„æ­¢æŸä½å»ºè®®")
                    
                    # è¯¦ç»†é£é™©åˆ†æ
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("##### ğŸ¯ é£é™©å› ç´ ")
                        key_factors = risk.get('key_factors', [])
                        if key_factors and isinstance(key_factors, list):
                            for i, factor in enumerate(key_factors, 1):
                                if factor and str(factor).strip():
                                    st.warning(f"**{i}.** {factor}")
                        else:
                            st.info("ğŸ“ é£é™©å› ç´ è¯†åˆ«å¾…å®Œå–„")
                    
                    with col2:
                        st.markdown("##### ğŸ›¡ï¸ ç¼“é‡Šæªæ–½")
                        mitigation = risk.get('mitigation_measures', [])
                        if mitigation and isinstance(mitigation, list):
                            for i, measure in enumerate(mitigation, 1):
                                if measure and str(measure).strip():
                                    st.success(f"**{i}.** {measure}")
                        else:
                            st.info("ğŸ“ ç¼“é‡Šæªæ–½å»ºè®®å¾…å®Œå–„")
                    
                    # é£æ§ç»ç†ä¸“ä¸šæ„è§
                    with st.expander("ğŸ‘¨â€ğŸ’¼ é£æ§æ€»ç›‘ä¸“ä¸šæ„è§", expanded=True):
                        manager_opinion = risk.get('manager_opinion', 'æš‚æ— æ„è§')
                        if manager_opinion and manager_opinion != 'æš‚æ— æ„è§':
                            st.markdown(f"**ä¸“ä¸šè¯„ä¼°ï¼š**\n\n{manager_opinion}")
                        else:
                            st.warning("âš ï¸ é£æ§æ€»ç›‘æ„è§ä¸å®Œæ•´")
                
                # ğŸ”¥ CIOå†³ç­–ç°åœ¨ç”±ç‹¬ç«‹çš„ã€ŒCIOå†³ç­–ã€æ ‡ç­¾é¡µæ˜¾ç¤ºï¼Œé¿å…é‡å¤å†…å®¹
                st.info("ğŸ’¡ **æŸ¥çœ‹CIOæœ€ç»ˆæƒå¨å†³ç­–ï¼Œè¯·ç‚¹å‡»ä¸Šæ–¹ã€ŒCIOå†³ç­–ã€æ ‡ç­¾é¡µ**")
            
            else:
                st.error(f"å®Œæ•´å†³ç­–æµç¨‹å¤±è´¥: {debate_result.get('error', 'æœªçŸ¥é”™è¯¯')}")
    
    def _display_analyst_only_result(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºä»…åˆ†æå¸ˆç»“æœ"""
        
        st.success(f"âœ… {commodity} åˆ†æå¸ˆå›¢é˜Ÿåˆ†æå®Œæˆ")
        
        analysis_state = result.get("result")
        if analysis_state:
            st.write("ğŸ“Š 6å¤§åˆ†ææ¨¡å—å·²å®Œæˆåˆ†æ")
            
            # æ˜¾ç¤ºå„æ¨¡å—çŠ¶æ€
            modules = {
                "inventory_analysis": "åº“å­˜ä»“å•åˆ†æ",
                "positioning_analysis": "æŒä»“å¸­ä½åˆ†æ",
                "term_structure_analysis": "æœŸé™ç»“æ„åˆ†æ",
                "technical_analysis": "æŠ€æœ¯é¢åˆ†æ",
                "basis_analysis": "åŸºå·®åˆ†æ",
                "news_analysis": "æ–°é—»åˆ†æ"
            }
            
            col1, col2, col3 = st.columns(3)
            for i, (module_attr, module_name) in enumerate(modules.items()):
                with [col1, col2, col3][i % 3]:
                    module_result = getattr(analysis_state, module_attr, None) if analysis_state else None
                    
                    if module_result and hasattr(module_result, 'status'):
                        # è·å–çŠ¶æ€å€¼
                        status = module_result.status
                        
                        # æ”¯æŒå¤šç§çŠ¶æ€æ ¼å¼çš„åˆ¤æ–­
                        is_completed = False
                        
                        # æ£€æŸ¥å„ç§å¯èƒ½çš„çŠ¶æ€æ ¼å¼
                        if hasattr(status, 'value'):
                            # æšä¸¾ç±»å‹ï¼Œå¦‚ AnalysisStatus.COMPLETED
                            status_value = status.value.lower() if isinstance(status.value, str) else str(status.value).lower()
                            is_completed = status_value in ["completed", "success"]
                        elif isinstance(status, str):
                            # å­—ç¬¦ä¸²ç±»å‹
                            is_completed = status.lower() in ["completed", "success"]
                        else:
                            # å…¶ä»–ç±»å‹ï¼Œè½¬æ¢ä¸ºå­—ç¬¦ä¸²æ¯”è¾ƒ
                            status_str = str(status).lower()
                            is_completed = status_str in ["completed", "success"]
                        
                        # é¢å¤–æ£€æŸ¥ï¼šå¦‚æœæœ‰result_dataä¸”åŒ…å«successæ ‡å¿—
                        if not is_completed and hasattr(module_result, 'result_data') and module_result.result_data:
                            if isinstance(module_result.result_data, dict):
                                is_completed = module_result.result_data.get('success', False)
                        
                        status_icon = "âœ…" if is_completed else "âŒ"
                        st.write(f"{status_icon} {module_name}")
                        
                        # å¦‚æœåˆ†æå®Œæˆï¼Œæ˜¾ç¤ºåˆ†æå†…å®¹æ‘˜è¦
                        if is_completed and hasattr(module_result, 'result_data') and module_result.result_data:
                            result_data = module_result.result_data
                            analysis_content = None
                            
                            # å°è¯•ä»å¤šä¸ªå¯èƒ½çš„å­—æ®µè·å–åˆ†æå†…å®¹
                            if isinstance(result_data, dict):
                                # å°è¯•ä¸åŒçš„å†…å®¹å­—æ®µï¼ˆä¼˜å…ˆåº“å­˜åˆ†æå­—æ®µï¼‰
                                content_fields = ["ai_comprehensive_analysis", "ai_analysis", "analysis_content", "speculative_analysis", "content", "report", "analysis"]
                                for field in content_fields:
                                    if field in result_data and result_data[field]:
                                        analysis_content = result_data[field]
                                        break
                                
                                # å¦‚æœæ˜¯å­—å…¸ï¼Œå°è¯•è·å–åµŒå¥—çš„å†…å®¹
                                if not analysis_content:
                                    for key, value in result_data.items():
                                        if isinstance(value, dict) and "analysis_content" in value:
                                            analysis_content = value["analysis_content"]
                                            break
                                        elif isinstance(value, dict) and "ai_analysis" in value:
                                            analysis_content = value["ai_analysis"]
                                            break
                            
                            # æ˜¾ç¤ºåˆ†æå†…å®¹
                            if analysis_content and isinstance(analysis_content, str) and len(analysis_content.strip()) > 0:
                                # æ˜¾ç¤ºæ‘˜è¦ï¼ˆå‰200å­—ç¬¦ï¼‰
                                summary = analysis_content[:200] + ("..." if len(analysis_content) > 200 else "")
                                st.caption(f"æ‘˜è¦: {summary}")
                                
                                # æä¾›æŸ¥çœ‹å®Œæ•´æŠ¥å‘Šçš„æŒ‰é’®
                                if len(analysis_content) > 200:
                                    if st.button(f"æŸ¥çœ‹{module_name}å®Œæ•´æŠ¥å‘Š", key=f"view_{module_attr}_{commodity}"):
                                        st.text_area(
                                            f"{module_name}å®Œæ•´åˆ†ææŠ¥å‘Š", 
                                            analysis_content, 
                                            height=400, 
                                            key=f"content_{module_attr}_{commodity}"
                                        )
                        
                        # è°ƒè¯•ä¿¡æ¯ï¼ˆä»…åœ¨å¼€å‘æ¨¡å¼ä¸‹æ˜¾ç¤ºï¼‰
                        if st.session_state.get('debug_mode', False):
                            st.caption(f"è°ƒè¯•: {module_attr} = {status} ({type(status)})")
                            if hasattr(module_result, 'result_data'):
                                result_data = module_result.result_data
                                if result_data:
                                    if isinstance(result_data, dict):
                                        st.caption(f"è°ƒè¯•: result_dataé”® = {list(result_data.keys())}")
                                        # æ˜¾ç¤ºå‰å‡ ä¸ªå­—æ®µçš„å†…å®¹é•¿åº¦
                                        for key, value in list(result_data.items())[:3]:
                                            if isinstance(value, str):
                                                st.caption(f"è°ƒè¯•: {key} = {len(value)}å­—ç¬¦")
                                            else:
                                                st.caption(f"è°ƒè¯•: {key} = {type(value)}")
                                    else:
                                        st.caption(f"è°ƒè¯•: result_dataç±»å‹ = {type(result_data)}")
                                else:
                                    st.caption(f"è°ƒè¯•: result_data = None")
                            else:
                                st.caption(f"è°ƒè¯•: æ— result_dataå±æ€§")
                            
                    else:
                        st.write(f"âš ï¸ {module_name}")
                        if st.session_state.get('debug_mode', False):
                            st.caption(f"è°ƒè¯•: {module_attr} = None æˆ–æ— statuså±æ€§")
        else:
            st.write("âš ï¸ åˆ†æç»“æœæ•°æ®ä¸å®Œæ•´")
    
    def _display_complete_flow_result_paginated(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºå®Œæ•´æµç¨‹ç»“æœ - æ ‡ç­¾é¡µç‰ˆæœ¬"""
        
        # åˆ›å»ºæ°´å¹³æ ‡ç­¾é¡µ
        tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
            "ğŸ“Š æ‰§è¡Œæ‘˜è¦", 
            "ğŸ‘¨â€ğŸ’¼ åˆ†æå¸ˆå›¢é˜Ÿ", 
            "ğŸ­ æ¿€çƒˆè¾©è®º", 
            "ğŸ’¼ ä¸“ä¸šäº¤æ˜“å‘˜", 
            "ğŸ›¡ï¸ ä¸“ä¸šé£æ§", 
            "ğŸ‘” CIOå†³ç­–"
        ])
        
        # åœ¨æ¯ä¸ªæ ‡ç­¾é¡µä¸­æ˜¾ç¤ºå¯¹åº”å†…å®¹
        with tab1:
            self._display_execution_summary(commodity, result)
        
        with tab2:
            self._display_analyst_team_page(commodity, result)
        
        with tab3:
            # ä¼ é€’å®Œæ•´çš„è¾©è®ºé£æ§å†³ç­–ç»“æœ
            debate_result = result.get("result", {}).get("debate_result", {})
            self._display_debate_page(commodity, debate_result)
        
        with tab4:
            # ä¼ é€’å®Œæ•´çš„è¾©è®ºé£æ§å†³ç­–ç»“æœ  
            debate_result = result.get("result", {}).get("debate_result", {})
            self._display_trader_page(commodity, debate_result)
        
        with tab5:
            # ä¼ é€’å®Œæ•´çš„è¾©è®ºé£æ§å†³ç­–ç»“æœ
            debate_result = result.get("result", {}).get("debate_result", {})
            self._display_risk_page(commodity, debate_result)
        
        with tab6:
            # ä¼ é€’å®Œæ•´çš„è¾©è®ºé£æ§å†³ç­–ç»“æœ
            debate_result = result.get("result", {}).get("debate_result", {})
            self._display_cio_decision_page(commodity, debate_result)
    
    def _display_analyst_only_result_paginated(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºä»…åˆ†æå¸ˆç»“æœ - æ ‡ç­¾é¡µç‰ˆæœ¬"""
        
        # åˆ›å»ºæ°´å¹³æ ‡ç­¾é¡µ
        tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
            "ğŸ“Š æ€»è§ˆ", 
            "ğŸ“¦ åº“å­˜ä»“å•", 
            "ğŸ“ˆ æŠ€æœ¯åˆ†æ", 
            "ğŸ¯ æŒä»“å¸­ä½", 
            "ğŸ’° åŸºå·®åˆ†æ", 
            "ğŸ“Š æœŸé™ç»“æ„", 
            "ğŸ“° æ–°é—»åˆ†æ"
        ])
        
        # åœ¨æ¯ä¸ªæ ‡ç­¾é¡µä¸­æ˜¾ç¤ºå¯¹åº”å†…å®¹
        with tab1:
            self._display_analyst_overview(commodity, result)
        
        with tab2:
            self._display_inventory_analysis_page(commodity, result)
        
        with tab3:
            self._display_technical_analysis_page(commodity, result)
        
        with tab4:
            self._display_positioning_analysis_page(commodity, result)
        
        with tab5:
            self._display_basis_analysis_page(commodity, result)
        
        with tab6:
            self._display_term_structure_analysis_page(commodity, result)
        
        with tab7:
            self._display_news_analysis_page(commodity, result)
    
    def _display_execution_summary(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºæ‰§è¡Œæ‘˜è¦é¡µé¢"""
        st.markdown("## ğŸ“Š æ‰§è¡Œæ‘˜è¦")
        
        execution_summary = result.get("execution_summary", {})
        
        # æœ€ç»ˆå†³ç­–æ˜¾ç¤º - ğŸ”¥ ä¿®å¤ï¼šä»æ­£ç¡®çš„å±‚çº§è·å–æ•°æ®
        debate_result = result.get("result", {}).get("debate_result", {})
        
        # æ£€æŸ¥åˆ†ææ˜¯å¦æˆåŠŸ
        if not debate_result or not debate_result.get("success"):
            error_msg = debate_result.get("error", "æœªçŸ¥é”™è¯¯") if debate_result else "åˆ†æç»“æœä¸ºç©º"
            st.error(f"âŒ åˆ†æå¤±è´¥: {error_msg}")
            st.info("ğŸ’¡ æç¤ºï¼šå¯èƒ½æ˜¯ç½‘ç»œè¿æ¥é—®é¢˜ï¼Œè¯·æ£€æŸ¥ç½‘ç»œåé‡è¯•")
            return
        
        # ä»æ­£ç¡®çš„ä½ç½®è·å–CIOå†³ç­–æ•°æ®
        executive_decision = debate_result.get("decision_section", {})
        
        if not executive_decision:
            st.error("âŒ CIOå†³ç­–æ•°æ®ç¼ºå¤±ï¼Œè¯·é‡æ–°è¿è¡Œå®Œæ•´åˆ†ææµç¨‹")
            return
        
        # è·å–æ•°æ®
        operational_decision = executive_decision.get("operational_decision", "ä¸­æ€§ç­–ç•¥")
        directional_view = executive_decision.get("directional_view", "ä¸­æ€§")
        operational_confidence = executive_decision.get("operational_confidence", "ä¸­")
        directional_confidence = executive_decision.get("directional_confidence", "ä¸­")
        
        # æå–ä»·æ ¼ä¿¡æ¯ - ğŸ”¥ ä¼˜å…ˆä»CIOçš„statementè·å–ï¼ˆå·²æ ¼å¼åŒ–ï¼‰ï¼Œå†ä»trading_sectionè·å–
        import re
        
        # ä¼˜å…ˆä»CIOçš„statementä¸­æå–ï¼ˆæ ¼å¼åŒ–å¥½çš„ï¼‰
        cio_statement = executive_decision.get("cio_statement", "")
        trading_section = debate_result.get("trading_section", {})
        trader_reasoning = trading_section.get("reasoning", "")
        
        # åˆå¹¶ä¸¤ä¸ªæ–‡æœ¬æºï¼ŒCIOçš„statementä¼˜å…ˆ
        combined_text = cio_statement + "\n\n" + trader_reasoning
        
        # å½“å‰ä»·æ ¼ï¼ˆå¯é€‰ï¼Œå¦‚æœæ²¡æœ‰å°±ä¸æ˜¾ç¤ºï¼‰
        current_price_match = re.search(r'å½“å‰ä»·æ ¼[ï¼š:]\s*(\d+\.?\d*)å…ƒ', combined_text)
        current_price = f"{current_price_match.group(1)}å…ƒ" if current_price_match else None
        
        # ğŸ”¥ æ›´æ–°æ­£åˆ™è¡¨è¾¾å¼ï¼Œæ”¯æŒæ›´å¤šæ ¼å¼ï¼ˆä»combined_textæå–ï¼‰
        # æ ¼å¼1ï¼šè¿›åœºï¼š82800-83100å…ƒ æˆ– è¿›åœºåŒºé—´ï¼š82800-83100å…ƒ
        # æ ¼å¼2ï¼šè¿›åœº82800-83100å…ƒ
        # æ ¼å¼3ï¼šå…³é”®ä»·ä½ï¼šè¿›åœº82800-83100å…ƒ
        # æ ¼å¼4ï¼šå»ºè®®è¿›åœºåŒºé—´ï¼š12300-12400å…ƒï¼ˆCIOå†³ç­–æ ¼å¼ï¼‰ â­
        entry_range_match = re.search(r'å»ºè®®è¿›åœºåŒºé—´[ï¼š:]?\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', combined_text)
        if not entry_range_match:
            entry_range_match = re.search(r'è¿›åœº[åŒºé—´]?[ï¼š:]?\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', combined_text)
        if not entry_range_match:
            # åå¤‡æ–¹æ¡ˆï¼šä»"å…³é”®ä»·ä½"å¼€å§‹æœç´¢
            key_price_section = re.search(r'å…³é”®ä»·ä½[ï¼š:](.*?)(?=ï¼Œ|ã€‚|æŒä»“å‘¨æœŸ|é£é™©æ§åˆ¶|$)', combined_text, re.DOTALL)
            if key_price_section:
                entry_range_match = re.search(r'è¿›åœº[åŒºé—´]?\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', key_price_section.group(1))
        entry_range = f"{entry_range_match.group(1)}-{entry_range_match.group(2)}å…ƒ" if entry_range_match else "æœªæä¾›"
        
        # æ­¢æŸä½ - ğŸ”¥ æ‰©å±•å¤šç§æ ¼å¼ï¼ˆä»combined_textæå–ï¼‰
        # æ ¼å¼1ï¼šæ­¢æŸä½ï¼š>53148å…ƒï¼ˆCIOå†³ç­–æ ¼å¼ï¼‰ â­
        # æ ¼å¼2ï¼šæ­¢æŸ<81500å…ƒ æˆ– æ­¢æŸï¼š<81500å…ƒ
        # æ ¼å¼3ï¼šæ­¢æŸä½81500å…ƒé™„è¿‘
        # æ ¼å¼4ï¼šæ­¢æŸ>12800å…ƒï¼ˆåšç©ºæ—¶ï¼‰
        stop_loss_match = re.search(r'æ­¢æŸ[ä½]?[ï¼š:]?\s*([<ï¼œ>ï¼])\s*(\d+\.?\d*)å…ƒ', combined_text)
        if stop_loss_match:
            direction_symbol = '<' if stop_loss_match.group(1) in ['<', 'ï¼œ'] else '>'
            stop_loss = f"{direction_symbol}{stop_loss_match.group(2)}å…ƒ"
        else:
            # å°è¯•å…¶ä»–æ ¼å¼
            stop_loss_match2 = re.search(r'æ­¢æŸ[ä½]?[ï¼š:]?\s*(\d+\.?\d*)å…ƒ?\s*[ä»¥ä¸Šä»¥ä¸‹é™„è¿‘]', combined_text)
            if stop_loss_match2:
                # åˆ¤æ–­æ˜¯ä»¥ä¸Šè¿˜æ˜¯ä»¥ä¸‹
                context = combined_text[max(0, stop_loss_match2.start()-50):stop_loss_match2.end()+50]
                if 'ä»¥ä¸Š' in context or 'åšç©º' in operational_decision:
                    stop_loss = f">{stop_loss_match2.group(1)}å…ƒ"
                elif 'ä»¥ä¸‹' in context or 'åšå¤š' in operational_decision:
                    stop_loss = f"<{stop_loss_match2.group(1)}å…ƒ"
                else:
                    stop_loss = f"~{stop_loss_match2.group(1)}å…ƒ"
            else:
                stop_loss = "æœªæä¾›"
        
        # ç›®æ ‡ä»·ä½ - ä»combined_textæå–
        # æ ¼å¼1ï¼šç›®æ ‡ä»·ä½ï¼š48000-49000å…ƒï¼ˆCIOå†³ç­–æ ¼å¼ï¼‰ â­
        # æ ¼å¼2ï¼šç›®æ ‡84500-85000å…ƒ æˆ– ç›®æ ‡ï¼š84500-85000å…ƒ
        target_price_match = re.search(r'ç›®æ ‡[ä»·ä½]?[ï¼š:]?\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', combined_text)
        if not target_price_match:
            # åå¤‡æ–¹æ¡ˆï¼šä»"å…³é”®ä»·ä½"æœç´¢
            key_price_section = re.search(r'å…³é”®ä»·ä½[ï¼š:](.*?)(?=ï¼ŒæŒä»“å‘¨æœŸ|é£é™©æ§åˆ¶|$)', combined_text, re.DOTALL)
            if key_price_section:
                target_price_match = re.search(r'ç›®æ ‡\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', key_price_section.group(1))
        target_price = f"{target_price_match.group(1)}-{target_price_match.group(2)}å…ƒ" if target_price_match else "æœªæä¾›"
        
        # è·å–åˆ†ææ¨¡å‹
        analysis_model = execution_summary.get('ai_model', 'DeepSeek-Reasoner')
        
        # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
        # ç°ä»£åŒ–å¡ç‰‡å¼å¸ƒå±€
        # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
        
        # 1. å†³ç­–ç»“æœå¡ç‰‡
        with st.container():
            st.markdown('<div class="summary-section">', unsafe_allow_html=True)
            st.markdown("#### ğŸ¯ å†³ç­–ç»“æœ")
            
            col1, col2 = st.columns(2, gap="medium")
            
            with col1:
                st.markdown("""
                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                           border-radius: 12px; padding: 20px; margin: 10px 0;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1);">
                    <div style="color: white; opacity: 0.9; font-size: 0.9em; margin-bottom: 5px;">ğŸ“Š æ“ä½œå†³ç­–</div>
                    <div style="color: white; font-size: 1.8em; font-weight: bold;">{}</div>
                    <div style="color: white; opacity: 0.8; font-size: 0.85em; margin-top: 10px; padding-top: 10px; border-top: 1px solid rgba(255,255,255,0.3);">
                        ğŸ² æ“ä½œä¿¡å¿ƒåº¦ï¼š<span style="font-weight: bold;">{}</span>
                    </div>
                </div>
                """.format(operational_decision, operational_confidence), unsafe_allow_html=True)
            
            with col2:
                st.markdown("""
                <div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
                           border-radius: 12px; padding: 20px; margin: 10px 0;
                           box-shadow: 0 4px 15px rgba(0,0,0,0.1);">
                    <div style="color: white; opacity: 0.9; font-size: 0.9em; margin-bottom: 5px;">ğŸ§­ æ–¹å‘åˆ¤æ–­</div>
                    <div style="color: white; font-size: 1.8em; font-weight: bold;">{}</div>
                    <div style="color: white; opacity: 0.8; font-size: 0.85em; margin-top: 10px; padding-top: 10px; border-top: 1px solid rgba(255,255,255,0.3);">
                        ğŸ“ æ–¹å‘ä¿¡å¿ƒåº¦ï¼š<span style="font-weight: bold;">{}</span>
                    </div>
                </div>
                """.format(directional_view, directional_confidence), unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        # æ·»åŠ é—´éš”
        st.markdown("<br>", unsafe_allow_html=True)
        
        # 2. å…³é”®ä»·ä½å¡ç‰‡
        with st.container():
            st.markdown('<div class="summary-section">', unsafe_allow_html=True)
            st.markdown("#### ğŸ’° å…³é”®ä»·ä½")
            
            # æ ¹æ®æ˜¯å¦æœ‰å½“å‰ä»·æ ¼å†³å®šåˆ—æ•°
            if current_price:
                col1, col2, col3, col4 = st.columns(4, gap="small")
                with col1:
                    st.markdown("""
                    <div style="background: white; border-left: 4px solid #3b82f6;
                               border-radius: 8px; padding: 15px; margin: 5px 0;
                               box-shadow: 0 2px 8px rgba(0,0,0,0.08);">
                        <div style="color: #6b7280; font-size: 0.85em; margin-bottom: 5px;">ğŸ“ å½“å‰ä»·æ ¼</div>
                        <div style="color: #1f2937; font-size: 1.5em; font-weight: bold;">{}</div>
                    </div>
                    """.format(current_price), unsafe_allow_html=True)
            else:
                col2, col3, col4 = st.columns(3, gap="medium")
            
            with col2:
                st.markdown("""
                <div style="background: white; border-left: 4px solid #10b981;
                           border-radius: 8px; padding: 15px; margin: 5px 0;
                           box-shadow: 0 2px 8px rgba(0,0,0,0.08);">
                    <div style="color: #6b7280; font-size: 0.85em; margin-bottom: 5px;">ğŸ¯ è¿›åœºåŒºé—´</div>
                    <div style="color: #1f2937; font-size: 1.5em; font-weight: bold;">{}</div>
                </div>
                """.format(entry_range), unsafe_allow_html=True)
            
            with col3:
                st.markdown("""
                <div style="background: white; border-left: 4px solid #f59e0b;
                           border-radius: 8px; padding: 15px; margin: 5px 0;
                           box-shadow: 0 2px 8px rgba(0,0,0,0.08);">
                    <div style="color: #6b7280; font-size: 0.85em; margin-bottom: 5px;">ğŸ–ï¸ ç›®æ ‡ä»·ä½</div>
                    <div style="color: #1f2937; font-size: 1.5em; font-weight: bold;">{}</div>
                </div>
                """.format(target_price), unsafe_allow_html=True)
            
            with col4:
                st.markdown("""
                <div style="background: white; border-left: 4px solid #ef4444;
                           border-radius: 8px; padding: 15px; margin: 5px 0;
                           box-shadow: 0 2px 8px rgba(0,0,0,0.08);">
                    <div style="color: #6b7280; font-size: 0.85em; margin-bottom: 5px;">ğŸ›¡ï¸ æ­¢æŸä»·ä½</div>
                    <div style="color: #1f2937; font-size: 1.5em; font-weight: bold;">{}</div>
                </div>
                """.format(stop_loss), unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        # æ·»åŠ é—´éš”
        st.markdown("<br>", unsafe_allow_html=True)
        
        # 3. æ‰§è¡Œä¿¡æ¯å¡ç‰‡
        with st.container():
            st.markdown('<div class="summary-section">', unsafe_allow_html=True)
            st.markdown("#### â±ï¸ æ‰§è¡Œä¿¡æ¯")
            
            # ä½¿ç”¨å•è¡Œ4åˆ—å¸ƒå±€ï¼ˆç§»é™¤åˆ†æè€—æ—¶ï¼‰
            st.markdown("""
            <div style="background: linear-gradient(135deg, #e0f7fa 0%, #b2ebf2 100%);
                       border-radius: 12px; padding: 20px; margin: 10px 0;
                       box-shadow: 0 2px 10px rgba(0,0,0,0.08);">
                <div style="display: flex; justify-content: space-around; align-items: center;">
                    <div style="text-align: center;">
                        <div style="color: #00796b; font-size: 0.85em; margin-bottom: 5px;">å®Œæˆé˜¶æ®µ</div>
                        <div style="color: #004d40; font-size: 1.8em; font-weight: bold;">{}/5</div>
                    </div>
                    <div style="border-left: 2px solid rgba(0,121,107,0.2); height: 50px;"></div>
                    <div style="text-align: center;">
                        <div style="color: #00796b; font-size: 0.85em; margin-bottom: 5px;">åˆ†ææ¨¡å—</div>
                        <div style="color: #004d40; font-size: 1.8em; font-weight: bold;">{}ä¸ª</div>
                    </div>
                    <div style="border-left: 2px solid rgba(0,121,107,0.2); height: 50px;"></div>
                    <div style="text-align: center;">
                        <div style="color: #00796b; font-size: 0.85em; margin-bottom: 5px;">è¾©è®ºè½®æ•°</div>
                        <div style="color: #004d40; font-size: 1.8em; font-weight: bold;">{}è½®</div>
                    </div>
                    <div style="border-left: 2px solid rgba(0,121,107,0.2); height: 50px;"></div>
                    <div style="text-align: center;">
                        <div style="color: #00796b; font-size: 0.85em; margin-bottom: 5px;">åˆ†ææ¨¡å‹</div>
                        <div style="color: #004d40; font-size: 1.3em; font-weight: bold;">{}</div>
                    </div>
                </div>
            </div>
            """.format(
                execution_summary.get('stages_completed', 0),
                execution_summary.get('analysis_modules', 0),
                execution_summary.get('debate_rounds', 0),
                analysis_model
            ), unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    def _display_analyst_team_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºåˆ†æå¸ˆå›¢é˜Ÿé¡µé¢"""
        st.markdown("## ğŸ‘¨â€ğŸ’¼ åˆ†æå¸ˆå›¢é˜ŸæŠ¥å‘Š")
        
        analysis_result = result.get("result", {})
        analysis_state = analysis_result.get("analysis_state")
        
        if not analysis_state:
            st.error("âŒ æ— æ³•è·å–åˆ†æå¸ˆå›¢é˜Ÿæ•°æ®")
            return
        
        # å›¢é˜Ÿæˆæœæ¦‚è§ˆ
        st.markdown("### ğŸ“Š å›¢é˜Ÿæˆæœæ¦‚è§ˆ")
        
        module_attrs = [
            ("inventory_analysis", "ğŸ“¦ åº“å­˜ä»“å•åˆ†æå¸ˆ"),
            ("positioning_analysis", "ğŸ¯ æŒä»“å¸­ä½åˆ†æå¸ˆ"),
            ("technical_analysis", "ğŸ“ˆ æŠ€æœ¯é¢åˆ†æå¸ˆ"),
            ("basis_analysis", "ğŸ’° åŸºå·®åˆ†æå¸ˆ"),
            ("term_structure_analysis", "ğŸ“Š æœŸé™ç»“æ„åˆ†æå¸ˆ"),
            ("news_analysis", "ğŸ“° æ–°é—»åˆ†æå¸ˆ")
        ]
        
        successful_modules = 0
        for attr_name, module_name in module_attrs:
            if hasattr(analysis_state, attr_name):
                module_result = getattr(analysis_state, attr_name)
                if module_result and (
                    module_result.status == "completed" or 
                    (hasattr(module_result.status, 'value') and module_result.status.value == "completed") or
                    str(module_result.status).lower() == "completed"
                ):
                    successful_modules += 1
        
        st.success(f"âœ… {successful_modules}/6 ä½åˆ†æå¸ˆå®Œæˆä¸“ä¸šæŠ¥å‘Š")
        
        # å„åˆ†æå¸ˆè¯¦ç»†æŠ¥å‘Šæ ‡ç­¾é¡µ
        st.markdown("### ğŸ“‹ åˆ†æå¸ˆè¯¦ç»†æŠ¥å‘Š")
        
        tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
            "ğŸ“¦ åº“å­˜ä»“å•", 
            "ğŸ¯ æŒä»“å¸­ä½", 
            "ğŸ“ˆ æŠ€æœ¯åˆ†æ", 
            "ğŸ’° åŸºå·®åˆ†æ", 
            "ğŸ“Š æœŸé™ç»“æ„", 
            "ğŸ“° æ–°é—»åˆ†æ"
        ])
        
        with tab1:
            self._display_single_analyst_report(commodity, analysis_state, 'inventory_analysis', 'ğŸ“¦ åº“å­˜ä»“å•åˆ†æå¸ˆ')
        
        with tab2:
            self._display_single_analyst_report(commodity, analysis_state, 'positioning_analysis', 'ğŸ¯ æŒä»“å¸­ä½åˆ†æå¸ˆ')
        
        with tab3:
            self._display_single_analyst_report(commodity, analysis_state, 'technical_analysis', 'ğŸ“ˆ æŠ€æœ¯é¢åˆ†æå¸ˆ')
        
        with tab4:
            self._display_single_analyst_report(commodity, analysis_state, 'basis_analysis', 'ğŸ’° åŸºå·®åˆ†æå¸ˆ')
        
        with tab5:
            self._display_single_analyst_report(commodity, analysis_state, 'term_structure_analysis', 'ğŸ“Š æœŸé™ç»“æ„åˆ†æå¸ˆ')
        
        with tab6:
            self._display_single_analyst_report(commodity, analysis_state, 'news_analysis', 'ğŸ“° æ–°é—»åˆ†æå¸ˆ')
    
    def _display_single_analyst_report(self, commodity: str, analysis_state, attr_name: str, module_name: str):
        """æ˜¾ç¤ºå•ä¸ªåˆ†æå¸ˆçš„å®Œæ•´æŠ¥å‘Š"""
        if not hasattr(analysis_state, attr_name):
            st.error(f"âŒ æœªæ‰¾åˆ°{module_name}çš„åˆ†æç»“æœ")
            return
        
        module_result = getattr(analysis_state, attr_name)
        
        # åˆ†æå¸ˆä¿¡æ¯å¡ç‰‡
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                   color: white; padding: 20px; border-radius: 15px; margin: 20px 0;">
            <h2 style="margin: 0; text-align: center;">{module_name}</h2>
            <p style="margin: 5px 0 0 0; text-align: center; font-size: 1.1em;">ä¸“ä¸šæ·±åº¦åˆ†ææŠ¥å‘Š</p>
        </div>
        """, unsafe_allow_html=True)
        
        # æ£€æŸ¥åˆ†æçŠ¶æ€
        is_completed = False
        if module_result and hasattr(module_result, 'status'):
            status = module_result.status
            if hasattr(status, 'value'):
                status_value = status.value.lower() if isinstance(status.value, str) else str(status.value).lower()
                is_completed = status_value in ["completed", "success"]
            elif isinstance(status, str):
                is_completed = status.lower() in ["completed", "success"]
            else:
                status_str = str(status).lower()
                is_completed = status_str in ["completed", "success"]
        
        if not is_completed:
            st.error(f"âŒ {module_name}åˆ†ææœªå®Œæˆ")
            return
        
        # è·å–åˆ†æå†…å®¹
        result_data = module_result.result_data if hasattr(module_result, 'result_data') else None
        if not result_data:
            st.error(f"âŒ {module_name}æ— åˆ†ææ•°æ®")
            return
        
        # åˆ†æåŸºæœ¬ä¿¡æ¯
        col1, col2 = st.columns(2)
        with col1:
            if "confidence_score" in result_data:
                confidence = result_data['confidence_score']
                if isinstance(confidence, (int, float)):
                    # è½¬æ¢ä¸ºé«˜ä¸­ä½çº§åˆ«
                    if confidence >= 0.8:
                        confidence_level = "é«˜"
                    elif confidence >= 0.6:
                        confidence_level = "ä¸­"
                    else:
                        confidence_level = "ä½"
                    st.metric("åˆ†æä¿¡å¿ƒ", confidence_level)
                else:
                    st.metric("åˆ†æä¿¡å¿ƒ", "é«˜")
            else:
                st.metric("åˆ†æä¿¡å¿ƒ", "é«˜")
        
        with col2:
            # åˆ†ææ—¶é—´ - æ ¼å¼åŒ–ä¸ºä¸­æ–‡æ—¶é—´
            time_str = "å®æ—¶"
            if "analysis_time" in result_data:
                time_str = result_data['analysis_time']
            elif "timestamp" in result_data:
                time_str = result_data['timestamp']
            
            # å°è¯•æ ¼å¼åŒ–æ—¶é—´ä¸ºä¸­æ–‡æ ¼å¼
            try:
                if isinstance(time_str, str) and len(time_str) > 10:
                    from datetime import datetime
                    # å°è¯•è§£æä¸åŒçš„æ—¶é—´æ ¼å¼
                    for fmt in ['%Y-%m-%d %H:%M:%S', '%Y/%m/%d %H:%M:%S', '%Y-%m-%d']:
                        try:
                            dt = datetime.strptime(time_str, fmt)
                            time_str = f"{dt.year}å¹´{dt.month}æœˆ{dt.day}æ—¥ {dt.hour}æ—¶{dt.minute}åˆ†{dt.second}ç§’"
                            break
                        except:
                            continue
            except:
                pass
            
            st.metric("åˆ†ææ—¶é—´", time_str)
        
        # æ ¸å¿ƒåˆ†æå†…å®¹
        st.markdown("### ğŸ“‹ ä¸“ä¸šåˆ†æå†…å®¹")
        
        analysis_content = None
        # å°è¯•ä»å¤šä¸ªå¯èƒ½çš„å­—æ®µè·å–åˆ†æå†…å®¹
        content_fields = ["ai_comprehensive_analysis", "ai_analysis", "analysis_content", "speculative_analysis", "content", "report", "analysis"]
        for field in content_fields:
            if field in result_data and result_data[field]:
                analysis_content = result_data[field]
                break
        
        if analysis_content and isinstance(analysis_content, str) and len(analysis_content.strip()) > 0:
            # ä¸“ä¸šåˆ†ææŠ¥å‘Šæ ¼å¼
            st.markdown("#### ğŸ” ä¸“ä¸šåˆ†ææŠ¥å‘Š")
            
            # å¤„ç†åˆ†æå†…å®¹ - ç§»é™¤å†—ä½™å¼€å¤´
            content = analysis_content
            
            # ç§»é™¤å¸¸è§çš„å†—ä½™å¼€å¤´
            redundant_beginnings = [
                "ä½œä¸ºä¸–ç•Œé¡¶çº§æœŸè´§åº“å­˜åˆ†æä¸“å®¶",
                "### ä½œä¸ºä¸–ç•Œé¡¶çº§æœŸè´§åº“å­˜åˆ†æä¸“å®¶",
                "åŸºäºæä¾›çš„çœŸå®æ•°æ®ï¼Œæˆ‘å¯¹",
                "æœ¬åˆ†æä¸¥æ ¼éµå¾ªæ•°æ®é©±åŠ¨åŸåˆ™",
                "åˆ†ææ¡†æ¶åˆ†ä¸ºäº”ä¸ªéƒ¨åˆ†",
                "æ€»å­—æ•°çº¦2500å­—",
                "ç¬¬ä¸€éƒ¨åˆ†ï¼šæ•°æ®è´¨é‡è¯„ä¼°ä¸å¯ä¿¡åº¦åˆ†æ",
                "æ•°æ®è´¨é‡æ˜¯åˆ†æå¯é æ€§çš„åŸºçŸ³"
            ]
            
            # ç§»é™¤å†—ä½™å¼€å¤´
            for redundant in redundant_beginnings:
                if content.startswith(redundant):
                    # æ‰¾åˆ°ç¬¬ä¸€ä¸ªå®é™…å†…å®¹æ®µè½
                    lines = content.split('\n')
                    content_start_idx = 0
                    for i, line in enumerate(lines):
                        if line.strip() and not any(r in line for r in redundant_beginnings):
                            # æŸ¥æ‰¾æ ¸å¿ƒè§‚ç‚¹æˆ–å®é™…åˆ†æå†…å®¹
                            if any(keyword in line for keyword in ['æ ¸å¿ƒè§‚ç‚¹', 'å¸‚åœºç°çŠ¶', 'åº“å­˜çŠ¶å†µ', 'å½“å‰', 'å»ºè®®']):
                                content_start_idx = i
                                break
                    if content_start_idx > 0:
                        content = '\n'.join(lines[content_start_idx:])
                    break
            
            # æ›¿æ¢å¸¸è§è‹±æ–‡è¯æ±‡ä¸ºä¸­æ–‡
            english_to_chinese = {
                " low ": " ä½ä½ ",
                " high ": " é«˜ä½ ",
                " support ": " æ”¯æ’‘ ",
                " resistance ": " é˜»åŠ› ",
                " bullish ": " çœ‹å¤š ",
                " bearish ": " çœ‹ç©º ",
                " trend ": " è¶‹åŠ¿ ",
                " breakout ": " çªç ´ ",
                " volume ": " æˆäº¤é‡ ",
                " price ": " ä»·æ ¼ ",
                " market ": " å¸‚åœº ",
                " analysis ": " åˆ†æ ",
                " data ": " æ•°æ® ",
                " level ": " æ°´å¹³ ",
                " range ": " åŒºé—´ ",
                " strong ": " å¼ºåŠ¿ ",
                " weak ": " å¼±åŠ¿ "
            }
            
            for eng, chn in english_to_chinese.items():
                content = content.replace(eng, chn)
            
            # ä¸“ä¸šæŠ¥å‘Šæ ·å¼
            st.markdown(f"""
            <div style="background: linear-gradient(135deg, #f8f9ff 0%, #ffffff 100%); 
                       padding: 25px; border-radius: 12px; margin: 20px 0;
                       border: 1px solid #e3e8f0; box-shadow: 0 4px 12px rgba(0,0,0,0.05);">
                <div style="font-family: 'Microsoft YaHei', 'SimHei', sans-serif; 
                           line-height: 1.8; font-size: 15px; color: #2c3e50; 
                           text-align: justify; letter-spacing: 0.5px;">
{content}
                </div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.warning("âš ï¸ æš‚æ— è¯¦ç»†åˆ†æå†…å®¹")
        
        # å›¾è¡¨å±•ç¤º
        st.markdown("### ğŸ“Š ä¸“ä¸šå›¾è¡¨åˆ†æ")
        
        chart_displayed = False
        
        # å°è¯•å¤šç§æ–¹å¼æ˜¾ç¤ºå›¾è¡¨
        if "charts_html" in result_data and result_data["charts_html"]:
            try:
                charts_content = result_data["charts_html"]
                
                # ä½¿ç”¨å¢å¼ºçš„å›¾è¡¨æ˜¾ç¤ºåŠŸèƒ½ï¼ˆå¸¦ä¸‹è½½æŒ‰é’®ï¼‰
                if CHART_DOWNLOAD_AVAILABLE:
                    chart_displayed = display_charts_with_download(
                        charts_content, 
                        commodity, 
                        module_name, 
                        "charts_html"
                    )
                    if chart_displayed:
                        # æ·»åŠ æ‰¹é‡ä¸‹è½½æŒ‰é’®ï¼ˆå¦‚æœæœ‰å¤šä¸ªå›¾è¡¨ï¼‰
                        create_batch_download_button(charts_content, commodity, module_name)
                else:
                    # å›é€€åˆ°åŸºç¡€æ˜¾ç¤ºåŠŸèƒ½
                    if hasattr(charts_content, 'to_html'):  # å•ä¸ªPlotlyå›¾è¡¨å¯¹è±¡
                        st.plotly_chart(charts_content, use_container_width=True, height=650)
                        chart_displayed = True
                        
                    # å¦‚æœæ˜¯å­—å…¸æ ¼å¼ï¼ˆå¤šä¸ªå›¾è¡¨ï¼‰
                    elif isinstance(charts_content, dict):
                        chart_count = 0
                        for chart_name, chart_obj in charts_content.items():
                            if hasattr(chart_obj, 'to_html'):  # ç¡®è®¤æ˜¯Plotlyå¯¹è±¡
                                st.subheader(f"ğŸ“Š {chart_name}")
                                st.plotly_chart(chart_obj, use_container_width=True, height=600)
                                chart_count += 1
                        
                        if chart_count > 0:
                            chart_displayed = True
                            st.success(f"âœ… å·²æ˜¾ç¤º {chart_count} ä¸ªä¸“ä¸šå›¾è¡¨")
                        else:
                            # å¦‚æœå­—å…¸ä¸­æ²¡æœ‰æœ‰æ•ˆçš„å›¾è¡¨å¯¹è±¡ï¼Œæ˜¾ç¤ºè°ƒè¯•ä¿¡æ¯
                            st.warning("âš ï¸ å›¾è¡¨æ•°æ®æ ¼å¼ä¸æ­£ç¡®")
                            with st.expander("è°ƒè¯•ä¿¡æ¯"):
                                st.json(charts_content)
                    
                    # å¦‚æœæ˜¯åˆ—è¡¨æ ¼å¼ï¼ˆåŒ…å«å¤šä¸ªå›¾è¡¨å¯¹è±¡ï¼‰
                    elif isinstance(charts_content, list):
                        chart_count = 0
                        for chart_item in charts_content:
                            if isinstance(chart_item, dict):
                                # å¤„ç†åŒ…å«titleå’Œhtmlå­—æ®µçš„å­—å…¸æ ¼å¼
                                title = chart_item.get('title', f'å›¾è¡¨ {chart_count + 1}')
                                chart_obj = chart_item.get('html', chart_item.get('chart', chart_item.get('figure')))
                                
                                if hasattr(chart_obj, 'to_html'):  # ç¡®è®¤æ˜¯Plotlyå¯¹è±¡
                                    st.subheader(f"ğŸ“Š {title}")
                                    st.plotly_chart(chart_obj, use_container_width=True, height=600)
                                    chart_count += 1
                            elif hasattr(chart_item, 'to_html'):  # ç›´æ¥æ˜¯Plotlyå¯¹è±¡
                                st.subheader(f"ğŸ“Š å›¾è¡¨ {chart_count + 1}")
                                st.plotly_chart(chart_item, use_container_width=True, height=600)
                                chart_count += 1
                        
                        if chart_count > 0:
                            chart_displayed = True
                            st.success(f"âœ… å·²æ˜¾ç¤º {chart_count} ä¸ªä¸“ä¸šå›¾è¡¨")
                        else:
                            st.warning("âš ï¸ åˆ—è¡¨ä¸­æ²¡æœ‰æœ‰æ•ˆçš„å›¾è¡¨å¯¹è±¡")
                            with st.expander("è°ƒè¯•ä¿¡æ¯"):
                                st.json(charts_content)
                    
                    # å¦‚æœæ˜¯å­—ç¬¦ä¸²æ ¼å¼çš„HTML
                    elif isinstance(charts_content, str) and len(charts_content.strip()) > 0:
                        # æ£€æŸ¥HTMLå†…å®¹æ˜¯å¦åŒ…å«æœ‰æ•ˆçš„å›¾è¡¨ä»£ç 
                        if any(tag in charts_content.lower() for tag in ['<svg', '<canvas', '<div', 'plotly', 'chart']):
                            st.components.v1.html(charts_content, height=650, scrolling=True)
                            chart_displayed = True
                        else:
                            # å¦‚æœHTMLå†…å®¹æ— æ•ˆï¼Œå°è¯•ä½œä¸ºæ–‡æœ¬æ˜¾ç¤º
                            st.code(charts_content, language='html')
                            chart_displayed = True
                            
                    else:
                        # å…¶ä»–æ ¼å¼ï¼Œæ˜¾ç¤ºè°ƒè¯•ä¿¡æ¯
                        st.warning(f"âš ï¸ æœªçŸ¥çš„å›¾è¡¨æ ¼å¼: {type(charts_content)}")
                        with st.expander("è°ƒè¯•ä¿¡æ¯"):
                            st.text(str(charts_content)[:1000] + "..." if len(str(charts_content)) > 1000 else str(charts_content))
                    
            except Exception as e:
                st.warning(f"ğŸ“Š å›¾è¡¨æ˜¾ç¤ºå‡ºé”™: {str(e)}")
                with st.expander("é”™è¯¯è¯¦æƒ…"):
                    st.text(f"é”™è¯¯ç±»å‹: {type(e).__name__}")
                    st.text(f"é”™è¯¯ä¿¡æ¯: {str(e)}")
                    if "charts_html" in result_data:
                        st.text(f"æ•°æ®ç±»å‹: {type(result_data['charts_html'])}")
        
        # æ£€æŸ¥å…¶ä»–å›¾è¡¨å­—æ®µ
        if not chart_displayed:
            chart_fields = ["professional_charts", "charts", "chart_data", "visualizations"]
            for field in chart_fields:
                if field in result_data and result_data[field]:
                    chart_data = result_data[field]
                    try:
                        # ä½¿ç”¨å¢å¼ºçš„å›¾è¡¨æ˜¾ç¤ºåŠŸèƒ½ï¼ˆå¸¦ä¸‹è½½æŒ‰é’®ï¼‰
                        if CHART_DOWNLOAD_AVAILABLE:
                            chart_displayed = display_charts_with_download(
                                chart_data, 
                                commodity, 
                                module_name, 
                                field
                            )
                            if chart_displayed:
                                # æ·»åŠ æ‰¹é‡ä¸‹è½½æŒ‰é’®ï¼ˆå¦‚æœæœ‰å¤šä¸ªå›¾è¡¨ï¼‰
                                create_batch_download_button(chart_data, commodity, module_name)
                                break
                        else:
                            # å›é€€åˆ°åŸºç¡€æ˜¾ç¤ºåŠŸèƒ½
                            if hasattr(chart_data, 'to_html'):  # å•ä¸ªPlotlyå›¾è¡¨å¯¹è±¡
                                # ä½¿ç”¨unique keyé¿å…IDå†²çª
                                unique_key = f"{attr_name}_{field}_single_chart"
                                st.plotly_chart(chart_data, use_container_width=True, height=650, key=unique_key)
                                chart_displayed = True
                                break
                            elif isinstance(chart_data, dict):
                                # å¦‚æœæ˜¯å­—å…¸ï¼Œæ£€æŸ¥æ˜¯å¦åŒ…å«å¤šä¸ªå›¾è¡¨
                                chart_count = 0
                                for key, value in chart_data.items():
                                    if hasattr(value, 'to_html'):  # Plotlyå›¾è¡¨å¯¹è±¡
                                        # æ¸…ç†keyï¼Œç§»é™¤emojiå’Œç‰¹æ®Šå­—ç¬¦
                                        clean_key = key.replace('ğŸ“Š ', '').replace('_', ' ').strip()
                                        st.subheader(f"{clean_key}")
                                        # ä½¿ç”¨unique keyé¿å…IDå†²çª
                                        unique_key = f"{attr_name}_{field}_{key}_{chart_count}"
                                        st.plotly_chart(value, use_container_width=True, height=600, key=unique_key)
                                        chart_count += 1
                                        chart_displayed = True
                                if chart_count > 0:
                                    st.success(f"âœ… å·²æ˜¾ç¤º {chart_count} ä¸ªå›¾è¡¨")
                                    break
                                else:
                                    # å¦‚æœæ²¡æœ‰å›¾è¡¨å¯¹è±¡ï¼Œæ˜¾ç¤ºè°ƒè¯•ä¿¡æ¯è€Œä¸æ˜¯å®Œæ•´å­—å…¸
                                    st.warning(f"âš ï¸ å­—æ®µ {field} åŒ…å«æ•°æ®ä½†ä¸æ˜¯æœ‰æ•ˆçš„å›¾è¡¨å¯¹è±¡")
                                    with st.expander("è°ƒè¯•ä¿¡æ¯"):
                                        st.json(chart_data)
                                    chart_displayed = True
                                    break
                            elif isinstance(chart_data, str):
                                # å¦‚æœæ˜¯å­—ç¬¦ä¸²ï¼Œå°è¯•ä½œä¸ºHTMLæ˜¾ç¤º
                                if any(tag in chart_data.lower() for tag in ['<svg', '<canvas', '<div', 'plotly']):
                                    st.components.v1.html(chart_data, height=650, scrolling=True)
                                    chart_displayed = True
                                    break
                                else:
                                    # ä½œä¸ºä»£ç æ˜¾ç¤º
                                    st.code(chart_data)
                                chart_displayed = True
                                break
                            else:
                                # å…¶ä»–ç±»å‹ï¼Œè½¬ä¸ºå­—ç¬¦ä¸²æ˜¾ç¤º
                                st.text(str(chart_data))
                                chart_displayed = True
                                break
                    except Exception as e:
                        st.warning(f"å¤„ç†å›¾è¡¨å­—æ®µ {field} æ—¶å‡ºé”™: {str(e)}")
                        continue
        
        # å¦‚æœä»ç„¶æ²¡æœ‰æ˜¾ç¤ºå›¾è¡¨ï¼Œä¸æ˜¾ç¤ºä»»ä½•è¯´æ˜ï¼ˆåˆ é™¤å†—ä½™å†…å®¹ï¼‰
        pass
        
        # å¤–éƒ¨æ•°æ®æ¥æºå’Œå¼•ç”¨
        citations_found = False
        
        # æ£€æŸ¥å„ç§å¼•ç”¨å­—æ®µ
        citation_fields = ["external_citations", "references", "citations", "data_sources"]
        for field in citation_fields:
            if field in result_data and result_data[field]:
                citations_found = True
                st.markdown("### ğŸ“š æ•°æ®æ¥æºä¸å¤–éƒ¨å¼•ç”¨")
                st.markdown("""
                <div style="background-color: #f0f8ff; padding: 15px; border-radius: 8px; 
                           border-left: 4px solid #4CAF50; margin: 15px 0;">
                    <p style="margin: 0; color: #2c5530; font-weight: 500;">
                    ğŸ“Œ ä»¥ä¸‹æ•°æ®æ¥æºå‡ä¸ºå®æ—¶è·å–ï¼Œç¡®ä¿ä¿¡æ¯çš„å‡†ç¡®æ€§å’Œæ—¶æ•ˆæ€§
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                citations = result_data[field]
                if isinstance(citations, list):
                    for i, citation in enumerate(citations, 1):
                        if isinstance(citation, dict):
                            title = citation.get('title', f'æ•°æ®æº {i}')
                            url = citation.get('url', '#')
                            source = citation.get('source', 'æœªçŸ¥æ¥æº')
                            
                            # æ—¶æ•ˆæ€§æ£€æŸ¥
                            time_info = ""
                            if 'date' in citation or 'timestamp' in citation:
                                date_str = citation.get('date', citation.get('timestamp', ''))
                                try:
                                    from datetime import datetime, timedelta
                                    if date_str:
                                        # å°è¯•è§£ææ—¥æœŸ
                                        for fmt in ['%Y-%m-%d', '%Y/%m/%d', '%Y-%m-%d %H:%M:%S']:
                                            try:
                                                date_obj = datetime.strptime(date_str.split(' ')[0], fmt.split(' ')[0])
                                                days_ago = (datetime.now() - date_obj).days
                                                if days_ago == 0:
                                                    time_info = "ï¼ˆä»Šæ—¥æ•°æ®ï¼‰"
                                                elif days_ago == 1:
                                                    time_info = "ï¼ˆæ˜¨æ—¥æ•°æ®ï¼‰"
                                                elif days_ago <= 7:
                                                    time_info = f"ï¼ˆ{days_ago}å¤©å‰æ•°æ®ï¼‰"
                                                else:
                                                    time_info = f"ï¼ˆ{date_str}ï¼‰"
                                                break
                                            except:
                                                continue
                                except:
                                    pass
                            
                            st.markdown(f"""
                            <div style="margin: 10px 0; padding: 12px; background-color: white; 
                                       border-radius: 8px; border-left: 3px solid #2196F3;">
                                <strong>[{i}]</strong> 
                                <a href="{url}" target="_blank" style="text-decoration: none; color: #1976D2;">
                                    {title}
                                </a>
                                <span style="color: #666; font-size: 0.9em;">{time_info}</span><br>
                                <span style="color: #888; font-size: 0.85em;">æ¥æº: {source}</span>
                            </div>
                            """, unsafe_allow_html=True)
                        elif isinstance(citation, str):
                            if citation.startswith('http'):
                                st.markdown(f"[{i}] [æ•°æ®é“¾æ¥]({citation})")
                            else:
                                st.markdown(f"[{i}] {citation}")
                        else:
                            st.markdown(f"[{i}] {str(citation)}")
                else:
                    st.write(citations)
                break
        
        # å¦‚æœæ²¡æœ‰æ‰¾åˆ°å¼•ç”¨ï¼Œæä¾›è¯´æ˜
        if not citations_found:
            st.markdown("### ğŸ“š æ•°æ®æ¥æº")
            st.markdown("""
            <div style="background-color: #fff3e0; padding: 15px; border-radius: 8px; 
                       border-left: 4px solid #ff9800; margin: 15px 0;">
                <p style="margin: 0; color: #ef6c00;">
                ğŸ“Š æœ¬åˆ†æåŸºäºæœ¬åœ°æ•°æ®åº“ä¸­çš„å†å²æ•°æ®å’Œå®æ—¶è®¡ç®—ç»“æœ<br>
                ğŸ”„ æ•°æ®å·²é€šè¿‡å¤šé‡éªŒè¯ï¼Œç¡®ä¿åˆ†æç»“è®ºçš„å¯é æ€§
                </p>
            </div>
            """, unsafe_allow_html=True)
        
        # å…³é”®å‘ç°
        if "key_findings" in result_data and result_data["key_findings"]:
            st.markdown("### ğŸ¯ å…³é”®å‘ç°")
            findings = result_data["key_findings"]
            if isinstance(findings, list):
                for finding in findings:
                    st.markdown(f"â€¢ {finding}")
            else:
                st.write(findings)
    
    def _display_debate_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºæ¿€çƒˆè¾©è®ºé¡µé¢"""
        st.markdown("## ğŸ­ æ¿€çƒˆè¾©è®ºç¯èŠ‚")
        
        # æ£€æŸ¥è¾©è®ºç»“æœæ•°æ®æ˜¯å¦å­˜åœ¨ - ä¿®å¤ï¼šæ£€æŸ¥å®é™…çš„æ•°æ®å†…å®¹è€Œä¸æ˜¯successå­—æ®µ
        if not result or not result.get("debate_section"):
            st.error("âŒ è¾©è®ºç¯èŠ‚æ•°æ®ä¸å®Œæ•´")
            return
        
        debate_section = result.get("debate_section", {})
        
        # è¾©è®ºç¯èŠ‚æ ‡é¢˜
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #ff9800 0%, #f57c00 100%); 
                   color: white; padding: 25px; border-radius: 15px; margin: 20px 0;">
            <h2 style="margin: 0; text-align: center;">ğŸ—£ï¸ å¤šç©ºæ¿€çƒˆäº¤é”‹</h2>
            <p style="margin: 10px 0 0 0; text-align: center; font-size: 1.1em;">
                æ™ºæ…§ç¢°æ’ï¼Œè§‚ç‚¹äº¤é”‹ï¼Œå¯»æ‰¾æŠ•èµ„çœŸç†
            </p>
            </div>
            """, unsafe_allow_html=True)
        
        # è¯¦ç»†è¾©è®ºè½®æ¬¡
        if "rounds" in debate_section and debate_section["rounds"]:
            st.markdown("### ğŸ—£ï¸ è¯¦ç»†è¾©è®ºè®°å½•")
            
            for i, round_data in enumerate(debate_section["rounds"], 1):
                # è½®æ¬¡æ ‡é¢˜
                st.markdown(f"""
                <div style="background: linear-gradient(90deg, #6c757d, #495057); color: white; 
                           padding: 10px; border-radius: 8px; text-align: center; margin: 20px 0 10px 0;">
                    <h4 style="margin: 0;">ç¬¬ {i} è½®è¾©è®º</h4>
                </div>
                """, unsafe_allow_html=True)
                
                # å¤šå¤´ç©ºå¤´å¹¶åˆ—æ˜¾ç¤º
                col_bull, col_bear = st.columns(2)
                
                with col_bull:
                    st.markdown("#### ğŸ‚ å¤šå¤´è§‚ç‚¹")
                    bull_argument = round_data.get("bull_argument", "æš‚æ— å†…å®¹")
                    st.markdown(f"""
                    <div style="background-color: #d4edda; padding: 15px; border-radius: 8px; 
                               border-left: 4px solid #28a745; margin: 10px 0;">
                        <div style="white-space: pre-wrap; line-height: 1.5;">
{bull_argument}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col_bear:
                    st.markdown("#### ğŸ» ç©ºå¤´è§‚ç‚¹")
                    bear_argument = round_data.get("bear_argument", "æš‚æ— å†…å®¹")
                    st.markdown(f"""
                    <div style="background-color: #f8d7da; padding: 15px; border-radius: 8px; 
                               border-left: 4px solid #dc3545; margin: 10px 0;">
                        <div style="white-space: pre-wrap; line-height: 1.5;">
{bear_argument}
                        </div>
                </div>
                """, unsafe_allow_html=True)
                
                if i < len(debate_section["rounds"]):
                    st.markdown("---")
        
        # è¾©è®ºæ•°æ®æ¥æºæ ‡æ³¨
        st.markdown("### ğŸ“š è¾©è®ºæ•°æ®æ¥æº")
        st.markdown("""
        <div style="background-color: #f0f8ff; padding: 15px; border-radius: 8px; 
                   border-left: 4px solid #4CAF50; margin: 15px 0;">
            <p style="margin: 0; color: #2c5530; font-weight: 500;">
            ğŸ“Š è¾©è®ºç¯èŠ‚å¼•ç”¨çš„å¸‚åœºæ•°æ®ã€æ–°é—»èµ„è®¯å’ŒæŠ€æœ¯æŒ‡æ ‡å‡æ¥è‡ªå®æ—¶æ•°æ®æº<br>
            ğŸ” æ‰€æœ‰è§‚ç‚¹åŸºäºå½“å‰å¸‚åœºçŠ¶å†µå’Œå†å²æ•°æ®åˆ†æå¾—å‡º<br>
            â° æ•°æ®æ—¶æ•ˆæ€§ï¼šå®æ—¶æ›´æ–°ï¼Œç¡®ä¿åˆ†æçš„å‡†ç¡®æ€§
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    def _display_trader_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºä¸“ä¸šäº¤æ˜“å‘˜é¡µé¢"""
        st.markdown("## ğŸ’¼ ä¸“ä¸šäº¤æ˜“å‘˜å†³ç­–")
        
        # æ£€æŸ¥äº¤æ˜“å‘˜ç»“æœæ•°æ®æ˜¯å¦å­˜åœ¨ - ä¿®å¤ï¼šæ£€æŸ¥å®é™…çš„æ•°æ®å†…å®¹è€Œä¸æ˜¯successå­—æ®µ
        if not result or not result.get("trading_section"):
            st.error("âŒ äº¤æ˜“å‘˜ç¯èŠ‚æ•°æ®ä¸å®Œæ•´")
            return
        
        trading_section = result.get("trading_section", {})
        
        # äº¤æ˜“å‘˜ä¿¡æ¯å¡ç‰‡
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #007bff 0%, #0056b3 100%); 
                   color: white; padding: 25px; border-radius: 15px; margin: 20px 0;">
            <h2 style="margin: 0; text-align: center;">ğŸ‘¨â€ğŸ’¼ èµ„æ·±äº¤æ˜“å‘˜åˆ†æ</h2>
            <p style="margin: 10px 0 0 0; text-align: center; font-size: 1.1em;">
                ç»¼åˆè¾©è®ºè§‚ç‚¹ï¼Œä¸“ä¸šè¯„åˆ¤ï¼Œç»™å‡ºäº¤æ˜“å»ºè®®
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # äº¤æ˜“é€»è¾‘ä¸æ‰§è¡Œè®¡åˆ’ï¼ˆåˆå¹¶æ˜¾ç¤ºï¼‰
        st.markdown("### ğŸ§  äº¤æ˜“é€»è¾‘ä¸æ‰§è¡Œè®¡åˆ’")
        
        reasoning = trading_section.get('reasoning', 'åŸºäºæŠ€æœ¯åˆ†æå’ŒåŸºæœ¬é¢åˆ†æï¼Œåˆ¶å®šä»¥ä¸‹äº¤æ˜“ç­–ç•¥')
        # ğŸ”¥ ä¿®å¤ï¼šä¸å†æ˜¾ç¤ºexecution_planå­—æ®µï¼Œå› ä¸ºå®ƒä¼šé‡å¤æ·»åŠ "æ‰§è¡Œè®¡åˆ’"
        # execution_plan = trading_section.get('execution_plan', '')
        
        # ğŸ”¥ ä¿®å¤ï¼šåªæ˜¾ç¤ºreasoningå­—æ®µï¼ˆå·²ç»è¢«åç«¯æ¸…ç†è¿‡ï¼Œä¸åŒ…å«æ‰§è¡Œè®¡åˆ’ï¼‰
        combined_content = reasoning
        # if execution_plan and execution_plan.strip():
        #     combined_content += f"\n\n**æ‰§è¡Œè®¡åˆ’ï¼š**\n{execution_plan}"
        
        st.markdown(f"""
        <div style="background-color: #e7f3ff; padding: 25px; border-radius: 12px; 
                   border-left: 5px solid #007bff; margin: 20px 0;">
            <div style="font-family: 'Microsoft YaHei'; line-height: 1.8; font-size: 15px; color: #2c3e50;">
                {combined_content.replace(chr(10), '<br>')}
            </div>
        </div>
        """, unsafe_allow_html=True)
        
    
    def _display_risk_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºä¸“ä¸šé£æ§é¡µé¢"""
        st.markdown("## ğŸ›¡ï¸ ä¸“ä¸šé£æ§è¯„ä¼°")
        
        # æ£€æŸ¥é£æ§ç»“æœæ•°æ®æ˜¯å¦å­˜åœ¨ - ä¿®å¤ï¼šæ£€æŸ¥å®é™…çš„æ•°æ®å†…å®¹è€Œä¸æ˜¯successå­—æ®µ
        if not result or not result.get("risk_section"):
            st.error("âŒ é£æ§ç¯èŠ‚æ•°æ®ä¸å®Œæ•´")
            return
        
        risk_section = result.get("risk_section", {})
        
        # é£æ§ç®¡ç†å¡ç‰‡
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #9c27b0 0%, #673ab7 100%); 
                   color: white; padding: 25px; border-radius: 15px; margin: 20px 0;">
            <h2 style="margin: 0; text-align: center;">ğŸ›¡ï¸ é£é™©ç®¡ç†å›¢é˜Ÿ</h2>
            <p style="margin: 10px 0 0 0; text-align: center; font-size: 1.1em;">
                ç‹¬ç«‹å®¢è§‚ï¼Œé£é™©ä¸ºå…ˆï¼Œç¨³å¥ç»è¥
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # é£é™©ç­‰çº§æŒ‡æ ‡
        col1, col2 = st.columns(2)
        with col1:
            overall_risk = risk_section.get('overall_risk', 'æœªçŸ¥')
            risk_color = "#dc3545" if "é«˜" in overall_risk else "#ffc107" if "ä¸­" in overall_risk else "#28a745"
            st.markdown(f"""
            <div style="background-color: {risk_color}; color: white; padding: 15px; 
                       border-radius: 8px; text-align: center;">
                <h4 style="margin: 0;">æ•´ä½“é£é™©ç­‰çº§</h4>
                <h3 style="margin: 5px 0 0 0;">{overall_risk}</h3>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            position_limit = risk_section.get('position_limit', 'æœªçŸ¥')
            st.markdown(f"""
            <div style="background-color: #17a2b8; color: white; padding: 15px; 
                       border-radius: 8px; text-align: center;">
                <h4 style="margin: 0;">å»ºè®®ä»“ä½ä¸Šé™</h4>
                <h3 style="margin: 5px 0 0 0;">{position_limit}</h3>
            </div>
            """, unsafe_allow_html=True)
        
        # æ­¢æŸå»ºè®®
        if risk_section.get('stop_loss'):
            st.markdown("### âš ï¸ æ­¢æŸå»ºè®®")
            st.markdown(f"""
            <div style="background-color: #f8d7da; padding: 15px; border-radius: 8px; 
                       border-left: 4px solid #dc3545;">
                <strong>å»ºè®®æ­¢æŸä½:</strong> {risk_section['stop_loss']}
            </div>
            """, unsafe_allow_html=True)
        
        # ğŸ”¥ åˆ é™¤é€šç”¨çš„é£é™©å› ç´ å’Œç¼“é‡Šæªæ–½æ¨¡æ¿ï¼Œè¿™äº›å†…å®¹ç°åœ¨æ•´åˆåˆ°é£æ§ç»ç†ä¸“ä¸šæ„è§ä¸­
        
        # é£æ§ç»ç†ä¸“ä¸šæ„è§
        if risk_section.get('manager_opinion'):
            st.markdown("### ğŸ‘¨â€ğŸ’¼ é£æ§ç»ç†ä¸“ä¸šæ„è§")
            st.markdown(f"""
            <div style="background-color: #e9ecef; padding: 20px; border-radius: 10px; 
                       border: 2px solid #6c757d; margin: 15px 0;">
                <div style="line-height: 1.6; font-style: italic;">
                    "{risk_section['manager_opinion']}"
                </div>
            </div>
            """, unsafe_allow_html=True)
    
    def _display_cio_decision_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºCIOå†³ç­–é¡µé¢"""
        st.markdown("## ğŸ‘” CIOæœ€ç»ˆæƒå¨å†³ç­–")
        
        # è·å–çœŸå®å¸‚åœºæ•°æ®
        market_data = get_market_data_info(commodity)
        
        # åˆ é™¤å¸‚åœºæ•°æ®åŸºç¡€éƒ¨åˆ†ï¼Œç›´æ¥è¿›å…¥CIOå†³ç­–
        
        # ğŸ”¥ ç›´æ¥ä»æ–°çš„æ•°æ®ç»“æ„è·å–CIOå†³ç­–æ•°æ®
        if not result.get("success"):
            error_msg = result.get("error", "æœªçŸ¥é”™è¯¯")
            st.error(f"âŒ åˆ†æå¤±è´¥: {error_msg}")
            st.info("ğŸ’¡ æç¤ºï¼šå¯èƒ½æ˜¯ç½‘ç»œè¿æ¥é—®é¢˜ï¼Œè¯·æ£€æŸ¥ç½‘ç»œåé‡è¯•")
            return
            
        executive_decision = result.get("decision_section", {})
        if not executive_decision:
            st.error("âŒ CIOå†³ç­–ç¯èŠ‚æ•°æ®ä¸å®Œæ•´ï¼Œè¯·é‡æ–°è¿è¡Œå®Œæ•´åˆ†ææµç¨‹")
            return
        
        # CIOå†³ç­–å¡ç‰‡
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%); 
                   color: white; padding: 30px; border-radius: 15px; margin: 20px 0;">
            <h2 style="margin: 0; text-align: center;">ğŸ‘” é¦–å¸­æŠ•èµ„å®˜</h2>
            <p style="margin: 10px 0 0 0; text-align: center; font-size: 1.2em;">
                æƒå¨æ‹æ¿ï¼Œä¸€é”¤å®šéŸ³ï¼Œå†³ç­–åˆ¶èƒœ
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # ç›´æ¥ä½¿ç”¨executive_decisionæ•°æ®
        final_decision = executive_decision.get('final_decision', 'æœªçŸ¥')
        # ç¡®ä¿final_decisionæ˜¯å­—ç¬¦ä¸²ç±»å‹
        if hasattr(final_decision, 'value'):
            final_decision_str = final_decision.value
        else:
            final_decision_str = str(final_decision)
        
        # è·å–å…¶ä»–å†³ç­–ä¿¡æ¯
        confidence_level = executive_decision.get('confidence_level', executive_decision.get('confidence', 0.5))
        # ç¡®ä¿confidence_levelæ˜¯æ•°å€¼ç±»å‹
        if isinstance(confidence_level, str):
            try:
                confidence_level = float(confidence_level.replace('%', '')) / 100 if '%' in confidence_level else float(confidence_level)
            except:
                confidence_level = 0.5
        
        directional_view = executive_decision.get('directional_view', 'ä¸­æ€§')
        directional_confidence = executive_decision.get('directional_confidence', 0.5)
        
        # ğŸ”¥ æ–°å¢ï¼šæ”¯æŒæ–‡æœ¬æ ¼å¼çš„ä¿¡å¿ƒåº¦
        directional_confidence_text = executive_decision.get('directional_confidence_text', "")
        operational_confidence_text = executive_decision.get('operational_confidence_text', "")
        
        # ç¡®ä¿directional_confidenceæ˜¯æ•°å€¼ç±»å‹
        if isinstance(directional_confidence, str):
            try:
                directional_confidence = float(directional_confidence.replace('%', '')) / 100 if '%' in directional_confidence else float(directional_confidence)
            except:
                directional_confidence = 0.5
        
        # ğŸ”§ è°ƒè¯•ä¿¡æ¯ - ä¸´æ—¶æ˜¾ç¤ºæ•°æ®è·å–æƒ…å†µ
        if st.sidebar.checkbox("ğŸ”§ æ˜¾ç¤ºè°ƒè¯•ä¿¡æ¯", key=f"debug_{commodity}"):
            st.sidebar.write("**è°ƒè¯•ä¿¡æ¯**:")
            st.sidebar.write(f"executive_decision keys: {list(executive_decision.keys())}")
            st.sidebar.write(f"confidence_level: {confidence_level} (type: {type(confidence_level)})")
            st.sidebar.write(f"directional_view: {directional_view}")
            st.sidebar.write(f"directional_confidence: {directional_confidence} (type: {type(directional_confidence)})")
            st.sidebar.write(f"final_decision: {final_decision}")
            if executive_decision:
                st.sidebar.json(executive_decision)
        
        # ğŸ”¥ åˆ é™¤æ­¢æŸæ ‡å‡†ç›¸å…³é€»è¾‘
        
        # âœ… ä¿®å¤ï¼šç›´æ¥ä½¿ç”¨æ–°çš„ç»Ÿä¸€å­—æ®µï¼Œä¸æ‰§è¡Œæ‘˜è¦ä¿æŒå®Œå…¨ä¸€è‡´
        operation_decision = executive_decision.get('operational_decision', 'ä¸­æ€§ç­–ç•¥')
        direction_judgment = executive_decision.get('directional_view', 'ä¸­æ€§')
        operational_confidence = executive_decision.get('operational_confidence', 'ä¸­')
        directional_confidence_display = executive_decision.get('directional_confidence', 'ä¸­')
        
        decision_color = "#28a745" if "åšå¤š" in operation_decision else "#dc3545" if "åšç©º" in operation_decision else "#ffc107"
        
        # ä½¿ç”¨StreamlitåŸç”Ÿç»„ä»¶æ›¿ä»£å¤æ‚HTML  
        with st.container():
            st.markdown("### ğŸ¯ æœ€ç»ˆå†³ç­–")
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric(
                    label="ğŸ“Š æ“ä½œå†³ç­–",
                    value=operation_decision,
                    help="åŸºäºç»¼åˆåˆ†æçš„æ“ä½œå»ºè®®"
                )
                st.metric(
                    label="ğŸ² æ“ä½œä¿¡å¿ƒåº¦", 
                    value=operational_confidence,
                    help="å¯¹æ“ä½œå†³ç­–çš„ä¿¡å¿ƒç¨‹åº¦ï¼ˆé«˜/ä¸­/ä½ï¼‰"
                )
                
            with col2:
                st.metric(
                    label="ğŸ§­ æ–¹å‘åˆ¤æ–­",
                    value=direction_judgment,
                    help="å¯¹å¸‚åœºæ–¹å‘çš„åˆ¤æ–­"
                )
                st.metric(
                    label="ğŸ“ æ–¹å‘ä¿¡å¿ƒåº¦",
                    value=directional_confidence_display,
                    help="å¯¹æ–¹å‘åˆ¤æ–­çš„ä¿¡å¿ƒç¨‹åº¦ï¼ˆé«˜/ä¸­/ä½ï¼‰"
                )
            
            # ğŸ”¥ åˆ é™¤æ­¢æŸæ ‡å‡†æ˜¾ç¤º
            st.divider()
        
        # åˆ é™¤é‡å¤çš„æ“ä½œå†³ç­–å’Œæ–¹å‘åˆ¤æ–­æ˜¾ç¤ºï¼Œå› ä¸ºåç»­çš„ç›‘æ§è¦ç‚¹å’ŒCIOå£°æ˜ä¸­å·²åŒ…å«è¿™äº›ä¿¡æ¯
        
        
        # å†³ç­–ç†ç”±
        if executive_decision.get('rationale'):
            st.markdown("### ğŸ§  å†³ç­–ç†ç”±")
            rationale_list = executive_decision['rationale']
            if isinstance(rationale_list, list):
                for i, reason in enumerate(rationale_list, 1):
                    st.markdown(f"""
                    <div style="background-color: #f8f9fa; padding: 15px; border-radius: 8px; 
                               margin: 10px 0; border-left: 4px solid #007bff;">
                        <strong>{i}.</strong> {reason}
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div style="background-color: #f8f9fa; padding: 15px; border-radius: 8px; 
                           border-left: 4px solid #007bff;">
                    {rationale_list}
                </div>
                """, unsafe_allow_html=True)
        
        # æ‰§è¡Œè®¡åˆ’å’Œç›‘æ§è¦ç‚¹å·²æ•´åˆåˆ°æœ€ç»ˆå†³ç­–ä¸­ï¼Œä¸å†å•ç‹¬æ˜¾ç¤º
        
        # CIOæƒå¨å£°æ˜ - ä¼˜åŒ–æ’ç‰ˆå’Œåº•è‰²
        if executive_decision.get('cio_statement'):
            st.markdown("### ğŸ“¢ CIOæƒå¨å£°æ˜")
            cio_statement = executive_decision['cio_statement']
            
            # ğŸ”¥ ä¼˜åŒ–æ’ç‰ˆï¼šå»é™¤markdownæ ¼å¼ï¼Œåˆ é™¤"ğŸ’¼ æŠ•èµ„å†³ç­–"æ ‡é¢˜
            if cio_statement and cio_statement.strip():
                # æ¸…ç†å¯èƒ½çš„markdownå­—ç¬¦å’Œä¸éœ€è¦çš„æ ‡é¢˜
                cleaned_statement = cio_statement.replace("**", "").replace("*", "").replace("#", "").replace("_", "")
                cleaned_statement = cleaned_statement.replace("ğŸ’¼ æŠ•èµ„å†³ç­–", "").replace("ğŸ’¼ **æŠ•èµ„å†³ç­–**", "")
                
                # âœ… ä½¿ç”¨å¸¦åº•è‰²çš„å®¹å™¨ä¼˜åŒ–æ˜¾ç¤ºï¼ˆåªæ˜¾ç¤ºä¸€æ¬¡ï¼‰
                st.markdown(f"""
                <div style="background-color: #f0f8ff; border: 2px solid #4a90e2; 
                           padding: 25px; margin: 20px 0; border-radius: 10px; 
                           line-height: 1.8; font-family: 'Microsoft YaHei', sans-serif;
                           box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
                <div style="color: #2c3e50; font-size: 16px; white-space: pre-line;">
                {cleaned_statement}
                </div>
                </div>
                """, unsafe_allow_html=True)
                
                # âœ… å·²åˆ é™¤é‡å¤çš„æŒ‰æ®µè½åˆ†å‰²æ˜¾ç¤ºé€»è¾‘
    
    def _display_analyst_overview(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºåˆ†æå¸ˆå›¢é˜Ÿæ€»è§ˆé¡µé¢"""
        st.markdown("## ğŸ“Š åˆ†æå¸ˆå›¢é˜Ÿæ€»è§ˆ")
        
        analysis_state = result.get("result")
        if not analysis_state:
            st.error("âŒ æ— æ³•è·å–åˆ†æå¸ˆå›¢é˜Ÿæ•°æ®")
            return
        
        # å›¢é˜Ÿæ¦‚è¿°å¡ç‰‡
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #4ECDC4 0%, #44A08D 100%); 
                   color: white; padding: 25px; border-radius: 15px; margin: 20px 0;">
            <h2 style="margin: 0; text-align: center;">ğŸ‘¥ ä¸“ä¸šåˆ†æå¸ˆå›¢é˜Ÿ</h2>
            <p style="margin: 10px 0 0 0; text-align: center; font-size: 1.1em;">
                6ä½èµ„æ·±ä¸“å®¶ï¼Œå¤šç»´åº¦æ·±åº¦åˆ†æï¼Œä¸“ä¸šå¯é 
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # åˆ†æå¸ˆçŠ¶æ€ç»Ÿè®¡
        module_attrs = [
            ("inventory_analysis", "ğŸ“¦ åº“å­˜ä»“å•åˆ†æå¸ˆ", "ä¸“æ³¨ä¾›éœ€å¹³è¡¡ä¸åº“å­˜å‘¨æœŸåˆ†æ"),
            ("positioning_analysis", "ğŸ¯ æŒä»“å¸­ä½åˆ†æå¸ˆ", "è§£è¯»èµ„é‡‘åŠ¨å‘ä¸ä¸»åŠ›æ„å›¾"),
            ("technical_analysis", "ğŸ“ˆ æŠ€æœ¯é¢åˆ†æå¸ˆ", "è¯†åˆ«ä»·æ ¼è¶‹åŠ¿ä¸äº¤æ˜“ä¿¡å·"),
            ("basis_analysis", "ğŸ’° åŸºå·®åˆ†æå¸ˆ", "å‰–æç°æœŸä»·å·®ä¸å¥—åˆ©æœºä¼š"),
            ("term_structure_analysis", "ğŸ“Š æœŸé™ç»“æ„åˆ†æå¸ˆ", "åˆ†æè¿œè¿‘æœˆä»·å·®ä¸æ—¶é—´ä»·å€¼"),
            ("news_analysis", "ğŸ“° æ–°é—»åˆ†æå¸ˆ", "è·Ÿè¸ªå¸‚åœºçƒ­ç‚¹ä¸æƒ…ç»ªå˜åŒ–")
        ]
        
        successful_modules = 0
        module_status = {}
        
        for attr_name, module_name, description in module_attrs:
            if hasattr(analysis_state, attr_name):
                module_result = getattr(analysis_state, attr_name)
                if module_result and (
                    module_result.status == "completed" or 
                    (hasattr(module_result.status, 'value') and module_result.status.value == "completed") or
                    str(module_result.status).lower() == "completed" or
                    str(module_result.status) == "AnalysisStatus.COMPLETED"
                ):
                    successful_modules += 1
                    module_status[attr_name] = "completed"
                else:
                    module_status[attr_name] = "failed"
            else:
                module_status[attr_name] = "missing"
        
        # æ€»ä½“å®Œæˆåº¦
        completion_rate = (successful_modules / 6) * 100
        color = "#28a745" if completion_rate >= 80 else "#ffc107" if completion_rate >= 60 else "#dc3545"
        
        st.markdown(f"""
        <div style="background-color: {color}; color: white; padding: 20px; 
                   border-radius: 10px; text-align: center; margin: 20px 0;">
            <h3 style="margin: 0;">å›¢é˜Ÿå®Œæˆåº¦</h3>
            <h1 style="margin: 10px 0; font-size: 2.5em;">{successful_modules}/6</h1>
            <p style="margin: 0; font-size: 1.1em;">å®Œæˆç‡: {completion_rate:.1f}%</p>
        </div>
        """, unsafe_allow_html=True)
        
        # å„åˆ†æå¸ˆè¯¦ç»†çŠ¶æ€
        st.markdown("### ğŸ‘¨â€ğŸ’¼ åˆ†æå¸ˆå›¢é˜ŸçŠ¶æ€")
        
        for attr_name, module_name, description in module_attrs:
            status = module_status.get(attr_name, "missing")
            
            if status == "completed":
                status_color = "#28a745"
                status_icon = "âœ…"
                status_text = "åˆ†æå®Œæˆ"
            elif status == "failed":
                status_color = "#dc3545"
                status_icon = "âŒ"
                status_text = "åˆ†æå¤±è´¥"
            else:
                status_color = "#6c757d"
                status_icon = "â³"
                status_text = "ç­‰å¾…ä¸­"
            
            st.markdown(f"""
            <div style="background-color: white; padding: 15px; border-radius: 10px; 
                       margin: 10px 0; border-left: 5px solid {status_color}; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                <div style="display: flex; align-items: center; margin-bottom: 8px;">
                    <span style="font-size: 1.2em; margin-right: 10px;">{status_icon}</span>
                    <strong style="font-size: 1.1em;">{module_name}</strong>
                    <span style="margin-left: auto; color: {status_color}; font-weight: bold;">{status_text}</span>
                </div>
                <p style="margin: 0; color: #666; font-size: 0.9em;">{description}</p>
            </div>
            """, unsafe_allow_html=True)
        
        # å¦‚æœæœ‰å®Œæˆçš„åˆ†æï¼Œæ˜¾ç¤ºå…³é”®å‘ç°æ±‡æ€»
        if successful_modules > 0:
            st.markdown("### ğŸ¯ å…³é”®å‘ç°æ±‡æ€»")
            
            key_findings = []
            for attr_name, module_name, _ in module_attrs:
                if module_status.get(attr_name) == "completed" and hasattr(analysis_state, attr_name):
                    module_result = getattr(analysis_state, attr_name)
                    if hasattr(module_result, 'result_data') and module_result.result_data:
                        result_data = module_result.result_data
                        
                        # å°è¯•è·å–å…³é”®å‘ç°æˆ–åˆ†ææ‘˜è¦
                        finding = None
                        if "key_findings" in result_data and result_data["key_findings"]:
                            findings = result_data["key_findings"]
                            if isinstance(findings, list) and findings:
                                finding = findings[0]  # å–ç¬¬ä¸€ä¸ªå…³é”®å‘ç°
                        
                        if not finding:
                            # å°è¯•ä»åˆ†æå†…å®¹ä¸­æå–æ‘˜è¦
                            content_fields = ["ai_comprehensive_analysis", "ai_analysis", "analysis_content", "speculative_analysis", "content"]
                            for field in content_fields:
                                if field in result_data and result_data[field]:
                                    content = result_data[field]
                                    if isinstance(content, str) and len(content) > 50:
                                        finding = content[:150] + "..."
                                        break
                        
                        if finding:
                            key_findings.append((module_name, finding))
            
            if key_findings:
                for module_name, finding in key_findings:
                    st.markdown(f"""
                    <div style="background-color: #f8f9fa; padding: 12px; border-radius: 8px; 
                               margin: 8px 0; border-left: 3px solid #007bff;">
                        <strong>{module_name}:</strong> {finding}
                    </div>
                    """, unsafe_allow_html=True)
        
        # å›¢é˜Ÿåä½œè¯´æ˜
        if successful_modules >= 4:
            st.markdown("### ğŸ¤ å›¢é˜Ÿåä½œä¼˜åŠ¿")
            st.markdown("""
            - **å¤šç»´åº¦åˆ†æ**: ä»åŸºæœ¬é¢ã€æŠ€æœ¯é¢ã€èµ„é‡‘é¢ç­‰å¤šè§’åº¦å…¨é¢åˆ†æ
            - **ä¸“ä¸šåˆ†å·¥**: æ¯ä½åˆ†æå¸ˆä¸“æ³¨è‡ªå·±çš„é¢†åŸŸï¼Œç¡®ä¿åˆ†æè´¨é‡
            - **ç›¸äº’éªŒè¯**: ä¸åŒåˆ†æå¸ˆçš„ç»“è®ºå¯ä»¥ç›¸äº’å°è¯ï¼Œæé«˜å¯é æ€§
            - **é£é™©æ§åˆ¶**: å¤šä¸“ä¸šè§†è§’æœ‰åŠ©äºè¯†åˆ«æ½œåœ¨é£é™©ç‚¹
            - **å†³ç­–æ”¯æŒ**: ä¸ºåç»­çš„äº¤æ˜“å†³ç­–æä¾›å…¨é¢çš„æ•°æ®æ”¯æ’‘
            """)
    
    # å„ä¸ªå…·ä½“åˆ†æå¸ˆé¡µé¢çš„æ˜¾ç¤ºæ–¹æ³•
    def _display_inventory_analysis_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºåº“å­˜ä»“å•åˆ†æé¡µé¢"""
        analysis_state = result.get("result")
        if hasattr(analysis_state, 'inventory_analysis'):
            self._display_single_analyst_report(commodity, analysis_state, 'inventory_analysis', 'ğŸ“¦ åº“å­˜ä»“å•åˆ†æå¸ˆ')
        else:
            st.error("âŒ åº“å­˜ä»“å•åˆ†ææ•°æ®ä¸å¯ç”¨")
    
    def _display_technical_analysis_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºæŠ€æœ¯åˆ†æé¡µé¢"""
        analysis_state = result.get("result")
        if hasattr(analysis_state, 'technical_analysis'):
            self._display_single_analyst_report(commodity, analysis_state, 'technical_analysis', 'ğŸ“ˆ æŠ€æœ¯é¢åˆ†æå¸ˆ')
        else:
            st.error("âŒ æŠ€æœ¯åˆ†ææ•°æ®ä¸å¯ç”¨")
    
    def _display_positioning_analysis_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºæŒä»“å¸­ä½åˆ†æé¡µé¢"""
        analysis_state = result.get("result")
        if hasattr(analysis_state, 'positioning_analysis'):
            self._display_single_analyst_report(commodity, analysis_state, 'positioning_analysis', 'ğŸ¯ æŒä»“å¸­ä½åˆ†æå¸ˆ')
        else:
            st.error("âŒ æŒä»“å¸­ä½åˆ†ææ•°æ®ä¸å¯ç”¨")
    
    def _display_basis_analysis_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºåŸºå·®åˆ†æé¡µé¢"""
        analysis_state = result.get("result")
        if hasattr(analysis_state, 'basis_analysis'):
            self._display_single_analyst_report(commodity, analysis_state, 'basis_analysis', 'ğŸ’° åŸºå·®åˆ†æå¸ˆ')
        else:
            st.error("âŒ åŸºå·®åˆ†ææ•°æ®ä¸å¯ç”¨")
    
    def _display_term_structure_analysis_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºæœŸé™ç»“æ„åˆ†æé¡µé¢"""
        analysis_state = result.get("result")
        if hasattr(analysis_state, 'term_structure_analysis'):
            self._display_single_analyst_report(commodity, analysis_state, 'term_structure_analysis', 'ğŸ“Š æœŸé™ç»“æ„åˆ†æå¸ˆ')
        else:
            st.error("âŒ æœŸé™ç»“æ„åˆ†ææ•°æ®ä¸å¯ç”¨")
    
    def _display_news_analysis_page(self, commodity: str, result: Dict):
        """æ˜¾ç¤ºæ–°é—»åˆ†æé¡µé¢"""
        analysis_state = result.get("result")
        if hasattr(analysis_state, 'news_analysis'):
            self._display_single_analyst_report(commodity, analysis_state, 'news_analysis', 'ğŸ“° æ–°é—»åˆ†æå¸ˆ')
            
            # æ–°é—»åˆ†æç‰¹æ®Šå¤„ç†ï¼šç¡®ä¿æ–°é—»é“¾æ¥å®Œæ•´æ˜¾ç¤º
            news_result = getattr(analysis_state, 'news_analysis')
            if hasattr(news_result, 'result_data') and news_result.result_data:
                result_data = news_result.result_data
                
                # æ£€æŸ¥æ˜¯å¦æœ‰æ–°é—»æ•°æ®ä½†ç¼ºå°‘é“¾æ¥æ˜¾ç¤º
                news_count = result_data.get('news_count', 0)
                if news_count > 0:
                    st.markdown("### ğŸ“° æ–°é—»æ•°æ®ç»Ÿè®¡")
                    
                    # æ˜¾ç¤ºæ–°é—»æ•°é‡ç»Ÿè®¡
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        akshare_count = result_data.get('akshare_count', 0)
                        st.metric("AkShareæ–°é—»", f"{akshare_count}æ¡")
                    with col2:
                        search_count = result_data.get('search_count', 0)  
                        st.metric("æœç´¢æ–°é—»", f"{search_count}æ¡")
                    with col3:
                        st.metric("æ€»è®¡æ–°é—»", f"{news_count}æ¡")
                    
                    # æ–°é—»æ—¶æ•ˆæ€§è¯´æ˜
                    st.markdown("""
                    <div style="background-color: #e8f5e8; padding: 15px; border-radius: 8px; 
                               border-left: 4px solid #4caf50; margin: 15px 0;">
                        <h5 style="color: #2e7d2e; margin-bottom: 10px;">ğŸ“… æ–°é—»æ—¶æ•ˆæ€§ä¿è¯</h5>
                        <p style="margin: 0; color: #2e7d2e;">
                        â€¢ <strong>å®æ—¶æ–°é—»æº</strong>ï¼šæ–°åè´¢ç»ã€ç”Ÿæ„ç¤¾ã€ä¸œæ–¹è´¢å¯Œç­‰æƒå¨è´¢ç»åª’ä½“<br>
                        â€¢ <strong>æ—¶é—´ç­›é€‰</strong>ï¼šä¼˜å…ˆä½¿ç”¨è¿‘7å¤©å†…çš„æœ€æ–°æ–°é—»èµ„è®¯<br>
                        â€¢ <strong>å†…å®¹è¿‡æ»¤</strong>ï¼šåªä¿ç•™ä¸{commodity}ç›¸å…³çš„é«˜è´¨é‡æ–°é—»<br>
                        â€¢ <strong>å¤šæºéªŒè¯</strong>ï¼šäº¤å‰éªŒè¯å¤šä¸ªæ–°é—»æºçš„ä¿¡æ¯å‡†ç¡®æ€§
                        </p>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.error("âŒ æ–°é—»åˆ†ææ•°æ®ä¸å¯ç”¨")

# ============================================================================
# WordæŠ¥å‘Šç”Ÿæˆå™¨
# ============================================================================

class WordReportGenerator:
    """ä¸“ä¸šç ”ç©¶æŠ¥å‘Šç”Ÿæˆå™¨"""
    
    def __init__(self):
        self.doc = None
        self.charts_data = {}  # å­˜å‚¨å›¾è¡¨æ•°æ®
    
    def create_comprehensive_report(self, analysis_results: Dict, analysis_config: Dict, 
                                   include_charts: bool = True, progress_callback=None) -> io.BytesIO:
        """åˆ›å»ºä¸“ä¸šç ”ç©¶æŠ¥å‘Š
        
        Args:
            analysis_results: åˆ†æç»“æœå­—å…¸
            analysis_config: åˆ†æé…ç½®
            include_charts: æ˜¯å¦åŒ…å«å›¾è¡¨ï¼ˆå›¾è¡¨è½¬æ¢å¾ˆè€—æ—¶ï¼Œå¿«é€Ÿæ¨¡å¼å¯è®¾ä¸ºFalseï¼‰
            progress_callback: è¿›åº¦å›è°ƒå‡½æ•°ï¼ˆç”¨äºæ˜¾ç¤ºè¿›åº¦ï¼‰
        """
        
        if not DOCX_AVAILABLE:
            raise ImportError("æœªå®‰è£…python-docxåº“")
        
        def update_progress(msg: str):
            """æ›´æ–°è¿›åº¦"""
            if progress_callback:
                progress_callback(msg)
            print(f"ğŸ“ {msg}")
        
        # åˆ›å»ºæ–‡æ¡£
        update_progress("åˆå§‹åŒ–Wordæ–‡æ¡£...")
        self.doc = Document()
        self.include_charts = include_charts  # ä¿å­˜é…ç½®
        
        # è®¾ç½®ä¸“ä¸šæ–‡æ¡£æ ·å¼
        self._setup_professional_styles()
        
        # è·å–åˆ†æçš„å“ç§ï¼ˆå–ç¬¬ä¸€ä¸ªå“ç§ä½œä¸ºä¸»è¦åˆ†æå¯¹è±¡ï¼‰
        commodity = list(analysis_results.keys())[0] if analysis_results else "æœªçŸ¥å“ç§"
        analysis_date = analysis_config.get("analysis_date", datetime.now().strftime('%Y-%m-%d'))
        
        # 1. æ·»åŠ å°é¢é¡µ
        update_progress("ç”Ÿæˆå°é¢é¡µ...")
        self._add_cover_page(commodity, analysis_date, analysis_config)
        
        # 2. æ·»åŠ æ‰§è¡Œæ‘˜è¦ï¼ˆæ–°å¢ï¼‰
        update_progress("ç”Ÿæˆæ‰§è¡Œæ‘˜è¦...")
        self._add_executive_summary_new(analysis_results, commodity)
        
        # 3. æ·»åŠ ä¸“ä¸šç›®å½•
        update_progress("ç”Ÿæˆç›®å½•...")
        self._add_professional_toc(analysis_results, commodity)
        
        # 4. æ·»åŠ å„åˆ†æå¸ˆæ¨¡å—å†…å®¹
        update_progress("ç”Ÿæˆåˆ†ææ¨¡å—å†…å®¹...")
        self._add_analyst_modules(analysis_results, commodity, progress_callback=update_progress)
        
        # 5. æ·»åŠ CIOå†³ç­–ç»“è®º
        update_progress("ç”ŸæˆCIOå†³ç­–...")
        self._add_cio_conclusion(analysis_results, commodity)
        
        # ä¿å­˜åˆ°å†…å­˜
        update_progress("ä¿å­˜Wordæ–‡æ¡£...")
        doc_io = io.BytesIO()
        self.doc.save(doc_io)
        doc_io.seek(0)
        
        update_progress("âœ… WordæŠ¥å‘Šç”Ÿæˆå®Œæˆï¼")
        
        return doc_io
    
    def _setup_professional_styles(self):
        """è®¾ç½®ä¸“ä¸šæ–‡æ¡£æ ·å¼"""
        
        # è®¾ç½®æ–‡æ¡£é»˜è®¤å­—ä½“
        from docx.oxml.shared import qn
        
        # æ ‡é¢˜æ ·å¼
        title_style = self.doc.styles['Title']
        title_font = title_style.font
        title_font.name = 'é»‘ä½“'
        title_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        title_font.size = Pt(22)
        title_font.bold = True
        
        # ä¸€çº§æ ‡é¢˜æ ·å¼
        heading1_style = self.doc.styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'é»‘ä½“'
        heading1_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        heading1_font.size = Pt(16)
        heading1_font.bold = True
        
        # äºŒçº§æ ‡é¢˜æ ·å¼
        heading2_style = self.doc.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = 'é»‘ä½“'
        heading2_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        heading2_font.size = Pt(14)
        heading2_font.bold = True
        
        # ä¸‰çº§æ ‡é¢˜æ ·å¼
        heading3_style = self.doc.styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.name = 'é»‘ä½“'
        heading3_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        heading3_font.size = Pt(12)
        heading3_font.bold = True
        
        # æ­£æ–‡æ ·å¼
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'å®‹ä½“'
        normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
        normal_font.size = Pt(12)
        
        # è®¾ç½®æ®µè½é—´è·
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.space_after = Pt(6)
        normal_paragraph.line_spacing = 1.2
    
    def _add_cover_page(self, commodity: str, analysis_date: str, config: Dict):
        """æ·»åŠ ä¸“ä¸šå°é¢é¡µ"""
        
        # æ·»åŠ ç©ºè¡Œç”¨äºä¸Šè¾¹è·
        for _ in range(8):
            self.doc.add_paragraph()
        
        # ä¸»æ ‡é¢˜
        title = self.doc.add_heading(f'{commodity}å“ç§åˆ†ææŠ¥å‘Š', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # æ·»åŠ ç©ºè¡Œ
        for _ in range(3):
            self.doc.add_paragraph()
        
        # æŠ¥å‘Šæ—¥æœŸ
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'æŠ¥å‘Šæ—¥æœŸï¼š{analysis_date}')
        date_run.font.size = Pt(16)
        date_run.font.name = 'é»‘ä½“'
        
        # æ·»åŠ ç©ºè¡Œ
        for _ in range(8):
            self.doc.add_paragraph()
        
        # æŠ¥å‘Šä¿¡æ¯
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_lines = [
            "å•†å“æœŸè´§ Trading Agentsç³»ç»Ÿ",
            "ä¸“ä¸šæŠ•èµ„ç ”ç©¶æŠ¥å‘Š",
            "",
            f"åˆ†æå“ç§ï¼š{commodity}",
            f"åˆ†ææ¨¡å¼ï¼š{config.get('analysis_mode', 'å®Œæ•´åˆ†æ')}",
            f"æŠ¥å‘Šç”Ÿæˆæ—¶é—´ï¼š{datetime.now().strftime('%Yå¹´%mæœˆ%dæ—¥ %H:%M')}"
        ]
        
        for line in info_lines:
            info_para.add_run(line + '\n')
        
        # æ·»åŠ åˆ†é¡µç¬¦
        self.doc.add_page_break()
    
    def _add_executive_summary_new(self, analysis_results: Dict, commodity: str):
        """æ·»åŠ æ‰§è¡Œæ‘˜è¦ï¼ˆæ–°å¢ - ç®€åŒ–ç‰ˆï¼Œç”¨äºå¿«é€ŸæŸ¥çœ‹ï¼‰"""
        
        # ğŸ”¥ å…¼å®¹æ€§æ£€æŸ¥ï¼šå¦‚æœcommodityæ˜¯å­—å…¸ï¼ˆæ—§ç‰ˆæœ¬çš„configå‚æ•°ï¼‰ï¼Œåˆ™è·³è¿‡
        if isinstance(commodity, dict):
            print("âš ï¸ æ£€æµ‹åˆ°æ—§ç‰ˆæœ¬è°ƒç”¨ï¼Œè¯·é‡å¯Streamlitä»¥ä½¿ç”¨æ–°ç‰ˆæœ¬")
            return
        
        self.doc.add_heading('æ‰§è¡Œæ‘˜è¦', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        try:
            # è·å–ç¬¬ä¸€ä¸ªå“ç§çš„åˆ†æç»“æœ
            result = list(analysis_results.values())[0]
            if result.get("status") != "success":
                self.doc.add_paragraph("åˆ†ææœªå®Œæˆï¼Œæ— æ‰§è¡Œæ‘˜è¦")
                self.doc.add_page_break()
                return
            
            # ğŸ”¥ ä¿®å¤ï¼šä½¿ç”¨æ­£ç¡®çš„æ•°æ®è·¯å¾„ï¼ˆä¸ç•Œé¢ä¸€è‡´ï¼‰
            debate_result = result.get("result", {}).get("debate_result", {})
            
            # æ£€æŸ¥åˆ†ææ˜¯å¦æˆåŠŸ
            if not debate_result or not debate_result.get("success"):
                self.doc.add_paragraph("åˆ†æå¤±è´¥ï¼Œæ— æ‰§è¡Œæ‘˜è¦")
                self.doc.add_page_break()
                return
            
            # ä»æ­£ç¡®çš„ä½ç½®è·å–CIOå†³ç­–æ•°æ®
            executive_decision = debate_result.get("decision_section", {})
            
            if not executive_decision:
                self.doc.add_paragraph("CIOå†³ç­–æ•°æ®ç¼ºå¤±ï¼Œæ— æ‰§è¡Œæ‘˜è¦")
                self.doc.add_page_break()
                return
            
            # æå–å…³é”®ä¿¡æ¯
            operational_decision = executive_decision.get("operational_decision", "æœªæ˜ç¡®")
            operational_confidence = executive_decision.get("operational_confidence", "æœªè¯„ä¼°")
            directional_view = executive_decision.get("directional_view", "æœªæ˜ç¡®")
            directional_confidence = executive_decision.get("directional_confidence", "æœªè¯„ä¼°")
            
            # 1. å†³ç­–ç»“æœ
            self.doc.add_heading('ä¸€ã€å†³ç­–ç»“æœ', level=2)
            
            decision_table = self.doc.add_table(rows=4, cols=2)
            decision_table.style = 'Light Grid Accent 1'
            
            decisions = [
                ("æ“ä½œå†³ç­–", operational_decision),
                ("æ“ä½œä¿¡å¿ƒåº¦", operational_confidence),
                ("æ–¹å‘åˆ¤æ–­", directional_view),
                ("æ–¹å‘ä¿¡å¿ƒåº¦", directional_confidence)
            ]
            
            for i, (label, value) in enumerate(decisions):
                row = decision_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(value)
            
            self.doc.add_paragraph()
            
            # 2. å…³é”®ä»·ä½ï¼ˆå¦‚æœæœ‰ï¼‰- ğŸ”¥ ä¿®å¤ï¼šä½¿ç”¨ä¸ç•Œé¢ç›¸åŒçš„æå–é€»è¾‘
            import re
            
            # ä¼˜å…ˆä»CIOçš„statementä¸­æå–ï¼ˆæ ¼å¼åŒ–å¥½çš„ï¼‰
            cio_statement = executive_decision.get("cio_statement", "")
            trading_section = debate_result.get("trading_section", {})
            trader_reasoning = trading_section.get("reasoning", "")
            
            # åˆå¹¶ä¸¤ä¸ªæ–‡æœ¬æºï¼ŒCIOçš„statementä¼˜å…ˆ
            combined_text = cio_statement + "\n\n" + trader_reasoning
            
            entry_range = None
            stop_loss = None
            target_price = None
            
            # æå–è¿›åœºåŒºé—´
            entry_match = re.search(r'å»ºè®®è¿›åœºåŒºé—´[ï¼š:]?\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', combined_text)
            if not entry_match:
                entry_match = re.search(r'è¿›åœº[åŒºé—´]?[ï¼š:]?\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', combined_text)
            if entry_match:
                entry_range = f"{entry_match.group(1)}-{entry_match.group(2)}å…ƒ"
            
            # æå–æ­¢æŸä½
            stop_match = re.search(r'æ­¢æŸä½?[ï¼š:]?\s*([<>~â‰¤â‰¥]\s*\d+\.?\d*)å…ƒ', combined_text)
            if not stop_match:
                stop_match = re.search(r'æ­¢æŸ[ï¼š:]?\s*([<>~]\d+\.?\d*)å…ƒ', combined_text)
            if stop_match:
                stop_loss = stop_match.group(1).replace(' ', '') + "å…ƒ"
            
            # æå–ç›®æ ‡ä»·ä½
            target_match = re.search(r'ç›®æ ‡ä»·ä½[ï¼š:]?\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', combined_text)
            if not target_match:
                target_match = re.search(r'ç›®æ ‡[ï¼š:]?\s*(\d+\.?\d*)[~ï½-](\d+\.?\d*)å…ƒ', combined_text)
            if target_match:
                target_price = f"{target_match.group(1)}-{target_match.group(2)}å…ƒ"
            
            if entry_range or stop_loss or target_price:
                self.doc.add_heading('äºŒã€å…³é”®ä»·ä½', level=2)
                
                price_table = self.doc.add_table(rows=3, cols=2)
                price_table.style = 'Light Grid Accent 1'
                
                prices = [
                    ("è¿›åœºåŒºé—´", entry_range or "å¾…å®š"),
                    ("ç›®æ ‡ä»·ä½", target_price or "å¾…å®š"),
                    ("æ­¢æŸä»·ä½", stop_loss or "å¾…å®š")
                ]
                
                for i, (label, value) in enumerate(prices):
                    row = price_table.rows[i]
                    row.cells[0].text = label
                    row.cells[1].text = str(value)
                
                self.doc.add_paragraph()
            
            # 3. åˆ†ææ¦‚å†µ
            self.doc.add_heading('ä¸‰ã€åˆ†ææ¦‚å†µ', level=2)
            
            system_metadata = debate_result.get('system_metadata', {})
            
            summary_info = [
                f"â€¢ å®Œæˆé˜¶æ®µï¼š5/5",
                f"â€¢ åˆ†ææ¨¡å—ï¼š{system_metadata.get('modules_count', 6)}ä¸ª",
                f"â€¢ è¾©è®ºè½®æ•°ï¼š{system_metadata.get('debate_rounds', 3)}è½®",
                f"â€¢ åˆ†ææ¨¡å‹ï¼š{system_metadata.get('model_used', 'DeepSeek-Reasoner')}"
            ]
            
            for info in summary_info:
                para = self.doc.add_paragraph(info)
                para.paragraph_format.left_indent = Inches(0.25)
            
            self.doc.add_paragraph()
            
            # 4. é£é™©æç¤º
            self.doc.add_heading('å››ã€é£é™©æç¤º', level=2)
            
            # ğŸ”¥ ä¿®å¤ï¼šä»æ­£ç¡®ä½ç½®è·å–é£é™©åˆ†ææ•°æ®
            risk_section = debate_result.get("risk_section", {})
            risk_level = risk_section.get("risk_level", "æœªè¯„ä¼°")
            position_advice = risk_section.get("position_advice", "è¯·è°¨æ…å†³ç­–")
            
            risk_para = self.doc.add_paragraph()
            risk_para.add_run(f"é£é™©è¯„çº§ï¼š{risk_level}\n").bold = True
            risk_para.add_run(f"ä»“ä½å»ºè®®ï¼š{position_advice}\n")
            risk_para.add_run("\næœ¬æŠ¥å‘Šä»…ä¾›å‚è€ƒï¼Œä¸æ„æˆæŠ•èµ„å»ºè®®ã€‚æŠ•èµ„è€…åº”æ ¹æ®è‡ªèº«æƒ…å†µè°¨æ…å†³ç­–ï¼Œå¹¶è‡ªè¡Œæ‰¿æ‹…æŠ•èµ„é£é™©ã€‚")
            
        except Exception as e:
            self.doc.add_paragraph(f"æ‰§è¡Œæ‘˜è¦ç”Ÿæˆå‡ºé”™ï¼š{str(e)}")
        
        self.doc.add_page_break()
    
    def _add_professional_toc(self, analysis_results: Dict, commodity: str):
        """æ·»åŠ ä¸“ä¸šç›®å½•"""
        
        self.doc.add_heading('ç›®  å½•', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ç›®å½•é¡¹ç›®
        toc_items = [
            ("ä¸€ã€åº“å­˜ä»“å•åˆ†æ", "3"),
            ("äºŒã€æŒä»“å¸­ä½åˆ†æ", ""),
            ("ä¸‰ã€æœŸé™ç»“æ„åˆ†æ", ""),
            ("å››ã€æŠ€æœ¯é¢åˆ†æ", ""),
            ("äº”ã€åŸºå·®åˆ†æ", ""),
            ("å…­ã€æ–°é—»åˆ†æ", ""),
            ("ä¸ƒã€æŠ•èµ„å†³ç­–ç»“è®º", ""),
        ]
        
        # æ·»åŠ ç›®å½•è¡¨æ ¼
        table = self.doc.add_table(rows=len(toc_items), cols=2)
        table.style = 'Table Grid'
        
        for i, (title, page) in enumerate(toc_items):
            row = table.rows[i]
            row.cells[0].text = title
            row.cells[1].text = page
            
            # è®¾ç½®å­—ä½“
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = 'å®‹ä½“'
                cell.paragraphs[0].runs[0].font.size = Pt(12)
        
        self.doc.add_page_break()
    
    def _add_analyst_modules(self, analysis_results: Dict, commodity: str, progress_callback=None):
        """æ·»åŠ å„åˆ†æå¸ˆæ¨¡å—å†…å®¹"""
        
        # è·å–ç¬¬ä¸€ä¸ªå“ç§çš„åˆ†æç»“æœ
        result = list(analysis_results.values())[0]
        if result.get("status") != "success":
            return
        
        # è·å–åˆ†æçŠ¶æ€
        if result.get("type") == "complete_flow":
            analysis_state = result["result"]["analysis_state"]
        else:
            analysis_state = result["result"]["analysis_state"]
        
        # æ¨¡å—æ˜ å°„
        modules_info = [
            ("inventory_analysis", "ä¸€ã€åº“å­˜ä»“å•åˆ†æ", "inventory_analysis"),
            ("positioning_analysis", "äºŒã€æŒä»“å¸­ä½åˆ†æ", "positioning_analysis"),
            ("term_structure_analysis", "ä¸‰ã€æœŸé™ç»“æ„åˆ†æ", "term_structure_analysis"),
            ("technical_analysis", "å››ã€æŠ€æœ¯é¢åˆ†æ", "technical_analysis"),
            ("basis_analysis", "äº”ã€åŸºå·®åˆ†æ", "basis_analysis"),
            ("news_analysis", "å…­ã€æ–°é—»åˆ†æ", "news_analysis")
        ]
        
        for idx, (module_attr, title, module_key) in enumerate(modules_info, 1):
            if progress_callback:
                progress_callback(f"å¤„ç†æ¨¡å— {idx}/{len(modules_info)}: {title}")
            
            module_result = getattr(analysis_state, module_attr, None)
            if module_result and hasattr(module_result, 'status'):
                self.doc.add_heading(title, level=1)
                
                # è·å–æ¨¡å—çš„è¯¦ç»†åˆ†æå†…å®¹
                self._add_module_content(module_result, module_key)
                
                # æ·»åŠ å›¾è¡¨ï¼ˆå¦‚æœæœ‰ä¸”å¯ç”¨äº†å›¾è¡¨ï¼‰
                if self.include_charts:
                    if progress_callback:
                        progress_callback(f"  â””â”€ è½¬æ¢{title}çš„å›¾è¡¨...")
                    self._add_module_charts(module_result, module_key)
                else:
                    # å¿«é€Ÿæ¨¡å¼ï¼šæ·»åŠ å›¾è¡¨è¯´æ˜
                    note_para = self.doc.add_paragraph()
                    note_run = note_para.add_run("æ³¨ï¼šå¿«é€Ÿæ¨¡å¼å·²è·³è¿‡å›¾è¡¨ï¼Œè¯¦ç»†å›¾è¡¨è¯·æŸ¥çœ‹ç³»ç»Ÿç•Œé¢ã€‚")
                    note_run.font.name = 'å®‹ä½“'
                    note_run.font.size = Pt(10)
                    note_run.font.italic = True
                
                self.doc.add_page_break()
    
    def _add_module_content(self, module_result, module_key: str):
        """æ·»åŠ æ¨¡å—åˆ†æå†…å®¹"""
        
        try:
            # è·å–åˆ†æå†…å®¹
            if hasattr(module_result, 'result_data') and module_result.result_data:
                result_data = module_result.result_data
                
                # è·å–AIåˆ†æå†…å®¹
                ai_content = None
                content_keys = [
                    'ai_comprehensive_analysis',  # åº“å­˜åˆ†æ
                    'ai_analysis',                # æŒä»“ã€åŸºå·®åˆ†æ
                    'professional_analysis',      # æŠ€æœ¯åˆ†æ
                    'comprehensive_analysis'      # å…¶ä»–åˆ†æ
                ]
                
                for key in content_keys:
                    if key in result_data and result_data[key]:
                        ai_content = result_data[key]
                        break
                
                if ai_content:
                    # æ¸…ç†å†…å®¹ï¼ˆå»é™¤Markdownç¬¦å·ç­‰ï¼‰
                    cleaned_content = self._clean_analysis_content(ai_content)
                    
                    # æ·»åŠ åˆ†æå†…å®¹
                    content_para = self.doc.add_paragraph()
                    content_run = content_para.add_run(cleaned_content)
                    content_run.font.name = 'å®‹ä½“'
                    content_run.font.size = Pt(12)
                else:
                    # å¦‚æœæ²¡æœ‰AIå†…å®¹ï¼Œæ·»åŠ åŸºæœ¬ä¿¡æ¯
                    basic_info = f"""
                    æ¨¡å—çŠ¶æ€ï¼š{getattr(module_result, 'status', 'æœªçŸ¥')}
                    ç½®ä¿¡åº¦ï¼š{getattr(module_result, 'confidence_score', 'æœªçŸ¥')}
                    æ‰§è¡Œæ—¶é—´ï¼š{getattr(module_result, 'execution_time', 'æœªçŸ¥')}ç§’
                    """
                    self.doc.add_paragraph(basic_info.strip())
            else:
                self.doc.add_paragraph("æš‚æ— è¯¦ç»†åˆ†æå†…å®¹")
                
        except Exception as e:
            self.doc.add_paragraph(f"å†…å®¹è§£æå‡ºé”™ï¼š{str(e)}")
    
    def _clean_analysis_content(self, content: str) -> str:
        """æ¸…ç†åˆ†æå†…å®¹ï¼Œå»é™¤Markdownç¬¦å·ç­‰"""
        
        if not isinstance(content, str):
            return str(content)
        
        import re
        
        # å»é™¤Markdownç¬¦å·
        content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)  # ç²—ä½“
        content = re.sub(r'\*([^*]+)\*', r'\1', content)      # æ–œä½“
        content = re.sub(r'`([^`]+)`', r'\1', content)        # ä»£ç 
        content = re.sub(r'#+\s*', '', content)               # æ ‡é¢˜ç¬¦å·
        content = re.sub(r'[-*+]\s+', 'â€¢ ', content)          # åˆ—è¡¨ç¬¦å·
        
        # æ¸…ç†å¤šä½™çš„ç©ºè¡Œ
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        return content.strip()
    
    def _add_module_charts(self, module_result, module_key: str):
        """æ·»åŠ æ¨¡å—å›¾è¡¨ - å¢å¼ºç‰ˆæœ¬ï¼Œæ”¯æŒPlotlyå›¾è¡¨è½¬æ¢"""
        
        try:
            if hasattr(module_result, 'result_data') and module_result.result_data:
                result_data = module_result.result_data
                
                # ğŸ”¥ æ–°å¢ï¼šå°è¯•ä½¿ç”¨å›¾è¡¨å¢å¼ºå™¨å¤„ç†Plotlyå›¾è¡¨
                charts_added = self._process_plotly_charts(result_data, module_key)
                
                if not charts_added:
                    # æ£€æŸ¥æ˜¯å¦æœ‰ä¼ ç»Ÿçš„å›¾è¡¨HTMLæˆ–å›¾è¡¨æ•°æ®
                    charts_html = result_data.get('charts_html')
                    charts = result_data.get('charts', [])
                    
                    if charts_html or charts:
                        # æ·»åŠ å›¾è¡¨æ ‡é¢˜
                        chart_title = self.doc.add_heading(f'{module_key}æ¨¡å—å›¾è¡¨', level=3)
                        
                        # å°è¯•æ’å…¥å®é™…çš„å›¾è¡¨æ–‡ä»¶
                        chart_files_added = self._insert_chart_files(module_key, result_data)
                        
                        if not chart_files_added:
                            # å¦‚æœæ²¡æœ‰æ‰¾åˆ°å›¾è¡¨æ–‡ä»¶ï¼Œæ·»åŠ è¯´æ˜
                            chart_para = self.doc.add_paragraph()
                            chart_run = chart_para.add_run(f"æ³¨ï¼š{module_key}æ¨¡å—åŒ…å«äº¤äº’å¼å›¾è¡¨ï¼Œè¯¦ç»†å›¾è¡¨è¯·æŸ¥çœ‹ç³»ç»Ÿç•Œé¢ã€‚")
                            chart_run.font.name = 'å®‹ä½“'
                            chart_run.font.size = Pt(10)
                            chart_run.font.italic = True
                
        except Exception as e:
            # æ·»åŠ é”™è¯¯è¯´æ˜è€Œä¸æ˜¯é™é»˜å¿½ç•¥
            try:
                error_para = self.doc.add_paragraph()
                error_run = error_para.add_run(f"æ³¨ï¼š{module_key}æ¨¡å—å›¾è¡¨åŠ è½½é‡åˆ°é—®é¢˜ï¼Œè¯·æŸ¥çœ‹ç³»ç»Ÿç•Œé¢è·å–å›¾è¡¨ã€‚")
                error_run.font.name = 'å®‹ä½“'
                error_run.font.size = Pt(9)
                error_run.font.italic = True
                error_run.font.color.rgb = RGBColor(128, 128, 128)
            except:
                pass
    
    def _process_plotly_charts(self, result_data: Dict, module_key: str) -> bool:
        """å¤„ç†Plotlyå›¾è¡¨å¹¶è½¬æ¢ä¸ºPNGå›¾ç‰‡"""
        
        try:
            import plotly.io as pio
            import tempfile
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            charts_added = 0
            
            # æŸ¥æ‰¾å›¾è¡¨æ•°æ®çš„å¤šç§å¯èƒ½å­—æ®µ
            chart_fields = [
                'charts', 'professional_charts', 'chart_data', 'figures', 
                'visualizations', 'plots', 'graphs'
            ]
            
            for field in chart_fields:
                if isinstance(result_data, dict) and field in result_data:
                    charts_data = result_data[field]
                    if charts_data:
                        # æ·»åŠ æ¨¡å—å›¾è¡¨æ ‡é¢˜
                        if charts_added == 0:  # åªæ·»åŠ ä¸€æ¬¡æ ‡é¢˜
                            heading = self.doc.add_heading(f'{module_key}æ¨¡å—å›¾è¡¨', level=3)
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # å¤„ç†å›¾è¡¨æ•°æ®
                        if isinstance(charts_data, dict):
                            # å¤„ç†å­—å…¸å½¢å¼çš„å›¾è¡¨æ•°æ®
                            for chart_key, chart_obj in charts_data.items():
                                if self._is_plotly_figure(chart_obj):
                                    success = self._convert_and_add_plotly_chart(chart_obj, chart_key, module_key)
                                    if success:
                                        charts_added += 1
                        
                        elif isinstance(charts_data, list):
                            # å¤„ç†åˆ—è¡¨å½¢å¼çš„å›¾è¡¨æ•°æ®
                            for i, chart_obj in enumerate(charts_data):
                                if self._is_plotly_figure(chart_obj):
                                    chart_name = f"å›¾è¡¨_{i+1}"
                                    success = self._convert_and_add_plotly_chart(chart_obj, chart_name, module_key)
                                    if success:
                                        charts_added += 1
                        
                        elif self._is_plotly_figure(charts_data):
                            # å¤„ç†å•ä¸ªå›¾è¡¨å¯¹è±¡
                            success = self._convert_and_add_plotly_chart(charts_data, "ä¸»è¦å›¾è¡¨", module_key)
                            if success:
                                charts_added += 1
                        
                        if charts_added > 0:
                            break  # æ‰¾åˆ°å›¾è¡¨åå°±ä¸å†æŸ¥æ‰¾å…¶ä»–å­—æ®µ
            
            if charts_added > 0:
                print(f"âœ… {module_key}: æˆåŠŸæ·»åŠ  {charts_added} ä¸ªå›¾è¡¨åˆ°WordæŠ¥å‘Š")
            
            return charts_added > 0
            
        except Exception as e:
            print(f"âŒ å¤„ç† {module_key} Plotlyå›¾è¡¨æ—¶å‡ºé”™: {e}")
            return False
    
    def _is_plotly_figure(self, obj: Any) -> bool:
        """æ£€æŸ¥æ˜¯å¦ä¸ºPlotlyå›¾è¡¨å¯¹è±¡"""
        return hasattr(obj, 'to_image') or hasattr(obj, 'to_html') or str(type(obj)).find('plotly') != -1
    
    def _convert_and_add_plotly_chart(self, chart_obj: Any, chart_name: str, module_name: str) -> bool:
        """å°†Plotlyå›¾è¡¨è½¬æ¢ä¸ºPNGå¹¶æ·»åŠ åˆ°Wordæ–‡æ¡£ï¼ˆå¸¦è¶…æ—¶ä¿æŠ¤ï¼‰"""
        
        try:
            import plotly.io as pio
            import tempfile
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import signal
            
            # æ¸…ç†å›¾è¡¨åç§°ï¼Œç§»é™¤æŠ€æœ¯ä»£ç 
            clean_name = self._clean_chart_name(chart_name)
            
            # ğŸ”¥ è¶…æ—¶ä¿æŠ¤ï¼šå¦‚æœè½¬æ¢è¶…è¿‡30ç§’ï¼Œåˆ™è·³è¿‡
            def timeout_handler(signum, frame):
                raise TimeoutError("å›¾è¡¨è½¬æ¢è¶…æ—¶")
            
            # Windowsä¸æ”¯æŒsignal.alarmï¼Œä½¿ç”¨try-exceptåŒ…è£¹
            try:
                # å°†Plotlyå›¾è¡¨è½¬æ¢ä¸ºPNGå­—èŠ‚
                img_bytes = pio.to_image(chart_obj, format='png', width=800, height=600, scale=2)
            except Exception as convert_error:
                print(f"âš ï¸ å›¾è¡¨è½¬æ¢å¤±è´¥ ({chart_name}): {convert_error}")
                # æ·»åŠ å›¾è¡¨è½¬æ¢å¤±è´¥è¯´æ˜
                note_para = self.doc.add_paragraph()
                note_run = note_para.add_run(f"æ³¨ï¼š{clean_name} å›¾è¡¨è½¬æ¢å¤±è´¥ï¼Œè¯·åœ¨ç³»ç»Ÿç•Œé¢æŸ¥çœ‹äº¤äº’å¼å›¾è¡¨ã€‚")
                note_run.font.name = 'å®‹ä½“'
                note_run.font.size = Pt(9)
                note_run.font.italic = True
                note_run.font.color.rgb = RGBColor(128, 128, 128)
                return False
            
            # åˆ›å»ºä¸´æ—¶æ–‡ä»¶
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                tmp_file.write(img_bytes)
                tmp_file_path = tmp_file.name
            
            try:
                # æ·»åŠ å›¾è¡¨æ ‡é¢˜
                if clean_name and clean_name != "ä¸»è¦å›¾è¡¨":
                    chart_title_para = self.doc.add_paragraph()
                    chart_title_run = chart_title_para.add_run(f"å›¾: {clean_name}")
                    chart_title_run.font.name = 'å®‹ä½“'
                    chart_title_run.font.size = Pt(12)
                    chart_title_run.bold = True
                    chart_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # æ·»åŠ å›¾ç‰‡åˆ°æ–‡æ¡£
                chart_para = self.doc.add_paragraph()
                chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_para.add_run().add_picture(tmp_file_path, width=Inches(6))
                
                # æ·»åŠ å›¾è¡¨è¯´æ˜
                caption_para = self.doc.add_paragraph()
                caption_run = caption_para.add_run(f"æ•°æ®æ¥æº: {module_name}æ¨¡å—åˆ†æç»“æœ")
                caption_run.font.name = 'å®‹ä½“'
                caption_run.font.size = Pt(9)
                caption_run.italic = True
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # æ·»åŠ é—´è·
                self.doc.add_paragraph()
                
                print(f"âœ… æˆåŠŸæ·»åŠ å›¾è¡¨: {clean_name}")
                return True
                
            finally:
                # æ¸…ç†ä¸´æ—¶æ–‡ä»¶
                try:
                    os.unlink(tmp_file_path)
                except:
                    pass
                    
        except Exception as e:
            print(f"âŒ è½¬æ¢Plotlyå›¾è¡¨å¤±è´¥ ({chart_name}): {e}")
            # æ·»åŠ é”™è¯¯è¯´æ˜
            try:
                note_para = self.doc.add_paragraph()
                note_run = note_para.add_run(f"æ³¨ï¼š{chart_name} å›¾è¡¨å¤„ç†å‡ºé”™ï¼Œè¯·åœ¨ç³»ç»Ÿç•Œé¢æŸ¥çœ‹ã€‚")
                note_run.font.name = 'å®‹ä½“'
                note_run.font.size = Pt(9)
                note_run.font.italic = True
                note_run.font.color.rgb = RGBColor(128, 128, 128)
            except:
                pass
            return False
    
    def _clean_chart_name(self, chart_name: str) -> str:
        """æ¸…ç†å›¾è¡¨åç§°ï¼Œç§»é™¤æŠ€æœ¯ä»£ç """
        
        if not isinstance(chart_name, str):
            return "å›¾è¡¨"
        
        # ç§»é™¤å¸¸è§çš„æŠ€æœ¯å‰ç¼€
        prefixes_to_remove = ['ğŸ“Š ', 'ğŸ“ˆ ', 'ğŸ“‰ ', 'ğŸ” ', 'ğŸ’¹ ']
        clean_name = chart_name
        for prefix in prefixes_to_remove:
            if clean_name.startswith(prefix):
                clean_name = clean_name[len(prefix):]
        
        # ç§»é™¤æˆ–æ›¿æ¢æŠ€æœ¯æœ¯è¯­
        replacements = {
            'price_trend': 'ä»·æ ¼èµ°åŠ¿',
            'spider_web': 'èœ˜è››ç½‘ç­–ç•¥åˆ†æ',
            'smart_money': 'èªæ˜é’±åˆ†æ',
            'family_reverse': 'å®¶äººå¸­ä½åå‘åˆ†æ',
            'seat_behavior': 'å¸­ä½è¡Œä¸ºåˆ†æ',
            'force_change': 'å¤šç©ºåŠ›é‡å˜åŒ–',
            'concentration': 'æŒä»“é›†ä¸­åº¦åˆ†æ',
            'basis_analysis': 'åŸºå·®åˆ†æ',
            'term_structure': 'æœŸé™ç»“æ„åˆ†æ',
            'technical_indicators': 'æŠ€æœ¯æŒ‡æ ‡åˆ†æ',
            'inventory_trend': 'åº“å­˜è¶‹åŠ¿åˆ†æ'
        }
        
        for tech_term, chinese_name in replacements.items():
            if tech_term in clean_name.lower():
                clean_name = chinese_name
                break
        
        return clean_name.strip()
    
    def _insert_chart_files(self, module_key: str, result_data: Dict) -> bool:
        """æ’å…¥å›¾è¡¨æ–‡ä»¶åˆ°Wordæ–‡æ¡£"""
        
        try:
            from docx.shared import Inches
            
            chart_files_added = 0
            
            # 1. æ£€æŸ¥result_dataä¸­æ˜¯å¦æœ‰å›¾è¡¨è·¯å¾„ä¿¡æ¯
            charts = result_data.get('charts', [])
            if isinstance(charts, list) and charts:
                for chart_info in charts:
                    if isinstance(chart_info, dict) and 'path' in chart_info:
                        chart_path = chart_info['path']
                        if os.path.exists(chart_path):
                            self._add_image_to_doc(chart_path, chart_info.get('title', 'å›¾è¡¨'))
                            chart_files_added += 1
            
            # 2. æ ¹æ®æ¨¡å—ç±»å‹å’Œå“ç§æŸ¥æ‰¾å¯¹åº”çš„å›¾è¡¨æ–‡ä»¶
            commodity = result_data.get('commodity', 'UNKNOWN')
            chart_patterns = self._get_chart_patterns(module_key, commodity)
            
            for pattern in chart_patterns:
                matching_files = []
                # åœ¨å½“å‰ç›®å½•æŸ¥æ‰¾åŒ¹é…çš„æ–‡ä»¶
                for filename in os.listdir('.'):
                    if pattern.lower() in filename.lower() and filename.endswith(('.png', '.jpg', '.jpeg')):
                        matching_files.append(filename)
                
                # æ’å…¥æ‰¾åˆ°çš„å›¾è¡¨æ–‡ä»¶
                for chart_file in matching_files[:3]:  # é™åˆ¶æ¯ç§ç±»å‹æœ€å¤š3ä¸ªå›¾è¡¨
                    if os.path.exists(chart_file):
                        # ä»æ–‡ä»¶åæ¨æ–­å›¾è¡¨æ ‡é¢˜
                        chart_title = self._generate_chart_title(chart_file, module_key)
                        self._add_image_to_doc(chart_file, chart_title)
                        chart_files_added += 1
            
            return chart_files_added > 0
            
        except Exception as e:
            print(f"æ’å…¥å›¾è¡¨æ–‡ä»¶æ—¶å‡ºé”™: {e}")
            return False
    
    def _get_chart_patterns(self, module_key: str, commodity: str) -> list:
        """è·å–æ¨¡å—å¯¹åº”çš„å›¾è¡¨æ–‡ä»¶åæ¨¡å¼"""
        
        patterns = []
        
        # æ ¹æ®æ¨¡å—ç±»å‹å®šä¹‰å›¾è¡¨æ–‡ä»¶åæ¨¡å¼
        if module_key == "åº“å­˜ä»“å•åˆ†æ":
            patterns = [
                f"{commodity}_inventory",
                f"{commodity}_receipt", 
                f"{commodity}_ratio",
                "inventory_trend",
                "receipt_trend",
                "inventory_receipt"
            ]
        elif module_key == "æŒä»“å¸­ä½åˆ†æ":
            patterns = [
                f"{commodity}_positioning",
                f"{commodity}_spider",
                f"{commodity}_smart_money",
                "positioning_analysis",
                "seat_analysis"
            ]
        elif module_key == "æŠ€æœ¯é¢åˆ†æ":
            patterns = [
                f"{commodity}_technical",
                f"{commodity}_price",
                "technical_analysis"
            ]
        elif module_key == "æœŸé™ç»“æ„åˆ†æ":
            patterns = [
                f"{commodity}_term_structure",
                f"{commodity}_curve",
                "term_structure"
            ]
        elif module_key == "åŸºå·®åˆ†æ":
            patterns = [
                f"{commodity}_basis",
                "basis_analysis"
            ]
        
        return patterns
    
    def _generate_chart_title(self, filename: str, module_key: str) -> str:
        """ä»æ–‡ä»¶åç”Ÿæˆå›¾è¡¨æ ‡é¢˜"""
        
        # ç§»é™¤æ–‡ä»¶æ‰©å±•å
        name = os.path.splitext(filename)[0]
        
        # æ ¹æ®æ–‡ä»¶åå…³é”®è¯ç”Ÿæˆæ ‡é¢˜
        if "inventory" in name.lower():
            if "trend" in name.lower():
                return "åº“å­˜è¶‹åŠ¿å›¾"
            elif "ratio" in name.lower():
                return "åº“å­˜ä»“å•æ¯”ç‡å›¾"
            elif "comparison" in name.lower():
                return "ä»·æ ¼åº“å­˜å¯¹æ¯”å›¾"
            else:
                return "åº“å­˜åˆ†æå›¾"
        elif "receipt" in name.lower():
            return "ä»“å•åˆ†æå›¾"
        elif "positioning" in name.lower():
            return "æŒä»“åˆ†æå›¾"
        elif "technical" in name.lower():
            return "æŠ€æœ¯åˆ†æå›¾"
        elif "term_structure" in name.lower():
            return "æœŸé™ç»“æ„å›¾"
        elif "basis" in name.lower():
            return "åŸºå·®åˆ†æå›¾"
        else:
            return f"{module_key}å›¾è¡¨"
    
    def _add_image_to_doc(self, image_path: str, title: str):
        """æ·»åŠ å›¾ç‰‡åˆ°Wordæ–‡æ¡£"""
        
        try:
            from docx.shared import Inches
            
            # æ·»åŠ å›¾è¡¨æ ‡é¢˜
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.name = 'é»‘ä½“'
            title_run.font.size = Pt(12)
            title_run.bold = True
            title_para.alignment = 1  # å±…ä¸­å¯¹é½
            
            # æ·»åŠ å›¾ç‰‡
            img_para = self.doc.add_paragraph()
            img_para.alignment = 1  # å±…ä¸­å¯¹é½
            
            # æ’å…¥å›¾ç‰‡ï¼Œè®¾ç½®åˆé€‚çš„å°ºå¯¸
            run = img_para.add_run()
            run.add_picture(image_path, width=Inches(6))  # 6è‹±å¯¸å®½åº¦
            
            # æ·»åŠ ç©ºè¡Œ
            self.doc.add_paragraph()
            
        except Exception as e:
            # å¦‚æœå›¾ç‰‡æ’å…¥å¤±è´¥ï¼Œæ·»åŠ æ–‡å­—è¯´æ˜
            error_para = self.doc.add_paragraph()
            error_run = error_para.add_run(f"[å›¾è¡¨: {title}] - å›¾ç‰‡åŠ è½½å¤±è´¥")
            error_run.font.name = 'å®‹ä½“'
            error_run.font.size = Pt(10)
            error_run.font.italic = True
    
    def _add_cio_conclusion(self, analysis_results: Dict, commodity: str):
        """æ·»åŠ CIOå†³ç­–ç»“è®º"""
        
        self.doc.add_heading('ä¸ƒã€æŠ•èµ„å†³ç­–ç»“è®º', level=1)
        
        # è·å–ç¬¬ä¸€ä¸ªå“ç§çš„åˆ†æç»“æœ
        result = list(analysis_results.values())[0]
        if result.get("status") != "success":
            self.doc.add_paragraph("åˆ†ææœªæˆåŠŸå®Œæˆï¼Œæ— æ³•ç”ŸæˆæŠ•èµ„å†³ç­–ç»“è®ºã€‚")
            return
        
        # æ£€æŸ¥æ˜¯å¦æœ‰è¾©è®ºé£æ§å†³ç­–ç»“æœ
        if result.get("type") == "complete_flow":
            # ğŸ”¥ ä¿®å¤ï¼šç›´æ¥ä»æ–°æ•°æ®ç»“æ„è·å–decision_section
            debate_result = result.get("result", {}).get("debate_result", {})
            
            # æ£€æŸ¥åˆ†ææ˜¯å¦æˆåŠŸ
            if not debate_result or not debate_result.get("success"):
                print(f"âŒ æŠ¥å‘Šç”Ÿæˆé”™è¯¯ï¼šåˆ†æå¤±è´¥ - {debate_result.get('error', 'æœªçŸ¥é”™è¯¯') if debate_result else 'åˆ†æç»“æœä¸ºç©º'}")
                return
            
            executive_decision = debate_result.get("decision_section", {})
            
            # ç¡®ä¿æ•°æ®å®Œæ•´æ€§
            if not executive_decision:
                print("âŒ æŠ¥å‘Šç”Ÿæˆé”™è¯¯ï¼šCIOå†³ç­–æ•°æ®ç¼ºå¤±")
                return
            
            if debate_result.get("success") and executive_decision:
                decision = executive_decision
                
                # CIOæœ€ç»ˆå†³ç­–
                self.doc.add_heading('7.1 CIOæœ€ç»ˆæŠ•èµ„å†³ç­–', level=2)
                
                decision_info = f"""
                æœ€ç»ˆå†³ç­–ï¼š{decision.get('final_decision', 'æœªçŸ¥')}
                å†³ç­–ä¿¡å¿ƒåº¦ï¼š{decision.get('confidence', 'æœªçŸ¥')}
                å»ºè®®ä»“ä½ï¼š{decision.get('position_size', 'æœªçŸ¥')}
                å†³ç­–ä¾æ®ï¼š
                {decision.get('rationale', 'æš‚æ— è¯¦ç»†ä¾æ®')}
                
                æ‰§è¡Œè®¡åˆ’ï¼š
                {decision.get('execution_plan', 'æš‚æ— æ‰§è¡Œè®¡åˆ’')}
                
                ç›‘æ§è¦ç‚¹ï¼š
                {' | '.join(decision.get('monitoring_points', ['æš‚æ— ç›‘æ§è¦ç‚¹']))}
                """
                
                decision_para = self.doc.add_paragraph()
                decision_run = decision_para.add_run(decision_info.strip())
                decision_run.font.name = 'å®‹ä½“'
                decision_run.font.size = Pt(12)
                
                # CIOæƒå¨å£°æ˜
                if decision.get('cio_statement'):
                    self.doc.add_heading('7.2 CIOæƒå¨å£°æ˜', level=2)
                    
                    statement_para = self.doc.add_paragraph()
                    statement_run = statement_para.add_run(decision['cio_statement'])
                    statement_run.font.name = 'å®‹ä½“'
                    statement_run.font.size = Pt(12)
                    statement_run.font.bold = True
            else:
                self.doc.add_paragraph("æœªå®Œæˆå®Œæ•´å†³ç­–æµç¨‹ï¼Œæ— CIOå†³ç­–ç»“è®ºã€‚")
        else:
            self.doc.add_paragraph("æœ¬æ¬¡åˆ†æä¸ºåŸºç¡€åˆ†ææ¨¡å¼ï¼ŒæœªåŒ…å«CIOå†³ç­–ç¯èŠ‚ã€‚")
    
    def _add_title_page(self, config: Dict):
        """æ·»åŠ æ ‡é¢˜é¡µ"""
        
        # ä¸»æ ‡é¢˜
        title = self.doc.add_heading('æœŸè´§Trading Agentsç³»ç»Ÿ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # å‰¯æ ‡é¢˜
        subtitle = self.doc.add_heading('ä¸“ä¸šå®Œæ•´åˆ†ææŠ¥å‘Šï¼ˆå«äº¤æ˜“å‘˜ç¯èŠ‚ï¼‰', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # åˆ†æä¿¡æ¯
        self.doc.add_paragraph()
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        analysis_date = config.get("analysis_date", "æœªçŸ¥æ—¥æœŸ")
        analysis_mode = config.get("analysis_mode", "æœªçŸ¥æ¨¡å¼")
        commodities = config.get("commodities", [])
        
        info_text = f"""
        åˆ†ææ—¥æœŸï¼š{analysis_date}
        åˆ†ææ¨¡å¼ï¼š{analysis_mode}
        åˆ†æå“ç§ï¼š{', '.join(commodities)}
        æŠ¥å‘Šç”Ÿæˆæ—¶é—´ï¼š{datetime.now().strftime('%Yå¹´%mæœˆ%dæ—¥ %H:%M:%S')}
        
        å®Œæ•´5é˜¶æ®µæµç¨‹ï¼š
        é˜¶æ®µ1ï¼šåˆ†æå¸ˆå›¢é˜Ÿï¼ˆ6å¤§åˆ†ææ¨¡å—ï¼‰
        é˜¶æ®µ2ï¼šæ¿€çƒˆè¾©è®ºï¼ˆå¤šç©ºè§‚ç‚¹äº¤é”‹ï¼‰
        é˜¶æ®µ3ï¼šä¸“ä¸šäº¤æ˜“å‘˜ï¼ˆå®¢è§‚ç†æ€§åˆ¶å®šç­–ç•¥ï¼‰
        é˜¶æ®µ4ï¼šä¸“ä¸šé£æ§ï¼ˆç‹¬ç«‹é£é™©è¯„ä¼°ï¼‰
        é˜¶æ®µ5ï¼šCIOå†³ç­–ï¼ˆæƒå¨æœ€ç»ˆå†³ç­–ï¼‰
        """
        
        info_para.add_run(info_text)
        
        # æ·»åŠ åˆ†é¡µç¬¦
        self.doc.add_page_break()
    
    def _add_table_of_contents(self, analysis_results: Dict):
        """æ·»åŠ ç›®å½•"""
        
        self.doc.add_heading('ç›®å½•', level=1)
        
        toc_items = [
            "1. æ‰§è¡Œæ‘˜è¦",
            "2. å“ç§åˆ†æè¯¦æƒ…"
        ]
        
        for i, commodity in enumerate(analysis_results.keys(), 3):
            toc_items.append(f"  2.{i-2} {commodity}å®Œæ•´åˆ†ææŠ¥å‘Š")
        
        toc_items.append(f"{len(analysis_results) + 3}. ç»¼åˆç»“è®ºä¸æŠ•èµ„å»ºè®®")
        
        for item in toc_items:
            para = self.doc.add_paragraph(item)
            if not item.startswith("  "):
                para.style = 'List Bullet'
        
        self.doc.add_page_break()
    
    def _add_executive_summary(self, analysis_results: Dict, config: Dict):
        """æ·»åŠ æ‰§è¡Œæ‘˜è¦"""
        
        self.doc.add_heading('1. æ‰§è¡Œæ‘˜è¦', level=1)
        
        # åˆ†ææ¦‚è§ˆ
        self.doc.add_heading('1.1 åˆ†ææ¦‚è§ˆ', level=2)
        
        total_commodities = len(analysis_results)
        successful_analyses = sum(1 for result in analysis_results.values() 
                                if result.get("status") == "success")
        
        overview_text = f"""
        æœ¬æ¬¡åˆ†æå…±æ¶µç›–{total_commodities}ä¸ªæœŸè´§å“ç§ï¼ŒæˆåŠŸå®Œæˆ{successful_analyses}ä¸ªå“ç§çš„æ·±åº¦åˆ†æã€‚
        åˆ†ææ¨¡å¼ï¼š{config.get('analysis_mode', 'æœªçŸ¥')}
        åˆ†ææ—¥æœŸï¼š{config.get('analysis_date', 'æœªçŸ¥')}
        ä½¿ç”¨æ¨¡å—ï¼š{', '.join(config.get('modules', []))}
        
        ç³»ç»Ÿé‡‡ç”¨å®Œæ•´çš„5é˜¶æ®µä¸“ä¸šåˆ†ææµç¨‹ï¼Œç¡®ä¿æŠ•èµ„å†³ç­–çš„ç§‘å­¦æ€§å’Œä¸“ä¸šæ€§ã€‚
        æ¯ä¸ªå“ç§éƒ½ç»è¿‡äº†ä»åŸºç¡€æ•°æ®åˆ†æåˆ°æœ€ç»ˆæŠ•èµ„å†³ç­–çš„å®Œæ•´æµç¨‹ã€‚
        """
        
        self.doc.add_paragraph(overview_text)
        
        # ä¸»è¦å‘ç°
        self.doc.add_heading('1.2 ä¸»è¦å‘ç°', level=2)
        
        findings = []
        for commodity, result in analysis_results.items():
            if result.get("status") == "success":
                if result.get("type") == "optimized_debate":
                    # ä¿®æ­£ï¼šç›´æ¥ä»resultè·å–æ•°æ®
                    debate_result = result.get("result", {}).get("debate_result", {})
                    if debate_result and debate_result.get("success"):
                        # ğŸ”¥ ä¿®å¤ï¼šä»æ­£ç¡®å±‚çº§è·å–å†³ç­–æ•°æ®
                        executive_dec = result.get("result", {}).get("executive_decision", {})
                        decision = executive_dec.get("final_decision", "æœªçŸ¥")
                        confidence = executive_dec.get("confidence_level", "æœªçŸ¥")
                        findings.append(f"{commodity}ï¼š{decision}ï¼ˆä¿¡å¿ƒåº¦ï¼š{confidence}ï¼‰")
                else:
                    findings.append(f"{commodity}ï¼šåˆ†æå®Œæˆ")
        
        for finding in findings:
            para = self.doc.add_paragraph(finding)
            para.style = 'List Bullet'
        
        self.doc.add_page_break()
    
    def _add_commodity_analysis(self, commodity: str, result: Dict, config: Dict):
        """æ·»åŠ å“ç§åˆ†æ"""
        
        self.doc.add_heading(f'2. {commodity}å®Œæ•´åˆ†ææŠ¥å‘Š', level=1)
        
        if result.get("status") != "success":
            error_para = self.doc.add_paragraph(f"åˆ†æå¤±è´¥ï¼š{result.get('message', 'æœªçŸ¥é”™è¯¯')}")
            return
        
        result_type = result.get("type")
        
        if result_type == "optimized_debate":
            self._add_optimized_debate_analysis(commodity, result)
        elif result_type == "complete_flow":
            self._add_complete_flow_analysis(commodity, result)
        else:
            self._add_analyst_only_analysis(commodity, result)
    
    def _add_optimized_debate_analysis(self, commodity: str, result: Dict):
        """æ·»åŠ ä¼˜åŒ–ç‰ˆè¾©è®ºåˆ†æ"""
        
        analysis_state = result["result"]["analysis_state"]
        debate_result = result["result"]["debate_result"]
        
        # åˆ†æå¸ˆå›¢é˜Ÿç»“æœ
        self.doc.add_heading('2.1 åˆ†æå¸ˆå›¢é˜Ÿä¸“ä¸šæŠ¥å‘Š', level=2)
        
        # æ·»åŠ å„æ¨¡å—åˆ†æç»“æœ
        modules = {
            "inventory_analysis": "åº“å­˜ä»“å•åˆ†æ",
            "positioning_analysis": "æŒä»“å¸­ä½åˆ†æ",
            "term_structure_analysis": "æœŸé™ç»“æ„åˆ†æ",
            "technical_analysis": "æŠ€æœ¯é¢åˆ†æ",
            "basis_analysis": "åŸºå·®åˆ†æ",
            "news_analysis": "æ–°é—»åˆ†æ"
        }
        
        for module_attr, module_name in modules.items():
            module_result = getattr(analysis_state, module_attr, None)
            if module_result and hasattr(module_result, 'status'):
                self.doc.add_heading(f'2.1.{list(modules.keys()).index(module_attr) + 1} {module_name}', level=3)
                
                # è¿™é‡Œå¯ä»¥æ·»åŠ æ›´è¯¦ç»†çš„åˆ†æå†…å®¹
                if hasattr(module_result, 'confidence_score'):
                    try:
                        confidence = module_result.confidence_score
                        if isinstance(confidence, (int, float)):
                            module_text = f"{module_name}å·²å®Œæˆï¼Œç½®ä¿¡åº¦ï¼š{confidence:.1%}"
                        else:
                            module_text = f"{module_name}å·²å®Œæˆï¼Œç½®ä¿¡åº¦ï¼š{confidence}"
                    except (ValueError, TypeError):
                        module_text = f"{module_name}å·²å®Œæˆ"
                else:
                    module_text = f"{module_name}å·²å®Œæˆ"
                self.doc.add_paragraph(module_text)
        
        # ä¼˜åŒ–ç‰ˆè¾©è®ºé£æ§å†³ç­–ç»“æœ
        if debate_result.get("success"):
            self.doc.add_heading('2.2 å®Œæ•´å†³ç­–æµç¨‹æŠ¥å‘Š', level=2)
            
            # è¾©è®ºç»“æœ
            if "debate_section" in debate_result:
                debate = debate_result["debate_section"]
                self.doc.add_heading('2.2.1 æ¿€çƒˆè¾©è®ºç¯èŠ‚', level=3)
                
                debate_text = f"""
                è¾©è®ºè½®æ•°ï¼š{debate.get('title', 'æœªçŸ¥')}
                è¾©è®ºèƒœè€…ï¼š{debate.get('winner', 'æœªçŸ¥')}
                å¤šå¤´å¾—åˆ†ï¼š{debate.get('scores', {}).get('bull', 0):.1f}åˆ†
                ç©ºå¤´å¾—åˆ†ï¼š{debate.get('scores', {}).get('bear', 0):.1f}åˆ†
                
                è¾©è®ºæ€»ç»“ï¼š
                {debate.get('summary', 'æš‚æ— æ€»ç»“')}
                
                è¾¾æˆå…±è¯†ï¼š
                {' | '.join(debate.get('consensus_points', []))}
                
                äº‰è®®ç„¦ç‚¹ï¼š
                {' | '.join(debate.get('unresolved_issues', []))}
                """
                
                self.doc.add_paragraph(debate_text)
            
            # äº¤æ˜“å‘˜å†³ç­–
            if "trading_section" in debate_result:
                trading = debate_result["trading_section"]
                self.doc.add_heading('2.2.2 ä¸“ä¸šäº¤æ˜“å‘˜å†³ç­–', level=3)
                
                trading_text = f"""
                ç­–ç•¥ç±»å‹ï¼š{trading.get('strategy_type', 'æœªçŸ¥')}
                äº¤æ˜“é€»è¾‘ï¼š{trading.get('reasoning', 'æš‚æ— é€»è¾‘')}
                è¿›åœºç‚¹ä½ï¼š{' | '.join(trading.get('entry_points', []))}
                å‡ºåœºç‚¹ä½ï¼š{' | '.join(trading.get('exit_points', []))}
                ä»“ä½ç®¡ç†ï¼š{trading.get('position_size', 'æœªçŸ¥')}
                é£é™©æ”¶ç›Šæ¯”ï¼š{trading.get('risk_reward_ratio', 'æœªçŸ¥')}
                å…·ä½“åˆçº¦ï¼š{' | '.join(trading.get('specific_contracts', []))}
                æ‰§è¡Œè®¡åˆ’ï¼š{trading.get('execution_plan', 'æš‚æ— è®¡åˆ’')}
                å¸‚åœºæ¡ä»¶ï¼š{trading.get('market_conditions', 'æœªçŸ¥')}
                å¤‡é€‰æ–¹æ¡ˆï¼š{' | '.join(trading.get('alternative_scenarios', []))}
                """
                
                self.doc.add_paragraph(trading_text)
            
            # é£é™©è¯„ä¼°
            if "risk_section" in debate_result:
                risk = debate_result["risk_section"]
                self.doc.add_heading('2.2.3 ä¸“ä¸šé£æ§è¯„ä¼°', level=3)
                
                risk_text = f"""
                æ•´ä½“é£é™©ç­‰çº§ï¼š{risk.get('overall_risk', 'æœªçŸ¥')}
                å»ºè®®ä»“ä½ä¸Šé™ï¼š{risk.get('position_limit', 'æœªçŸ¥')}
                å»ºè®®æ­¢æŸä½ï¼š{risk.get('stop_loss', 'æœªçŸ¥')}
                
                é£é™©ç­‰çº§ä¾æ®ï¼šåŸºäºäº¤æ˜“å‘˜å†³ç­–å’Œè¾©è®ºåˆ†æ­§åº¦è¿›è¡Œä¸“ä¸šè¯„ä¼°
                
                é£æ§ç»ç†ä¸“ä¸šæ„è§ï¼š
                {risk.get('manager_opinion', 'æš‚æ— æ„è§')}
                """
                
                self.doc.add_paragraph(risk_text)
            
            # CIOå†³ç­– - ç›´æ¥ä»æ–°æ•°æ®ç»“æ„è·å–
            executive_decision = result.get("decision_section", {}) if result.get("success") else {}
            decision = executive_decision if executive_decision else None
            
            if decision:
                self.doc.add_heading('2.2.4 CIOæœ€ç»ˆæƒå¨å†³ç­–', level=3)
                
                decision_text = f"""
                æœ€ç»ˆå†³ç­–ï¼š{decision.get('final_decision', 'æœªçŸ¥')}
                å†³ç­–ä¿¡å¿ƒï¼š{decision.get('confidence_level', decision.get('confidence', 'æœªçŸ¥'))}
                å»ºè®®ä»“ä½ï¼š{decision.get('position_size', 'æœªçŸ¥')}
                
                å†³ç­–ç†ç”±ï¼š
                {' | '.join(decision.get('key_rationale', decision.get('rationale', [])))}
                
                æ‰§è¡Œè®¡åˆ’ï¼š
                {decision.get('execution_plan', 'æš‚æ— è®¡åˆ’')}
                
                ç›‘æ§è¦ç‚¹ï¼š
                {' | '.join(decision.get('monitoring_points', []))}
                
                CIOæƒå¨å£°æ˜ï¼š
                {decision.get('cio_statement', 'æš‚æ— å£°æ˜')}
                """
                
                self.doc.add_paragraph(decision_text)
        else:
            error_text = f"å®Œæ•´å†³ç­–æµç¨‹åˆ†æå¤±è´¥ï¼š{debate_result.get('error', 'æœªçŸ¥é”™è¯¯')}"
            self.doc.add_paragraph(error_text)
        
        self.doc.add_page_break()
    
    def _add_complete_flow_analysis(self, commodity: str, result: Dict):
        """æ·»åŠ å®Œæ•´æµç¨‹åˆ†æ"""
        
        self.doc.add_heading('2.1 ä¼ ç»Ÿå®Œæ•´Trading Agentsæµç¨‹ç»“æœ', level=2)
        
        complete_result = result["result"]
        execution_summary = result.get("execution_summary", {})
        
        confidence = execution_summary.get('confidence_score', '0%')
        try:
            if isinstance(confidence, (int, float)):
                confidence_text = f"{confidence:.1%}"
            else:
                confidence_text = str(confidence)
        except (ValueError, TypeError):
            confidence_text = str(confidence)
            
        summary_text = f"""
        å·²å®Œæˆé˜¶æ®µï¼š{execution_summary.get('stages_completed', 0)}
        æœ€ç»ˆå†³ç­–ï¼š{execution_summary.get('final_decision', 'æœªçŸ¥')}
        ç½®ä¿¡åº¦ï¼š{confidence_text}
        æ€»è€—æ—¶ï¼š{execution_summary.get('total_duration', 0):.2f}ç§’
        """
        
        self.doc.add_paragraph(summary_text)
        
        self.doc.add_page_break()
    
    def _add_analyst_only_analysis(self, commodity: str, result: Dict):
        """æ·»åŠ ä»…åˆ†æå¸ˆåˆ†æ"""
        
        self.doc.add_heading('2.1 åˆ†æå¸ˆå›¢é˜Ÿä¸“ä¸šæŠ¥å‘Š', level=2)
        
        analysis_state = result["result"]
        
        # åˆ†æè¿›åº¦
        progress_text = f"åˆ†æå¸ˆå›¢é˜Ÿå·²å®Œæˆ{commodity}çš„6å¤§æ¨¡å—ä¸“ä¸šåˆ†æ"
        self.doc.add_paragraph(progress_text)
        
        self.doc.add_page_break()
    
    def _add_comprehensive_conclusion(self, analysis_results: Dict):
        """æ·»åŠ ç»¼åˆç»“è®º"""
        
        self.doc.add_heading('3. ç»¼åˆç»“è®ºä¸ä¸“ä¸šæŠ•èµ„å»ºè®®', level=1)
        
        # ç»Ÿè®¡åˆ†æç»“æœ
        total_count = len(analysis_results)
        success_count = sum(1 for result in analysis_results.values() 
                          if result.get("status") == "success")
        
        conclusion_text = f"""
        æœ¬æ¬¡ä¸“ä¸šåˆ†æå…±æ¶µç›–{total_count}ä¸ªæœŸè´§å“ç§ï¼ŒæˆåŠŸå®Œæˆ{success_count}ä¸ªå“ç§çš„æ·±åº¦åˆ†æã€‚
        
        ä¸“ä¸šç»“è®ºï¼š
        1. ç³»ç»Ÿè¿è¡Œç¨³å®šï¼Œ5é˜¶æ®µåˆ†ææµç¨‹å®Œæ•´æœ‰æ•ˆ
        2. å„åˆ†ææ¨¡å—åè°ƒå·¥ä½œï¼Œæ•°æ®è´¨é‡ä¼˜è‰¯
        3. è¾©è®ºé£æ§å†³ç­–æµç¨‹ä¸“ä¸šï¼Œæä¾›äº†å…·ä½“å¯æ“ä½œçš„æŠ•èµ„å»ºè®®
        4. äº¤æ˜“å‘˜ç¯èŠ‚æœ‰æ•ˆæ•´åˆäº†å¤šç©ºè§‚ç‚¹ï¼Œåˆ¶å®šäº†å®¢è§‚ç†æ€§çš„äº¤æ˜“ç­–ç•¥
        5. é£æ§è¯„ä¼°ç‹¬ç«‹å®¢è§‚ï¼ŒCIOå†³ç­–æƒå¨ä¸“ä¸š
        
        ä¸“ä¸šæŠ•èµ„å»ºè®®ï¼š
        1. ä¸¥æ ¼æŒ‰ç…§é£æ§å»ºè®®æ‰§è¡Œä»“ä½ç®¡ç†ï¼Œæ§åˆ¶æŠ•èµ„é£é™©
        2. å¯†åˆ‡å…³æ³¨å¸‚åœºå˜åŒ–ï¼Œæ ¹æ®äº¤æ˜“å‘˜åˆ¶å®šçš„å…·ä½“ç­–ç•¥æ‰§è¡Œ
        3. å®šæœŸå›é¡¾åˆ†æç»“æœï¼Œä¼˜åŒ–æŠ•èµ„å†³ç­–æµç¨‹
        4. é‡è§†äº¤æ˜“å‘˜æå‡ºçš„å¤‡é€‰æ–¹æ¡ˆï¼Œåšå¥½é£é™©é¢„æ¡ˆ
        5. éµå¾ªCIOçš„æƒå¨å†³ç­–ï¼Œç¡®ä¿æŠ•èµ„çºªå¾‹æ€§
        
        ä¸“ä¸šé£é™©æç¤ºï¼š
        1. æœŸè´§å¸‚åœºå­˜åœ¨é«˜åº¦ä¸ç¡®å®šæ€§ï¼ŒæŠ•èµ„éœ€è°¨æ…
        2. æœ¬åˆ†æç»“æœåŸºäºå†å²æ•°æ®å’Œå½“å‰ä¿¡æ¯ï¼Œå¸‚åœºæƒ…å†µå¯èƒ½å‘ç”Ÿå˜åŒ–
        3. è¯·ç»“åˆå®é™…èµ„é‡‘æƒ…å†µå’Œé£é™©æ‰¿å—èƒ½åŠ›åšå‡ºæŠ•èµ„å†³ç­–
        4. å»ºè®®é…åˆä¸“ä¸šçš„é£é™©ç®¡ç†ç³»ç»Ÿè¿›è¡ŒæŠ•èµ„
        5. å®šæœŸè¯„ä¼°æŠ•èµ„ç»„åˆï¼ŒåŠæ—¶è°ƒæ•´ç­–ç•¥
        
        ç³»ç»Ÿä¼˜åŠ¿ï¼š
        1. å®Œæ•´çš„5é˜¶æ®µä¸“ä¸šåˆ†ææµç¨‹
        2. å¤šç»´åº¦æ•°æ®åˆ†æå’ŒéªŒè¯
        3. å®¢è§‚ç†æ€§çš„äº¤æ˜“ç­–ç•¥åˆ¶å®š
        4. ç‹¬ç«‹ä¸“ä¸šçš„é£é™©è¯„ä¼°
        5. æƒå¨çš„æœ€ç»ˆæŠ•èµ„å†³ç­–
        """
        
        self.doc.add_paragraph(conclusion_text)
        
        # æ·»åŠ ç”Ÿæˆä¿¡æ¯
        self.doc.add_paragraph()
        footer_text = f"æŠ¥å‘Šç”±æœŸè´§Trading Agentsç³»ç»Ÿä¸“ä¸šç‰ˆè‡ªåŠ¨ç”Ÿæˆ\nç”Ÿæˆæ—¶é—´ï¼š{datetime.now().strftime('%Yå¹´%mæœˆ%dæ—¥ %H:%M:%S')}"
        footer_para = self.doc.add_paragraph(footer_text)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================================
# ä¸»ç•Œé¢
# ============================================================================

def main():
    """ä¸»å‡½æ•°"""
    
    # é¡µé¢æ ‡é¢˜
    st.markdown('<h1 class="main-header">å•†å“æœŸè´§ Trading Agentsç³»ç»Ÿ</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div style="text-align: center; margin-bottom: 2rem;">
        <p style="font-size: 1.2rem; color: #666;">
            åˆ†æå¸ˆå›¢é˜Ÿ + å¤šç©ºè¾©è®º + äº¤æ˜“å‘˜åˆ†æ+ é£æ§ç®¡ç† + CIOå†³ç­– 
        </p>
        <p style="color: #888;">
            å®Œæ•´çš„æœŸè´§æŠ•èµ„å†³ç­–æµç¨‹ï¼Œä»æ•°æ®åˆ†æåˆ°æœ€ç»ˆæ‹æ¿
        </p>
        <p style="font-size: 1rem; color: #999; margin-top: 1rem;">
            åˆ¶ä½œäººï¼š7haoge &nbsp;&nbsp;&nbsp;&nbsp; è”ç³»æ–¹å¼ï¼š953534947@qq.com
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # åˆå§‹åŒ–ç®¡ç†å™¨
    if "data_manager" not in st.session_state:
        st.session_state.data_manager = StreamlitDataManager()
    
    if "analysis_manager" not in st.session_state:
        st.session_state.analysis_manager = StreamlitAnalysisManager()
    
    if "word_generator" not in st.session_state and DOCX_AVAILABLE:
        st.session_state.word_generator = WordReportGenerator()
    
    # ä¾§è¾¹æ 
    with st.sidebar:
        
        # ç³»ç»ŸçŠ¶æ€
        st.subheader("ğŸ“Š ç³»ç»ŸçŠ¶æ€")
        
        # æ£€æŸ¥APIå¯†é’¥é…ç½®
        try:
            config = FuturesTradingAgentsConfig("æœŸè´§TradingAgentsç³»ç»Ÿ_é…ç½®æ–‡ä»¶.json")
            
            # DeepSeek APIçŠ¶æ€
            deepseek_api_key = config.get("api_settings", {}).get("deepseek", {}).get("api_key")
            if deepseek_api_key:
                st.success("âœ… DeepSeek APIå·²é…ç½®")
            else:
                st.error("âŒ æœªé…ç½®DeepSeek APIå¯†é’¥")
                st.info("è¯·åœ¨é…ç½®æ–‡ä»¶ä¸­è®¾ç½® api_settings.deepseek.api_key")
            
            # Serper APIçŠ¶æ€
            serper_api_key = config.get("api_settings", {}).get("serper", {}).get("api_key")
            if serper_api_key:
                st.success("âœ… Serper APIå·²é…ç½®")
            else:
                st.error("âŒ æœªé…ç½®Serper APIå¯†é’¥")
                st.info("è¯·åœ¨é…ç½®æ–‡ä»¶ä¸­è®¾ç½® api_settings.serper.api_key")
                
        except Exception as e:
            st.error(f"âŒ é…ç½®æ–‡ä»¶é”™è¯¯: {e}")
        
        # WordæŠ¥å‘ŠåŠŸèƒ½çŠ¶æ€
        if not DOCX_AVAILABLE:
            st.error("âŒ WordæŠ¥å‘ŠåŠŸèƒ½ä¸å¯ç”¨")
            st.info("è¯·å®‰è£…ï¼špip install python-docx")
        
        # ç³»ç»Ÿç‰¹è‰²
        st.subheader("ğŸŒŸ ç³»ç»Ÿç‰¹è‰²")
        st.info("""
        **6å¤§æ¨¡å—ç»¼åˆåˆ†æ**
        â€¢ åŸºäºçœŸå®æ•°æ®
        â€¢ AIé©±åŠ¨æ™ºèƒ½å†³ç­–
        """)
        
        # å¤‡æ³¨
        st.subheader("ğŸ“ å¤‡æ³¨")
        st.info("""
        **æœ¬ç³»ç»Ÿçš„åˆ†æåŠŸèƒ½ä¾èµ–äºDeepSeekï¼Œè”ç½‘æœç´¢åŠŸèƒ½ä¾èµ–äºSerperï¼Œç³»ç»Ÿæ­£åœ¨å®Œå–„ä¸­ï¼Œå¦‚æœ‰ä¸è¶³è¯·å¤šåŒ…æ¶µï¼æ¬¢è¿æä¾›æ”¹è¿›æ„è§ï¼**
        """)
        
        # æ„Ÿè°¢ä¿¡æ¯
        st.subheader("ğŸ™ æ„Ÿè°¢")
        st.info("""
        **Trading Agentsé¡¹ç›®æ–¹**
        
        **Akshareæ•°æ®æ¥å£**
        """)
    
    # ä¸»ç•Œé¢é€‰é¡¹å¡
    tab1, tab2, tab3, tab4 = st.tabs(["ğŸ“Š æ•°æ®ç®¡ç†", "ğŸ”„ æ•°æ®æ›´æ–°", "ğŸ”§ åˆ†æé…ç½®", "ğŸ­ åˆ†æç»“æœ"])
    
    # Tab 1: æ•°æ®ç®¡ç†
    with tab1:
        st.header("ğŸ“Š æ•°æ®ç®¡ç†")
        
        with st.spinner("æ­£åœ¨æ£€æŸ¥æ•°æ®çŠ¶æ€..."):
            data_status = st.session_state.data_manager.get_data_status()
        
        # æ•°æ®æ¦‚è§ˆ
        col1, col2 = st.columns(2)
        
        with col1:
            success_modules = sum(1 for m in data_status["modules"].values() if m["status"] == "success")
            st.metric("å¯ç”¨æ¨¡å—", f"{success_modules}/6", help="æ­£å¸¸å·¥ä½œçš„åˆ†ææ¨¡å—æ•°")
        
        with col2:
            total_records = sum(m.get("total_records", 0) for m in data_status["modules"].values())
            st.metric("æ€»è®°å½•æ•°", f"{total_records:,}", help="æ‰€æœ‰æ¨¡å—çš„æ•°æ®è®°å½•æ€»æ•°")
        
        # å„æ¨¡å—è¯¦ç»†çŠ¶æ€
        st.subheader("ğŸ“‹ å„æ¨¡å—æ•°æ®çŠ¶æ€")
        
        module_names = {
            "inventory": "åº“å­˜æ•°æ®",
            "positioning": "æŒä»“å¸­ä½", 
            "term_structure": "æœŸé™ç»“æ„",
            "technical_analysis": "æŠ€æœ¯é¢æŒ‡æ ‡",
            "basis": "åŸºå·®æ•°æ®",
            "receipt": "ä»“å•æ•°æ®"
        }
        
        for module_key, module_info in data_status["modules"].items():
            module_name = module_names.get(module_key, module_key)
            status_color = "success" if module_info["status"] == "success" else "error"
            
            with st.expander(f"**{module_name}** - {module_info['status']}", expanded=False):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("å“ç§æ•°", module_info.get("commodities_count", 0))
                with col2:
                    st.metric("è®°å½•æ•°", f"{module_info.get('total_records', 0):,}")
                with col3:
                    st.write(f"æœ€åæ›´æ–°: {module_info.get('last_update', 'æœªçŸ¥')}")
                
                if module_info["status"] == "error":
                    st.error(f"é”™è¯¯ä¿¡æ¯: {module_info.get('error', 'æœªçŸ¥é”™è¯¯')}")
                
                # æ˜¾ç¤ºæ•°æ®è·¯å¾„
                st.info(f"æ•°æ®è·¯å¾„: {module_info.get('path', 'æœªçŸ¥')}")
                
                # æ˜¾ç¤ºè¯¦ç»†æ•°æ®æƒ…å†µ
                if module_info["status"] == "success" and module_info.get("commodities_count", 0) > 0:
                    st.write("**ğŸ“Š æ•°æ®è¯¦æƒ…:**")
                    
                    # è·å–è¯¥æ¨¡å—çš„è¯¦ç»†æ•°æ®ä¿¡æ¯
                    detailed_info = st.session_state.data_manager._get_module_detailed_info(module_key)
                    
                    if detailed_info:
                        # æ˜¾ç¤ºå“ç§åˆ—è¡¨å’Œæ•°æ®èŒƒå›´ - æ˜¾ç¤ºæ‰€æœ‰å“ç§
                        st.write("**å“ç§æ•°æ®èŒƒå›´:**")
                        varieties_info = []
                        for variety, info in detailed_info.items():  # æ˜¾ç¤ºæ‰€æœ‰å“ç§
                            start_date = info.get('start_date', 'æœªçŸ¥')
                            end_date = info.get('end_date', 'æœªçŸ¥')
                            record_count = info.get('record_count', 0)
                            varieties_info.append({
                                'å“ç§': variety,
                                'èµ·å§‹æ—¶é—´': start_date,
                                'ç»“æŸæ—¶é—´': end_date,
                                'æ•°æ®é‡': record_count
                            })
                        
                        if varieties_info:
                            # æŒ‰å“ç§åç§°æ’åº
                            varieties_info.sort(key=lambda x: x['å“ç§'])
                            df_info = pd.DataFrame(varieties_info)
                            st.dataframe(df_info, use_container_width=True, height=min(400, len(varieties_info) * 35 + 100))
        
    
    # Tab 2: æ•°æ®æ›´æ–°
    with tab2:
        st.header("ğŸ”„ æ•°æ®æ›´æ–°")
        
        st.info("ğŸ’¡ é€‰æ‹©éœ€è¦æ›´æ–°çš„æ•°æ®æ¨¡å—ï¼Œç³»ç»Ÿå°†å¯åŠ¨å¯¹åº”çš„æ›´æ–°ç¨‹åº")
        
        # æ•°æ®æ›´æ–°é€‰é¡¹
        col1, col2 = st.columns(2)
        
        with col1:
            # åº“å­˜æ•°æ®æ›´æ–°
            if st.button("ğŸ“¦ æ›´æ–°åº“å­˜æ•°æ®", use_container_width=True):
                with st.spinner("å¯åŠ¨åº“å­˜æ•°æ®æ›´æ–°..."):
                    result = st.session_state.data_manager.run_data_update_direct("inventory")
                    if result["status"] == "success":
                        st.success(result["message"])
                        st.info(result["details"])
                    else:
                        st.error(result["message"])
            
            # æŒä»“æ•°æ®æ›´æ–°
            if st.button("ğŸ¯ æ›´æ–°æŒä»“å¸­ä½æ•°æ®", use_container_width=True):
                with st.spinner("å¯åŠ¨æŒä»“æ•°æ®æ›´æ–°..."):
                    result = st.session_state.data_manager.run_data_update_direct("positioning")
                    if result["status"] == "success":
                        st.success(result["message"])
                        st.info(result["details"])
                    else:
                        st.error(result["message"])
            
            # æœŸé™ç»“æ„æ›´æ–°
            if st.button("ğŸ“ˆ æ›´æ–°æœŸé™ç»“æ„æ•°æ®", use_container_width=True):
                with st.spinner("å¯åŠ¨æœŸé™ç»“æ„æ›´æ–°..."):
                    result = st.session_state.data_manager.run_data_update_direct("term_structure")
                    if result["status"] == "success":
                        st.success(result["message"])
                        st.info(result["details"])
                    else:
                        st.error(result["message"])
        
        with col2:
            # æŠ€æœ¯åˆ†ææ•°æ®æ›´æ–°
            if st.button("ğŸ“Š æ›´æ–°æŠ€æœ¯åˆ†ææ•°æ®", use_container_width=True):
                with st.spinner("å¯åŠ¨æŠ€æœ¯æ•°æ®æ›´æ–°..."):
                    result = st.session_state.data_manager.run_data_update_direct("technical_analysis")
                    if result["status"] == "success":
                        st.success(result["message"])
                        st.info(result["details"])
                    else:
                        st.error(result["message"])
            
            # åŸºå·®æ•°æ®æ›´æ–°
            if st.button("ğŸ’° æ›´æ–°åŸºå·®åˆ†ææ•°æ®", use_container_width=True):
                with st.spinner("å¯åŠ¨åŸºå·®æ•°æ®æ›´æ–°..."):
                    result = st.session_state.data_manager.run_data_update_direct("basis")
                    if result["status"] == "success":
                        st.success(result["message"])
                        st.info(result["details"])
                    else:
                        st.error(result["message"])
            
            # ä»“å•æ•°æ®æ›´æ–°
            if st.button("ğŸ“œ æ›´æ–°ä»“å•æ•°æ®", use_container_width=True):
                with st.spinner("å¯åŠ¨ä»“å•æ•°æ®æ›´æ–°..."):
                    result = st.session_state.data_manager.run_data_update_direct("receipt")
                    if result["status"] == "success":
                        st.success(result["message"])
                        st.info(result["details"])
                    else:
                        st.error(result["message"])
        
        st.markdown("---")
        st.warning("âš ï¸ æ•°æ®æ›´æ–°å°†åœ¨æ–°çš„å‘½ä»¤è¡Œçª—å£ä¸­è¿è¡Œï¼Œè¯·æŒ‰ç…§æç¤ºå®Œæˆæ“ä½œ")
    
    # Tab 3: åˆ†æé…ç½®
    with tab3:
        st.header("ğŸ”§ åˆ†æé…ç½®")
        
        # åˆ†æå‚æ•°é…ç½®
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("ğŸ“Š åŸºç¡€é…ç½®")
            
            # åˆ†ææ¨¡å—é€‰æ‹© - å…ˆå®šä¹‰ï¼Œåé¢ä¼šç”¨åˆ°
            st.subheader("ğŸ”§ åˆ†ææ¨¡å—")
            analysis_modules = st.multiselect(
                "é€‰æ‹©åˆ†ææ¨¡å—",
                options=["inventory", "positioning", "term_structure", "technical", "basis", "news"],
                default=["inventory", "positioning", "term_structure", "technical", "basis", "news"],
                format_func=lambda x: {
                    "inventory": "åº“å­˜ä»“å•åˆ†æ",
                    "positioning": "æŒä»“å¸­ä½åˆ†æ", 
                    "term_structure": "æœŸé™ç»“æ„åˆ†æ",
                    "technical": "æŠ€æœ¯é¢åˆ†æ",
                    "basis": "åŸºå·®åˆ†æ",
                    "news": "æ–°é—»åˆ†æ"
                }.get(x, x),
                help="é€‰æ‹©è¦æ‰§è¡Œçš„åˆ†ææ¨¡å—"
            )
            
            # æ ¹æ®é€‰æ‹©çš„æ¨¡å—æ˜¾ç¤ºæ”¯æŒçš„å“ç§
            module_name_mapping = {
                "inventory": "åº“å­˜ä»“å•åˆ†æ",
                "positioning": "æŒä»“å¸­ä½åˆ†æ", 
                "term_structure": "æœŸé™ç»“æ„åˆ†æ",
                "technical": "æŠ€æœ¯åˆ†æ",
                "basis": "åŸºå·®åˆ†æ",
                "news": "æ–°é—»åˆ†æ"
            }
            
            if analysis_modules:
                module_varieties_info = st.session_state.data_manager.get_module_supported_varieties()
                supported_varieties = set()
                for module in analysis_modules:
                    module_name = module_name_mapping.get(module, module)
                    module_varieties = module_varieties_info.get(module_name, set())
                    supported_varieties.update(module_varieties)
                
            else:
                supported_varieties = set()
            
            # å“ç§é€‰æ‹© - æ”¹ä¸ºæ–‡æœ¬è¾“å…¥æ–¹å¼
            with st.spinner("è·å–å¯ç”¨å“ç§..."):
                data_status = st.session_state.data_manager.get_data_status()
                available_commodities = sorted(list(data_status["summary"]["common_commodities"]))
            
            # æ–‡æœ¬è¾“å…¥å“ç§ä»£ç 
            commodity_input = st.text_input(
                "è¾“å…¥åˆ†æå“ç§",
                value="AU",
                placeholder="å¦‚ï¼šAU",
                help="è¾“å…¥æœŸè´§å“ç§ä»£ç ï¼Œå¤šä¸ªå“ç§ç”¨ç©ºæ ¼åˆ†éš”"
            )
            
            # è§£æè¾“å…¥çš„å“ç§ä»£ç å¹¶æ˜¾ç¤ºæœ¬åœ°æ•°æ®æƒ…å†µ
            if commodity_input.strip():
                # åˆ†å‰²å¹¶æ¸…ç†å“ç§ä»£ç ï¼ˆæ”¯æŒç©ºæ ¼ã€é€—å·åˆ†éš”ï¼‰
                input_commodities = []
                for separator in [' ', ',']:
                    if separator in commodity_input:
                        input_commodities = [code.strip().upper() for code in commodity_input.split(separator) if code.strip()]
                        break
                else:
                    # å¦‚æœæ²¡æœ‰åˆ†éš”ç¬¦ï¼Œå½“ä½œå•ä¸ªå“ç§
                    input_commodities = [commodity_input.strip().upper()]
                
                selected_commodities = input_commodities
                
                # æ˜¾ç¤ºæ¯ä¸ªå“ç§çš„æœ¬åœ°æ•°æ®æƒ…å†µ
                for commodity in input_commodities:
                    st.write(f"**{commodity} æœ¬åœ°æ•°æ®æƒ…å†µï¼š**")
                    
                    # æ£€æŸ¥å„ä¸ªåˆ†æå¸ˆæ¨¡å—çš„æœ¬åœ°æ•°æ®æƒ…å†µï¼ˆé™¤äº†æ–°é—»åˆ†æï¼‰
                    module_data_status = check_commodity_local_data(commodity, data_status)
                    
                    # åˆ›å»ºåˆ—æ¥æ˜¾ç¤ºæ•°æ®çŠ¶æ€
                    cols = st.columns(6)  # 6ä¸ªæ¨¡å—ï¼ˆé™¤äº†æ–°é—»åˆ†æï¼‰
                    
                    module_display_names = {
                        "technical": "æŠ€æœ¯åˆ†æ",
                        "basis": "åŸºå·®åˆ†æ", 
                        "inventory": "åº“å­˜åˆ†æ",
                        "positioning": "æŒä»“åˆ†æ",
                        "term_structure": "æœŸé™ç»“æ„",
                        "receipt": "ä»“å•åˆ†æ"
                    }
                    
                    for i, (module_key, display_name) in enumerate(module_display_names.items()):
                        with cols[i]:
                            has_data = module_data_status.get(module_key, False)
                            if has_data:
                                st.success(f"âœ… {display_name}\næœ¬åœ°æ•°æ®âˆš")
                            else:
                                st.error(f"âŒ {display_name}\næœ¬åœ°æ•°æ®âœ—")
                    
                    # æ·»åŠ è”ç½‘æœç´¢å¤‡æ³¨
                    st.caption("*è‹¥æ— æœ¬åœ°æ•°æ®ï¼Œåˆ™é€šè¿‡è”ç½‘æœç´¢æ•°æ®")
                    
                    st.write("---")  # åˆ†éš”çº¿
            else:
                selected_commodities = []
                st.info("ğŸ’¡ è¯·è¾“å…¥è¦åˆ†æçš„æœŸè´§å“ç§ä»£ç ")
            
            # åˆ†ææ—¥æœŸ
            analysis_date = st.date_input(
                "åˆ†ææ—¥æœŸ",
                value=datetime.now().date() - timedelta(days=1),
                help="é€‰æ‹©åˆ†æçš„ç›®æ ‡æ—¥æœŸ"
            )
            
        
        with col2:
            st.subheader("âš™ï¸ é«˜çº§é…ç½®")
            
            # åˆ†ææ¨¡å¼é€‰æ‹© - ç®€åŒ–ä¸ºä¸¤ç§æ¨¡å¼
            analysis_mode = st.radio(
                "é€‰æ‹©åˆ†ææ¨¡å¼",
                options=["analyst_only", "complete_flow"],
                index=1,
                format_func=lambda x: {
                    "analyst_only": "ä»…åˆ†æå¸ˆå›¢é˜Ÿ",
                    "complete_flow": "å®Œæ•´æµç¨‹åˆ†æï¼ˆåˆ†æå¸ˆ+è¾©è®º+äº¤æ˜“å‘˜+é£æ§+å†³ç­–ï¼‰"
                }.get(x, x),
                help="é€‰æ‹©ä¸åŒçš„åˆ†ææ¨¡å¼"
            )
            
            # è°ƒè¯•æ¨¡å¼å¼€å…³
            debug_mode = st.checkbox("ğŸ”§ è°ƒè¯•æ¨¡å¼", help="æ˜¾ç¤ºè¯¦ç»†çš„çŠ¶æ€è°ƒè¯•ä¿¡æ¯")
            if debug_mode:
                st.session_state.debug_mode = True
            else:
                st.session_state.debug_mode = False
            
            # æ ¹æ®æ¨¡å¼æ˜¾ç¤ºä¸åŒé€‰é¡¹
            if analysis_mode == "complete_flow":
                # æ˜¾ç¤ºè¯¦ç»†çš„æµç¨‹è¯´æ˜
                with st.expander("ğŸ­ å®Œæ•´5é˜¶æ®µæµç¨‹è¯¦è§£", expanded=False):
                    st.markdown("""
### ğŸ“Š ç³»ç»Ÿæ¶æ„ä¸åˆ†ææµç¨‹

æœ¬ç³»ç»Ÿé‡‡ç”¨**5é˜¶æ®µæ¸è¿›å¼å†³ç­–æ¶æ„**ï¼Œé€šè¿‡å¤šå±‚æ¬¡ã€å¤šè§’åº¦çš„ä¸“ä¸šåˆ†æï¼Œç¡®ä¿æŠ•èµ„å†³ç­–çš„ç§‘å­¦æ€§å’Œå¯è¿½æº¯æ€§ã€‚

---

#### é˜¶æ®µ1ï¸âƒ£ï¼šå…­å¤§åˆ†æå¸ˆå›¢é˜Ÿ ğŸ“Š

**èŒè´£**ï¼šåŸºäºçœŸå®å¸‚åœºæ•°æ®è¿›è¡Œä¸“ä¸šåˆ†æ

- **ğŸ“¦ åº“å­˜ä»“å•åˆ†æå¸ˆ**
  - æ•°æ®æ¥æºï¼šæ—¥åº¦åº“å­˜æ•°æ®ã€ä»“å•ç»Ÿè®¡
  - åˆ†æè¦ç‚¹ï¼šä¾›éœ€å¹³è¡¡ã€åº“å­˜å‘¨æœŸã€åº“å­˜ä¸ä»·æ ¼ç›¸å…³æ€§
  - è¾“å‡ºï¼šåº“å­˜å‹åŠ›è¯„ä¼°ã€ä¸»åŠ¨/è¢«åŠ¨è¡¥åº“åˆ¤æ–­ã€çœ‹å¤š/çœ‹ç©ºè§‚ç‚¹

- **ğŸ“ˆ æŠ€æœ¯é¢åˆ†æå¸ˆ**
  - æ•°æ®æ¥æºï¼šä»·æ ¼Kçº¿ã€æŠ€æœ¯æŒ‡æ ‡ï¼ˆMACDã€RSIã€å¸ƒæ—å¸¦ç­‰ï¼‰
  - åˆ†æè¦ç‚¹ï¼šè¶‹åŠ¿åˆ¤æ–­ã€æ”¯æ’‘é˜»åŠ›ä½ã€è¶…ä¹°è¶…å–ä¿¡å·
  - è¾“å‡ºï¼šå…³é”®ä»·ä½ã€äº¤æ˜“ä¿¡å·ã€æŠ€æœ¯é¢å¤šç©ºå€¾å‘

- **ğŸ¯ æŒä»“å¸­ä½åˆ†æå¸ˆ**
  - æ•°æ®æ¥æºï¼šå‰20åæŒä»“å¸­ä½æ•°æ®ã€ä¸»åŠ›èµ„é‡‘åŠ¨å‘
  - åˆ†æè¦ç‚¹ï¼šèµ„é‡‘æµå‘ã€ä¸»åŠ›æ„å›¾ã€æŒä»“é›†ä¸­åº¦
  - è¾“å‡ºï¼šæœºæ„åšå¤š/åšç©ºä¿¡å·ã€èµ„é‡‘åŠ¨å‘è¯„ä¼°

- **ğŸ’° åŸºå·®åˆ†æå¸ˆ**
  - æ•°æ®æ¥æºï¼šæœŸè´§ä»·æ ¼vsç°è´§ä»·æ ¼ã€åŸºå·®ç‡
  - åˆ†æè¦ç‚¹ï¼šæœŸç°å…³ç³»ã€å¥—åˆ©æœºä¼šã€å¸‚åœºç»“æ„
  - è¾“å‡ºï¼šåŸºå·®æ”¶æ•›/æ‰©å¤§è¶‹åŠ¿ã€æ­£å‘/è´Ÿå‘å¸‚åœºåˆ¤æ–­

- **ğŸ“Š æœŸé™ç»“æ„åˆ†æå¸ˆ**
  - æ•°æ®æ¥æºï¼šè¿‘æœˆ/è¿œæœˆåˆçº¦ä»·å·®ã€Contango/Backwardation
  - åˆ†æè¦ç‚¹ï¼šè¿œæœŸæ›²çº¿å½¢æ€ã€æ—¶é—´ä»·å€¼ã€å¸‚åœºé¢„æœŸ
  - è¾“å‡ºï¼šå¸‚åœºé¢„æœŸè¯„ä¼°ã€æœŸé™ç»“æ„é£é™©æç¤º

- **ğŸ“° æ–°é—»åˆ†æå¸ˆ**
  - æ•°æ®æ¥æºï¼šå®è§‚æ”¿ç­–ã€è¡Œä¸šæ–°é—»ã€å¸‚åœºæƒ…ç»ª
  - åˆ†æè¦ç‚¹ï¼šæ”¿ç­–å½±å“ã€çªå‘äº‹ä»¶ã€å¸‚åœºé¢„æœŸå˜åŒ–
  - è¾“å‡ºï¼šå®è§‚é¢åˆ©å¥½/åˆ©ç©ºå› ç´ ã€å¸‚åœºæƒ…ç»ªè¯„ä¼°

**é˜¶æ®µäº§å‡º**ï¼š6ä»½ç‹¬ç«‹çš„ä¸“ä¸šåˆ†ææŠ¥å‘Šï¼Œæ¯ä»½åŒ…å«æ•°æ®ã€è§‚ç‚¹ã€ä¿¡å¿ƒåº¦

---

#### é˜¶æ®µ2ï¸âƒ£ï¼šæ¿€æƒ…å¤šç©ºè¾©è®º ğŸ­

**èŒè´£**ï¼šé€šè¿‡å¯¹æŠ—æ€§è¾©è®ºæš´éœ²åˆ†æç›²ç‚¹ï¼Œæå‡å†³ç­–è´¨é‡

- **ğŸ‚ å¤šå¤´åˆ†æå¸ˆ**
  - åŸºäº6å¤§æ¨¡å—æ•°æ®ï¼Œè®ºè¯åšå¤šæœºä¼š
  - åå‡»ç©ºå¤´è§‚ç‚¹ï¼ŒæŒ–æ˜å¤šå¤´é€»è¾‘
  - å¼ºè°ƒä¸Šæ¶¨åŠ¨èƒ½ã€æ”¯æ’‘å› ç´ 

- **ğŸ» ç©ºå¤´åˆ†æå¸ˆ**
  - åŸºäº6å¤§æ¨¡å—æ•°æ®ï¼Œè¯†åˆ«åšç©ºé£é™©
  - ç—›å‡»å¤šå¤´é€»è¾‘ï¼Œæš´éœ²é£é™©éšæ‚£
  - å¼ºè°ƒå›è°ƒå‹åŠ›ã€é˜»åŠ›å› ç´ 

- **âš–ï¸ AIè¯„åˆ¤ç³»ç»Ÿ**
  - ä¸‰ç»´åº¦è¯„åˆ†ï¼šè®ºæ®è´¨é‡(40%) + é€»è¾‘ä¸¥å¯†æ€§(35%) + å®¢è§‚ç†æ€§åº¦(25%)
  - è®ºæ®å¯éªŒè¯æ€§æ£€æŸ¥ï¼šæ‰€æœ‰æ•°æ®å¿…é¡»å¯è¿½æº¯
  - åŠ åˆ†æœºåˆ¶ï¼šæ–¹å‘ä¸åˆ†æå¸ˆè§‚ç‚¹ä¸€è‡´ â†’ åŠ åˆ†

**è¾©è®ºè§„åˆ™**ï¼š
- æ¯è½®è¾©è®ºåŒ…å«ï¼šå¤šå¤´è®ºè¿° â†’ ç©ºå¤´åé©³ â†’ å¤šå¤´åå‡» â†’ ç©ºå¤´æ€»ç»“
- è¯„åˆ¤åŸºäºè®ºè¯è´¨é‡ï¼Œä¸é¢„è®¾ç«‹åœº
- è¾©è®ºè¯„åˆ¤å½±å“äº¤æ˜“å‘˜çš„ä¿¡å¿ƒåº¦è°ƒæ•´

**é˜¶æ®µäº§å‡º**ï¼šå¤šç©ºè¯„åˆ†ã€è¾©è®ºèƒœè´Ÿã€å…³é”®è®ºæ®æ€»ç»“

---

#### é˜¶æ®µ3ï¸âƒ£ï¼šAIé©±åŠ¨ä¸“ä¸šäº¤æ˜“å‘˜ ğŸ’¼

**èŒè´£**ï¼šç»¼åˆå‰ä¸¤é˜¶æ®µä¿¡æ¯ï¼Œåˆ¶å®šå…·ä½“äº¤æ˜“ç­–ç•¥

**æ ¸å¿ƒå†³ç­–é€»è¾‘**ï¼š

1. **ã€è¾©è®ºè¯„åˆ¤åˆ†æã€‘**
   - ç»Ÿè®¡å„æ¨¡å—è¾©è®ºç»“æœï¼šå¼ºåŠ¿/åå¼º/ä¸­æ€§
   - è¯†åˆ«å¤šç©ºä¼˜åŠ¿æ¨¡å—å’Œå…³é”®è®ºæ®
   - ç”Ÿæˆè¾©è®ºè´¨é‡è¯„ä¼°æŠ¥å‘Š

2. **ã€åˆ†æå¸ˆè§‚ç‚¹ç§‘å­¦è¯„ä¼°ã€‘**
   - ç»Ÿè®¡çœ‹å¤š/çœ‹ç©º/ä¸­æ€§æ¨¡å—æ•°é‡
   - è¯„ä¼°æ•°æ®æ”¯æ’‘ç¨‹åº¦ï¼ˆå……åˆ†/æœ‰é™ï¼‰
   - è®¡ç®—å¯ä¿¡åº¦ç­‰çº§ï¼ˆé«˜/ä¸­/ä½ï¼‰

3. **ã€å¤šç©ºç»¼åˆåˆ¤æ–­ç»“è®ºã€‘**
   - **æ–¹å‘åˆ¤æ–­é“å¾‹**ï¼šåˆ†æå¸ˆè§‚ç‚¹å¤šæ•°å†³ï¼ˆçœ‹å¤šæ•°é‡ vs çœ‹ç©ºæ•°é‡ï¼‰
   - **ä¿¡å¿ƒåº¦è®¡ç®—**ï¼šåŸºç¡€ä¿¡å¿ƒåº¦ï¼ˆä¼˜åŠ¿æ¨¡å—æ•°ï¼‰ + 3ç±»è°ƒæ•´å› ç´ 
     - è°ƒæ•´1ï¼šè¾©è®ºè¯„åˆ¤å€¾å‘
     - è°ƒæ•´2ï¼šåå‘é«˜ä¿¡åº¦æ¨¡å—
     - è°ƒæ•´3ï¼šä¸­æ€§æ¨¡å—è¾©è®ºå€¾å‘
   - è¾“å‡ºï¼šæ–¹å‘ï¼ˆçœ‹å¤š/çœ‹ç©º/ä¸­æ€§ï¼‰+ ä¿¡å¿ƒåº¦ï¼ˆé«˜/ä¸­/ä½ï¼‰

4. **ã€äº¤æ˜“ç­–ç•¥å»ºè®®ã€‘**
   - **ä»“ä½è®¡ç®—**ï¼šæ ¹æ®ä¿¡å¿ƒåº¦ç¡®å®šåŒºé—´ï¼ˆé«˜10-15%ã€ä¸­5-10%ã€ä½2-5%ï¼‰
   - **æ­¢æŸä½è®¾ç½®**ï¼šåŸºäºå…³é”®æŠ€æœ¯ä½ï¼ˆæ”¯æ’‘/é˜»åŠ›ï¼‰
   - **æŒä»“å‘¨æœŸ**ï¼šçŸ­çº¿/ä¸­çº¿/é•¿çº¿
   - **é£æ§æªæ–½**ï¼šä¸¥æ ¼æ­¢æŸã€åˆ†æ‰¹å»ºä»“ã€åŠ¨æ€è°ƒæ•´

**é€æ˜åŒ–åŸåˆ™**ï¼šæ‰€æœ‰è®¡ç®—è¿‡ç¨‹å®Œå…¨é€æ˜ï¼Œå†³ç­–ä¾æ®å¯è¿½æº¯

**é˜¶æ®µäº§å‡º**ï¼šäº¤æ˜“æ–¹å‘ã€å»ºè®®ä»“ä½ã€æ­¢æŸä½ã€é£é™©æç¤º

---

#### é˜¶æ®µ4ï¸âƒ£ï¼šä¸“ä¸šé£æ§ç®¡ç† ğŸ›¡ï¸

**èŒè´£**ï¼šç‹¬ç«‹é£é™©è¯„ä¼°ï¼Œä¸€ç¥¨å¦å†³æƒ

**é£æ§è¯„ä¼°æµç¨‹**ï¼š

1. **é£é™©çŸ©é˜µè¯„ä¼°**
   - æ½œåœ¨æœ€å¤§äºæŸï¼ˆåŸºäºæ­¢æŸä½è®¡ç®—ï¼‰
   - é£é™©å‘ç”Ÿæ¦‚ç‡ï¼ˆåŸºäºè¾©è®ºåˆ†æ­§åº¦ï¼‰
   - é£é™©ç­‰çº§åˆ¤å®šï¼šæé«˜/é«˜/ä¸­/ä½

2. **æ“ä½œä¿¡å¿ƒåº¦ç‹¬ç«‹åˆ¤æ–­**
   - åŸºäºé£é™©ç­‰çº§ç¡®å®šæ“ä½œä¿¡å¿ƒåº¦
   - é«˜é£é™© â†’ ä½ä¿¡å¿ƒ â†’ ä¿å®ˆä»“ä½
   - ä½é£é™© â†’ é«˜ä¿¡å¿ƒ â†’ é€‚åº¦é…ç½®

3. **é£æ§æ‰¹å‡†æ„è§**
   - ğŸŸ¢ æ‰¹å‡†ï¼šé£é™©å¯æ§ï¼Œæ­£å¸¸æ‰§è¡Œ
   - ğŸŸ¡ æœ‰æ¡ä»¶æ‰¹å‡†ï¼šé«˜é£é™©ï¼Œå¿…é¡»è½»ä»“
   - ğŸ”´ æ‹’ç»ï¼šé£é™©å¤±æ§ï¼Œç¦æ­¢æ“ä½œ

4. **ä»“ä½ä¸Šé™å¼ºåˆ¶çº¦æŸ**
   - æé«˜é£é™©ï¼šâ‰¤1%
   - é«˜é£é™©ï¼š2-5%
   - ä¸­é£é™©ï¼š5-10%
   - ä½é£é™©ï¼š10-15%

**é£æ§çº¢çº¿**ï¼š
- æ­¢æŸå¿…é¡»åŸºäºæŠ€æœ¯ä½ï¼ˆé€šå¸¸2.5-3.5%ï¼‰
- å•ç¬”äºæŸæç«¯æƒ…å†µä¸è¶…è¿‡5%
- é£é™©å¤±æ§ç«‹å³å‡ä»“

**é˜¶æ®µäº§å‡º**ï¼šé£é™©ç­‰çº§ã€æ“ä½œä¿¡å¿ƒåº¦ã€ä»“ä½ä¸Šé™ã€æ‰¹å‡†æ„è§

---

#### é˜¶æ®µ5ï¸âƒ£ï¼šCIOç»¼åˆå†³ç­– ğŸ‘”

**èŒè´£**ï¼šæœ€ç»ˆæ‹æ¿ï¼Œæƒå¨å†³ç­–

**å†³ç­–æ¡†æ¶**ï¼š

1. **æ–¹å‘åˆ¤æ–­ç¡®è®¤**
   - å¤æ ¸äº¤æ˜“å‘˜çš„æ–¹å‘åˆ¤æ–­
   - åŸºäºåˆ†æå¸ˆè§‚ç‚¹ï¼ˆé“å¾‹ï¼šå¤šæ•°å†³ï¼‰
   - æœ€ç»ˆæ–¹å‘ï¼šçœ‹å¤š/çœ‹ç©º/è§‚æœ›

2. **åŒä¿¡å¿ƒåº¦ä½“ç³»**
   - **æ–¹å‘ä¿¡å¿ƒåº¦**ï¼šå¯¹æ–¹å‘åˆ¤æ–­çš„ä¿¡å¿ƒï¼ˆåŸºäºåˆ†æå¸ˆå…±è¯†ï¼‰
   - **æ“ä½œä¿¡å¿ƒåº¦**ï¼šå¯¹å®‰å…¨æ‰§è¡Œçš„ä¿¡å¿ƒï¼ˆåŸºäºé£æ§è¯„ä¼°ï¼‰
   - ä¸¤è€…ä¸åŒç»´åº¦ï¼Œå¯ä»¥ä¸ä¸€è‡´ï¼ˆä¾‹å¦‚ï¼šæ–¹å‘ä¿¡å¿ƒä¸­ + æ“ä½œä¿¡å¿ƒä½ï¼‰

3. **ä»“ä½æœ€ç»ˆç¡®å®š**
   - äº¤æ˜“å‘˜å»ºè®® vs é£æ§ä¸Šé™ â†’ å–æ›´ä¿å®ˆå€¼
   - é£æ§æ‹¥æœ‰ä¸€ç¥¨å¦å†³æƒ
   - æœ€ç»ˆä»“ä½å¿…é¡»åœ¨é£æ§æ¡†æ¶å†…

4. **æ‰§è¡ŒæŒ‡å¯¼åˆ¶å®š**
   - ä»“ä½é…ç½®ï¼šåŒºé—´ã€åˆ†æ‰¹ç­–ç•¥
   - é£é™©æ§åˆ¶ï¼šæ­¢æŸçºªå¾‹ã€åŠ¨æ€è°ƒæ•´
   - ç›‘æ§è¦ç‚¹ï¼šå…³é”®ä»·ä½ã€é£é™©ä¿¡å·

**CIOæ ¸å¿ƒåˆ¤æ–­**ï¼š
- åˆ¤æ–­1ï¼šå¸‚åœºæ–¹å‘ä¸è¶‹åŠ¿åŸºç¡€
- åˆ¤æ–­2ï¼šé£é™©ç®¡æ§çš„ä¼˜å…ˆçº§
- åˆ¤æ–­3ï¼šæ“ä½œç­–ç•¥çš„ä¿å®ˆç¨‹åº¦

**é˜¶æ®µäº§å‡º**ï¼šæœ€ç»ˆå†³ç­–ã€æ“ä½œæŒ‡å¯¼ã€ç›‘æ§è¦ç‚¹ã€å†³ç­–è¿½æº¯

---

### ğŸ¯ ç³»ç»Ÿç‰¹è‰²

âœ… **æ•°æ®é©±åŠ¨**ï¼šæ‰€æœ‰åˆ¤æ–­åŸºäºçœŸå®å¸‚åœºæ•°æ®  
âœ… **å¤šå±‚éªŒè¯**ï¼š6å¤§åˆ†æå¸ˆ + è¾©è®º + äº¤æ˜“å‘˜ + é£æ§ + CIO  
âœ… **é€æ˜å¯è¿½æº¯**ï¼šæ¯ä¸ªå†³ç­–ç¯èŠ‚å®Œå…¨é€æ˜ï¼Œè®¡ç®—è¿‡ç¨‹å¯æŸ¥  
âœ… **é£æ§ä¼˜å…ˆ**ï¼šé£æ§æ‹¥æœ‰ä¸€ç¥¨å¦å†³æƒï¼Œä¿æŠ¤æœ¬é‡‘ç¬¬ä¸€  
âœ… **è§’è‰²åˆ†å·¥**ï¼šå„å¸å…¶èŒï¼Œåˆ¶è¡¡æœºåˆ¶ï¼Œé¿å…å•ä¸€è§†è§’ç›²ç‚¹  
                    """)
                
                debate_rounds = st.slider(
                    "è¾©è®ºè½®æ•°",
                    min_value=1,
                    max_value=5,
                    value=3,
                    help="è®¾ç½®å¤šç©ºåŒæ–¹çš„è¾©è®ºè½®æ•°"
                )
                
                st.warning("âš ï¸ å®Œæ•´æµç¨‹éœ€è¦æ¶ˆè€—æ›´å¤šAIèµ„æºï¼Œé¢„è®¡è€—æ—¶3-5åˆ†é’Ÿ")
            
            else:
                st.info("ğŸ“Š ä»…æ‰§è¡Œ6å¤§åˆ†ææ¨¡å—ï¼Œä¸åŒ…å«è¾©è®ºå’Œå†³ç­–")
                debate_rounds = 0  # ä¸éœ€è¦è¾©è®º
            
            # AIæ¨¡å‹é€‰æ‹©
            ai_model = st.selectbox(
                "AIåˆ†ææ¨¡å‹",
                options=["deepseek-reasoner", "deepseek-chat"],
                index=0,
                help="é€‰æ‹©ç”¨äºåˆ†æçš„AIæ¨¡å‹"
            )
            
            # ç§»é™¤åˆ†ææ·±åº¦é€‰é¡¹ - AIæ¨¡å‹é€‰æ‹©å·²è¶³å¤ŸåŒºåˆ†åˆ†æè´¨é‡
            
            # å®æ—¶æ•°æ®
            use_realtime = st.checkbox(
                "å¯ç”¨å®æ—¶æ•°æ®å¢å¼º",
                value=True,
                help="ä½¿ç”¨å®æ—¶æœç´¢æ•°æ®å¢å¼ºåˆ†æç»“æœ"
            )
        
        # ä¿å­˜é…ç½®åˆ°session state
        if st.button("ğŸ’¾ ä¿å­˜é…ç½®", type="primary"):
            st.session_state.analysis_config = {
                "commodities": selected_commodities,
                "analysis_date": analysis_date.strftime("%Y-%m-%d"),
                "modules": analysis_modules,
                "analysis_mode": analysis_mode,
                "ai_model": ai_model,
                "use_realtime": use_realtime,
                "debate_rounds": debate_rounds
            }
            st.success("âœ… é…ç½®å·²ä¿å­˜")
        
        # å¼€å§‹åˆ†ææŒ‰é’®
        if st.button("ğŸš€ å¼€å§‹åˆ†æ", type="primary", disabled=not selected_commodities):
            if selected_commodities and analysis_modules:
                # å¯åŠ¨åˆ†æ
                with st.spinner("æ­£åœ¨å¯åŠ¨åˆ†æ..."):
                    st.session_state.analysis_manager.start_analysis(
                        commodities=selected_commodities,
                        analysis_date=analysis_date.strftime("%Y-%m-%d"),
                        modules=analysis_modules,
                        analysis_mode=analysis_mode,
                        config={
                            "ai_model": ai_model,
                            "use_realtime": use_realtime,
                            "debate_rounds": debate_rounds
                        }
                    )
                st.success("âœ… åˆ†æå·²å¯åŠ¨ï¼Œè¯·åˆ‡æ¢åˆ°'åˆ†æç»“æœ'æ ‡ç­¾é¡µæŸ¥çœ‹è¿›åº¦")
                st.rerun()
            else:
                st.error("âŒ è¯·è‡³å°‘é€‰æ‹©ä¸€ä¸ªå“ç§å’Œä¸€ä¸ªåˆ†ææ¨¡å—")
    
    # Tab 4: åˆ†æç»“æœ
    with tab4:
        st.header("ğŸ­ åˆ†æç»“æœ")
        
        # æ£€æŸ¥æ˜¯å¦æœ‰æ­£åœ¨è¿›è¡Œçš„åˆ†æ
        if hasattr(st.session_state.analysis_manager, 'current_analysis') and st.session_state.analysis_manager.current_analysis:
            
            # å¦‚æœåˆ†æåˆšå¯åŠ¨ï¼Œå¼€å§‹å¤„ç†
            if (st.session_state.analysis_manager.current_analysis.get("status") == "running" and 
                not st.session_state.analysis_manager.analysis_results):
                
                st.info("ğŸš€ å¼€å§‹æ‰§è¡Œåˆ†æ...")
                
                # å¤„ç†æ‰€æœ‰å“ç§
                st.session_state.analysis_manager.process_all_commodities()
            
            # æ˜¾ç¤ºåˆ†æè¿›åº¦
            st.session_state.analysis_manager.display_progress()
            
            # æ˜¾ç¤ºåˆ†æç»“æœ
            st.session_state.analysis_manager.display_results()
            
            # WordæŠ¥å‘Šç”ŸæˆæŒ‰é’®
            if (st.session_state.analysis_manager.analysis_results and 
                st.session_state.analysis_manager.current_analysis.get("status") == "completed" and
                DOCX_AVAILABLE):
                
                st.markdown("---")
                
                # ğŸ”¥ WordæŠ¥å‘Šç”Ÿæˆé€‰é¡¹
                st.subheader("ğŸ“„ ç”ŸæˆWordæŠ¥å‘Š")
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    report_mode = st.radio(
                        "æŠ¥å‘Šæ¨¡å¼",
                        options=["å¿«é€Ÿæ¨¡å¼ï¼ˆä»…æ–‡å­—ï¼‰", "å®Œæ•´æ¨¡å¼ï¼ˆåŒ…å«å›¾è¡¨ï¼‰"],
                        help="å¿«é€Ÿæ¨¡å¼ï¼šä»…åŒ…å«æ–‡å­—åˆ†æï¼Œç”Ÿæˆé€Ÿåº¦å¿«ï¼ˆçº¦10ç§’ï¼‰\nå®Œæ•´æ¨¡å¼ï¼šåŒ…å«æ‰€æœ‰å›¾è¡¨ï¼Œéœ€è¦æ›´é•¿æ—¶é—´ï¼ˆçº¦2-5åˆ†é’Ÿï¼‰",
                        horizontal=True
                    )
                
                with col2:
                    include_charts = (report_mode == "å®Œæ•´æ¨¡å¼ï¼ˆåŒ…å«å›¾è¡¨ï¼‰")
                
                if st.button("ğŸ“„ ç”ŸæˆWordæŠ¥å‘Š", type="primary"):
                    # åˆ›å»ºè¿›åº¦æ˜¾ç¤ºå®¹å™¨
                    progress_placeholder = st.empty()
                    status_placeholder = st.empty()
                    
                    def update_progress(msg: str):
                        """æ›´æ–°è¿›åº¦æ˜¾ç¤º"""
                        status_placeholder.info(f"ğŸ“ {msg}")
                    
                    try:
                        update_progress("å¼€å§‹ç”ŸæˆWordæŠ¥å‘Š...")
                        
                        doc_io = st.session_state.word_generator.create_comprehensive_report(
                            st.session_state.analysis_manager.analysis_results,
                            st.session_state.analysis_manager.current_analysis,
                            include_charts=include_charts,
                            progress_callback=update_progress
                        )
                        
                        # æ¸…ç©ºè¿›åº¦æ˜¾ç¤º
                        progress_placeholder.empty()
                        status_placeholder.empty()
                        
                        # æä¾›ä¸‹è½½
                        st.download_button(
                            label="ğŸ“¥ ä¸‹è½½WordæŠ¥å‘Š",
                            data=doc_io.getvalue(),
                            file_name=f"æœŸè´§åˆ†ææŠ¥å‘Š_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success("âœ… WordæŠ¥å‘Šç”Ÿæˆå®Œæˆï¼")
                        
                        if not include_charts:
                            st.info("ğŸ’¡ æç¤ºï¼šå¿«é€Ÿæ¨¡å¼å·²è·³è¿‡å›¾è¡¨è½¬æ¢ï¼Œå¦‚éœ€æŸ¥çœ‹å›¾è¡¨è¯·é€‰æ‹©å®Œæ•´æ¨¡å¼æˆ–åœ¨ç³»ç»Ÿç•Œé¢æŸ¥çœ‹äº¤äº’å¼å›¾è¡¨ã€‚")
                        
                    except Exception as e:
                        progress_placeholder.empty()
                        status_placeholder.empty()
                        st.error(f"âŒ WordæŠ¥å‘Šç”Ÿæˆå¤±è´¥: {e}")
                        st.exception(e)
            
            # è‡ªåŠ¨åˆ·æ–°ï¼ˆä»…åœ¨åˆ†æè¿è¡Œæ—¶ï¼‰
            if st.session_state.analysis_manager.is_analysis_running():
                time.sleep(1)
                st.rerun()
                
        else:
            st.info("ğŸ‘ˆ è¯·åœ¨åˆ†æé…ç½®é¡µé¢è®¾ç½®å‚æ•°å¹¶å¼€å§‹åˆ†æ")

# ============================================================================
# åº”ç”¨å…¥å£
# ============================================================================

if __name__ == "__main__":
    main()
